"""
HỆ THỐNG PHÂN TÍCH KỸ THUẬT CHỨNG KHOÁN VIỆT NAM
Phiên bản: 3.2 - Đầy đủ 26 chỉ báo có trọng số
- 26 chỉ báo kỹ thuật (chọn từng cái, tất cả đều được tính điểm)
- Dự báo T0-T5, W1-W4, M1-M3
- Hành động dựa trên NHIỀU chỉ báo (không chỉ RSI)
- Hệ thống điểm có trọng số + giải thích chi tiết
- Phân biệt hành động: ĐANG GIỮ vs CHƯA CÓ
- Xuất Excel, Word, Vietstock, MetaStock, ZIP
- Chọn thời gian: Số ngày hoặc Từ ngày - Đến ngày
"""

import os
import time
import warnings
import zipfile
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
import gradio as gr

warnings.filterwarnings('ignore')

# Tắt log vnstock
import logging
logging.getLogger('vnstock').setLevel(logging.CRITICAL)

# ============================================================
# MODULE DỰ BÁO 12 PHƯƠNG PHÁP CHO 26 CHỈ BÁO
# Phiên bản: 1.0
# ============================================================

from scipy import stats
from scipy.signal import find_peaks, argrelextrema
from scipy.fft import fft, fftfreq
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import PolynomialFeatures
import warnings
warnings.filterwarnings('ignore')

# ============================================================
# CẤU HÌNH 26 CHỈ BÁO CHO DỰ BÁO
# ============================================================

INDICATOR_FORECAST_CONFIG = {
    # Nhóm ĐỘNG LƯỢNG (Oscillators) - có bounds cố định
    'RSI': {'bounds': (0, 100), 'type': 'oscillator', 'weight': 8, 'oversold': 30, 'overbought': 70},
    'Stoch_K': {'bounds': (0, 100), 'type': 'oscillator', 'weight': 6, 'oversold': 20, 'overbought': 80},
    'Stoch_D': {'bounds': (0, 100), 'type': 'oscillator', 'weight': 4, 'oversold': 20, 'overbought': 80},
    'StochRSI': {'bounds': (0, 100), 'type': 'oscillator', 'weight': 4, 'oversold': 20, 'overbought': 80},
    'MFI': {'bounds': (0, 100), 'type': 'oscillator', 'weight': 5, 'oversold': 20, 'overbought': 80},
    'Williams_R': {'bounds': (-100, 0), 'type': 'oscillator', 'weight': 4, 'oversold': -80, 'overbought': -20},
    'CCI': {'bounds': (-500, 500), 'type': 'oscillator', 'weight': 4, 'oversold': -100, 'overbought': 100},
    'ADX': {'bounds': (0, 100), 'type': 'trend_strength', 'weight': 6, 'weak': 20, 'strong': 40},
    
    # Nhóm XU HƯỚNG (Trend) - không có bounds cố định
    'MACD': {'bounds': None, 'type': 'trend', 'weight': 8},
    'MACD_Hist': {'bounds': None, 'type': 'momentum', 'weight': 7},
    'MACD_Signal': {'bounds': None, 'type': 'trend', 'weight': 5},
    'ROC': {'bounds': (-50, 50), 'type': 'momentum', 'weight': 3},
    'Momentum': {'bounds': None, 'type': 'momentum', 'weight': 4},
    
    # Nhóm ĐƯỜNG TRUNG BÌNH - giá trị theo giá
    'SMA_5': {'bounds': None, 'type': 'ma', 'weight': 2},
    'SMA_10': {'bounds': None, 'type': 'ma', 'weight': 2},
    'SMA_20': {'bounds': None, 'type': 'ma', 'weight': 3},
    'SMA_50': {'bounds': None, 'type': 'ma', 'weight': 4},
    'SMA_100': {'bounds': None, 'type': 'ma', 'weight': 3},
    'SMA_200': {'bounds': None, 'type': 'ma', 'weight': 5},
    'EMA_12': {'bounds': None, 'type': 'ma', 'weight': 3},
    'EMA_26': {'bounds': None, 'type': 'ma', 'weight': 3},
    'EMA_50': {'bounds': None, 'type': 'ma', 'weight': 3},
    
    # Nhóm KHỐI LƯỢNG
    'OBV': {'bounds': None, 'type': 'volume', 'weight': 5},
    'CMF': {'bounds': (-1, 1), 'type': 'volume', 'weight': 3},
    'FI': {'bounds': None, 'type': 'volume', 'weight': 2},
    
    # Nhóm BIẾN ĐỘNG
    'ATR': {'bounds': None, 'type': 'volatility', 'weight': 5},
    'BB_Upper': {'bounds': None, 'type': 'band', 'weight': 3},
    'BB_Middle': {'bounds': None, 'type': 'band', 'weight': 3},
    'BB_Lower': {'bounds': None, 'type': 'band', 'weight': 3},
}

# Fibonacci Levels
FIBONACCI_LEVELS = [0, 0.236, 0.382, 0.5, 0.618, 0.786, 1.0]

# ============================================================
# 12 PHƯƠNG PHÁP DỰ BÁO
# ============================================================

class IndicatorForecaster:
    """
    Lớp dự báo cho 1 chỉ báo với 12 phương pháp
    """
    
    def __init__(self, indicator_name, values, config=None):
        """
        indicator_name: Tên chỉ báo (RSI, MACD, etc.)
        values: Mảng giá trị lịch sử của chỉ báo
        config: Cấu hình cho chỉ báo
        """
        self.name = indicator_name
        self.values = np.array([v for v in values if pd.notna(v)])
        self.n = len(self.values)
        self.config = config or INDICATOR_FORECAST_CONFIG.get(indicator_name, {})
        self.bounds = self.config.get('bounds')
        self.weight = self.config.get('weight', 1)
        
        # Kết quả từ 12 phương pháp
        self.method_results = {}
        self.forecasts = {}  # T1-T5, W1-W4, M1-M3
        self.explanations = []
        
    def apply_bounds(self, value):
        """Giới hạn giá trị theo bounds của chỉ báo"""
        if self.bounds is None:
            return value
        return max(self.bounds[0], min(self.bounds[1], value))
    
    # ============================================================
    # PHƯƠNG PHÁP 1: LINEAR REGRESSION
    # ============================================================
    
    def method_linear_regression(self, lookback=20):
        """
        Tìm đường thẳng khớp nhất với dữ liệu
        Công dụng: Xác định xu hướng chính
        """
        if self.n < 10:
            return None
            
        lookback = min(lookback, self.n)
        y = self.values[-lookback:]
        X = np.arange(lookback).reshape(-1, 1)
        
        model = LinearRegression()
        model.fit(X, y)
        
        slope = model.coef_[0]
        intercept = model.intercept_
        r_squared = model.score(X, y)
        
        # Dự báo T1-T5
        forecasts = {}
        for i in range(1, 6):
            pred = model.predict([[lookback - 1 + i]])[0]
            forecasts[f'T{i}'] = self.apply_bounds(pred)
        
        # Xác định xu hướng
        if slope > 0.5:
            trend = 'TĂNG MẠNH'
        elif slope > 0.1:
            trend = 'TĂNG'
        elif slope < -0.5:
            trend = 'GIẢM MẠNH'
        elif slope < -0.1:
            trend = 'GIẢM'
        else:
            trend = 'ĐI NGANG'
        
        result = {
            'method': 'Linear Regression',
            'slope': round(slope, 4),
            'intercept': round(intercept, 4),
            'r_squared': round(r_squared, 4),
            'trend': trend,
            'forecasts': forecasts,
            'confidence': min(r_squared * 100, 95),
            'explanation': f"Hồi quy tuyến tính {lookback} ngày: slope={slope:.4f}, R²={r_squared:.2%}. "
                          f"Xu hướng {trend}. {'Khớp tốt' if r_squared > 0.7 else 'Khớp trung bình' if r_squared > 0.4 else 'Khớp yếu'}."
        }
        
        self.method_results['linear'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 2: POLYNOMIAL REGRESSION (BẬC 2)
    # ============================================================
    
    def method_polynomial_regression(self, lookback=20, degree=2):
        """
        Tìm đường cong bậc 2 khớp với dữ liệu
        Công dụng: Phát hiện điểm uốn, đảo chiều
        """
        if self.n < 15:
            return None
            
        lookback = min(lookback, self.n)
        y = self.values[-lookback:]
        X = np.arange(lookback).reshape(-1, 1)
        
        poly = PolynomialFeatures(degree=degree)
        X_poly = poly.fit_transform(X)
        
        model = LinearRegression()
        model.fit(X_poly, y)
        
        coeffs = model.coef_
        r_squared = model.score(X_poly, y)
        
        # Hệ số bậc 2 (a trong ax² + bx + c)
        a = coeffs[2] if len(coeffs) > 2 else 0
        
        # Dự báo T1-T5
        forecasts = {}
        for i in range(1, 6):
            X_pred = poly.transform([[lookback - 1 + i]])
            pred = model.predict(X_pred)[0]
            forecasts[f'T{i}'] = self.apply_bounds(pred)
        
        # Tìm điểm uốn (vertex của parabola)
        # Vertex tại x = -b/(2a)
        b = coeffs[1] if len(coeffs) > 1 else 0
        if abs(a) > 0.0001:
            vertex_x = -b / (2 * a)
            vertex_day = vertex_x - (lookback - 1)  # Số ngày từ hôm nay
            
            if a > 0:
                curvature = 'LÕM (đáy)'
                if 0 < vertex_day < 5:
                    turning_point = f"Có thể TẠO ĐÁY trong T{int(vertex_day)+1}"
                elif -3 < vertex_day <= 0:
                    turning_point = "VỪA TẠO ĐÁY gần đây"
                else:
                    turning_point = None
            else:
                curvature = 'LỒI (đỉnh)'
                if 0 < vertex_day < 5:
                    turning_point = f"Có thể TẠO ĐỈNH trong T{int(vertex_day)+1}"
                elif -3 < vertex_day <= 0:
                    turning_point = "VỪA TẠO ĐỈNH gần đây"
                else:
                    turning_point = None
        else:
            curvature = 'GẦN TUYẾN TÍNH'
            turning_point = None
            vertex_day = None
        
        result = {
            'method': 'Polynomial Regression',
            'degree': degree,
            'coefficients': {'a': round(a, 6), 'b': round(b, 4)},
            'r_squared': round(r_squared, 4),
            'curvature': curvature,
            'vertex_day': round(vertex_day, 1) if vertex_day else None,
            'turning_point': turning_point,
            'forecasts': forecasts,
            'confidence': min(r_squared * 100, 90),
            'explanation': f"Hồi quy đa thức bậc {degree}: a={a:.6f} ({curvature}). "
                          f"{'⚠️ ' + turning_point if turning_point else 'Không phát hiện điểm uốn gần'}. "
                          f"R²={r_squared:.2%}."
        }
        
        self.method_results['polynomial'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 3: ĐẠO HÀM BẬC 1 (VELOCITY)
    # ============================================================
    
    def method_first_derivative(self):
        """
        Tính tốc độ thay đổi (velocity)
        Công dụng: Chỉ báo đang tăng/giảm nhanh cỡ nào
        """
        if self.n < 5:
            return None
        
        # Tính đạo hàm xấp xỉ
        velocity = np.diff(self.values)
        
        # Các giá trị gần đây
        v_current = velocity[-1]
        v_avg_3d = np.mean(velocity[-3:]) if len(velocity) >= 3 else v_current
        v_avg_5d = np.mean(velocity[-5:]) if len(velocity) >= 5 else v_current
        
        # Xác định hướng và tốc độ
        if v_current > 0:
            if v_current > v_avg_5d * 1.5:
                speed = 'TĂNG TỐC NHANH'
            elif v_current > v_avg_5d:
                speed = 'TĂNG TỐC'
            else:
                speed = 'TĂNG ỔN ĐỊNH'
        elif v_current < 0:
            if v_current < v_avg_5d * 1.5:
                speed = 'GIẢM TỐC NHANH'
            elif v_current < v_avg_5d:
                speed = 'GIẢM TỐC'
            else:
                speed = 'GIẢM ỔN ĐỊNH'
        else:
            speed = 'ĐI NGANG'
        
        # Dự báo dựa trên velocity
        current_value = self.values[-1]
        forecasts = {}
        for i in range(1, 6):
            # Velocity có xu hướng giảm dần (mean reversion)
            decay = 0.8 ** i
            pred = current_value + v_avg_3d * i * decay
            forecasts[f'T{i}'] = self.apply_bounds(pred)
        
        result = {
            'method': 'First Derivative (Velocity)',
            'current_velocity': round(v_current, 4),
            'avg_velocity_3d': round(v_avg_3d, 4),
            'avg_velocity_5d': round(v_avg_5d, 4),
            'speed': speed,
            'direction': 'UP' if v_current > 0 else ('DOWN' if v_current < 0 else 'FLAT'),
            'forecasts': forecasts,
            'confidence': 70,
            'explanation': f"Đạo hàm bậc 1: velocity hiện tại={v_current:.4f}, TB 3 ngày={v_avg_3d:.4f}. "
                          f"Tốc độ: {speed}."
        }
        
        self.method_results['velocity'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 4: ĐẠO HÀM BẬC 2 (ACCELERATION)
    # ============================================================
    
    def method_second_derivative(self):
        """
        Tính gia tốc (acceleration)
        Công dụng: Đang tăng tốc hay chậm lại → dự đoán đảo chiều
        """
        if self.n < 7:
            return None
        
        # Đạo hàm bậc 1
        velocity = np.diff(self.values)
        
        # Đạo hàm bậc 2
        acceleration = np.diff(velocity)
        
        a_current = acceleration[-1]
        a_avg_3d = np.mean(acceleration[-3:]) if len(acceleration) >= 3 else a_current
        v_current = velocity[-1]
        
        # Phân tích gia tốc
        if v_current > 0:  # Đang tăng
            if a_current > 0:
                momentum = 'TĂNG TỐC (xu hướng tăng mạnh lên)'
                reversal_signal = False
            else:
                momentum = 'CHẬM LẠI (có thể sắp đảo chiều giảm)'
                reversal_signal = True
        elif v_current < 0:  # Đang giảm
            if a_current < 0:
                momentum = 'GIẢM TỐC (xu hướng giảm mạnh lên)'
                reversal_signal = False
            else:
                momentum = 'BỚT GIẢM (có thể sắp đảo chiều tăng)'
                reversal_signal = True
        else:
            momentum = 'ỔN ĐỊNH'
            reversal_signal = False
        
        # Dự báo
        current_value = self.values[-1]
        forecasts = {}
        for i in range(1, 6):
            # Mô hình vật lý: x = x0 + v*t + 0.5*a*t²
            pred = current_value + v_current * i + 0.5 * a_avg_3d * (i ** 2)
            forecasts[f'T{i}'] = self.apply_bounds(pred)
        
        result = {
            'method': 'Second Derivative (Acceleration)',
            'current_acceleration': round(a_current, 4),
            'avg_acceleration_3d': round(a_avg_3d, 4),
            'current_velocity': round(v_current, 4),
            'momentum': momentum,
            'reversal_signal': reversal_signal,
            'forecasts': forecasts,
            'confidence': 65,
            'explanation': f"Đạo hàm bậc 2: acceleration={a_current:.4f}, velocity={v_current:.4f}. "
                          f"{momentum}. {'⚠️ TÍN HIỆU ĐẢO CHIỀU' if reversal_signal else ''}."
        }
        
        self.method_results['acceleration'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 5: PEAK/TROUGH DETECTION
    # ============================================================
    
    def method_peak_trough_detection(self):
        """
        Tìm đỉnh và đáy lịch sử
        Công dụng: Xác định pha hiện tại, vùng đảo chiều
        """
        if self.n < 20:
            return None
        
        # Tìm đỉnh và đáy
        # Sử dụng scipy.signal.argrelextrema
        order = max(3, self.n // 10)  # Window size
        
        peaks_idx = argrelextrema(self.values, np.greater, order=order)[0]
        troughs_idx = argrelextrema(self.values, np.less, order=order)[0]
        
        peaks = self.values[peaks_idx] if len(peaks_idx) > 0 else []
        troughs = self.values[troughs_idx] if len(troughs_idx) > 0 else []
        
        current_value = self.values[-1]
        
        # Phân tích vùng
        if len(peaks) > 0 and len(troughs) > 0:
            avg_peak = np.mean(peaks)
            avg_trough = np.mean(troughs)
            range_size = avg_peak - avg_trough
            
            # Vị trí hiện tại trong range
            if range_size > 0:
                position_pct = (current_value - avg_trough) / range_size * 100
            else:
                position_pct = 50
            
            # Xác định pha
            if position_pct < 20:
                phase = 'GẦN ĐÁY'
                reversal_prob = 70
            elif position_pct < 40:
                phase = 'VÙNG THẤP'
                reversal_prob = 50
            elif position_pct > 80:
                phase = 'GẦN ĐỈNH'
                reversal_prob = 70
            elif position_pct > 60:
                phase = 'VÙNG CAO'
                reversal_prob = 50
            else:
                phase = 'VÙNG GIỮA'
                reversal_prob = 30
            
            # Kiểm tra chỉ báo có oversold/overbought không
            config = INDICATOR_FORECAST_CONFIG.get(self.name, {})
            oversold = config.get('oversold')
            overbought = config.get('overbought')
            
            zone_signal = None
            if oversold and current_value <= oversold:
                zone_signal = f'QUÁ BÁN (≤{oversold})'
                reversal_prob = 80
            elif overbought and current_value >= overbought:
                zone_signal = f'QUÁ MUA (≥{overbought})'
                reversal_prob = 80
        else:
            avg_peak = current_value * 1.1
            avg_trough = current_value * 0.9
            position_pct = 50
            phase = 'KHÔNG XÁC ĐỊNH'
            reversal_prob = 30
            zone_signal = None
        
        # Dự báo dựa trên pha
        forecasts = {}
        for i in range(1, 6):
            if phase in ['GẦN ĐÁY', 'VÙNG THẤP']:
                # Kỳ vọng tăng
                target = avg_trough + (avg_peak - avg_trough) * 0.3 * i / 5
                pred = current_value + (target - current_value) * i / 5
            elif phase in ['GẦN ĐỈNH', 'VÙNG CAO']:
                # Kỳ vọng giảm
                target = avg_peak - (avg_peak - avg_trough) * 0.3 * i / 5
                pred = current_value + (target - current_value) * i / 5
            else:
                pred = current_value
            forecasts[f'T{i}'] = self.apply_bounds(pred)
        
        result = {
            'method': 'Peak/Trough Detection',
            'num_peaks': len(peaks),
            'num_troughs': len(troughs),
            'avg_peak': round(avg_peak, 2),
            'avg_trough': round(avg_trough, 2),
            'current_value': round(current_value, 2),
            'position_pct': round(position_pct, 1),
            'phase': phase,
            'zone_signal': zone_signal,
            'reversal_probability': reversal_prob,
            'forecasts': forecasts,
            'confidence': reversal_prob,
            'explanation': f"Phát hiện {len(peaks)} đỉnh (TB={avg_peak:.2f}) và {len(troughs)} đáy (TB={avg_trough:.2f}). "
                          f"Hiện tại ở vị trí {position_pct:.0f}% ({phase}). "
                          f"{'⚠️ ' + zone_signal if zone_signal else ''} "
                          f"Xác suất đảo chiều: {reversal_prob}%."
        }
        
        self.method_results['peak_trough'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 6: MULTI-TIMEFRAME ANALYSIS
    # ============================================================
    
    def method_multi_timeframe(self):
        """
        Phân tích xu hướng theo nhiều khung thời gian: 10d, 30d, toàn bộ
        Công dụng: Xác định xu hướng ngắn/trung/dài hạn
        """
        if self.n < 10:
            return None
        
        timeframes = {}
        
        # Ngắn hạn (10 ngày)
        if self.n >= 10:
            short_data = self.values[-10:]
            short_change = (short_data[-1] - short_data[0]) / short_data[0] * 100 if short_data[0] != 0 else 0
            short_trend = 'UP' if short_change > 2 else ('DOWN' if short_change < -2 else 'FLAT')
            timeframes['short_10d'] = {
                'change_pct': round(short_change, 2),
                'trend': short_trend,
                'start': round(short_data[0], 2),
                'end': round(short_data[-1], 2)
            }
        
        # Trung hạn (30 ngày)
        if self.n >= 30:
            mid_data = self.values[-30:]
            mid_change = (mid_data[-1] - mid_data[0]) / mid_data[0] * 100 if mid_data[0] != 0 else 0
            mid_trend = 'UP' if mid_change > 5 else ('DOWN' if mid_change < -5 else 'FLAT')
            timeframes['mid_30d'] = {
                'change_pct': round(mid_change, 2),
                'trend': mid_trend,
                'start': round(mid_data[0], 2),
                'end': round(mid_data[-1], 2)
            }
        
        # Dài hạn (toàn bộ)
        long_data = self.values
        long_change = (long_data[-1] - long_data[0]) / long_data[0] * 100 if long_data[0] != 0 else 0
        long_trend = 'UP' if long_change > 10 else ('DOWN' if long_change < -10 else 'FLAT')
        timeframes['long_all'] = {
            'change_pct': round(long_change, 2),
            'trend': long_trend,
            'start': round(long_data[0], 2),
            'end': round(long_data[-1], 2),
            'days': self.n
        }
        
        # Phân tích divergence giữa các khung thời gian
        trends = [tf.get('trend') for tf in timeframes.values()]
        
        if all(t == 'UP' for t in trends):
            overall = 'ĐỒNG THUẬN TĂNG'
            confidence = 85
        elif all(t == 'DOWN' for t in trends):
            overall = 'ĐỒNG THUẬN GIẢM'
            confidence = 85
        elif trends[0] != trends[-1]:  # Ngắn hạn khác dài hạn
            overall = 'PHÂN KỲ (ngắn vs dài hạn)'
            confidence = 60
        else:
            overall = 'HỖN HỢP'
            confidence = 50
        
        # Dự báo
        current_value = self.values[-1]
        short_tf = timeframes.get('short_10d', {})
        forecasts = {}
        
        # Dùng xu hướng ngắn hạn làm chính
        daily_change = short_tf.get('change_pct', 0) / 10 / 100  # % change per day
        
        for i in range(1, 6):
            pred = current_value * (1 + daily_change * i)
            forecasts[f'T{i}'] = self.apply_bounds(pred)
        
        result = {
            'method': 'Multi-Timeframe Analysis',
            'timeframes': timeframes,
            'overall_trend': overall,
            'divergence': trends[0] != trends[-1] if len(trends) > 1 else False,
            'forecasts': forecasts,
            'confidence': confidence,
            'explanation': f"Phân tích đa khung thời gian: "
                          f"10d={timeframes.get('short_10d', {}).get('trend', 'N/A')}, "
                          f"30d={timeframes.get('mid_30d', {}).get('trend', 'N/A')}, "
                          f"All={timeframes.get('long_all', {}).get('trend', 'N/A')}. "
                          f"Kết luận: {overall}."
        }
        
        self.method_results['multi_timeframe'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 7: PATTERN MATCHING
    # ============================================================
    
    def method_pattern_matching(self, pattern_length=5, tolerance_pct=10):
        """
        So sánh mô hình hiện tại với lịch sử
        Công dụng: Tìm các tình huống tương tự trong quá khứ để dự đoán
        """
        if self.n < 30:
            return None
        
        # Lấy pattern hiện tại (5 ngày gần nhất)
        current_pattern = self.values[-pattern_length:]
        
        # Chuẩn hóa pattern (so sánh hình dạng, không phải giá trị tuyệt đối)
        def normalize_pattern(pattern):
            min_val = np.min(pattern)
            max_val = np.max(pattern)
            if max_val - min_val > 0:
                return (pattern - min_val) / (max_val - min_val)
            return np.zeros_like(pattern)
        
        current_norm = normalize_pattern(current_pattern)
        
        # Tìm các pattern tương tự trong lịch sử
        similar_patterns = []
        min_similarity = 1 - tolerance_pct / 100
        
        for i in range(pattern_length, self.n - pattern_length - 5):
            hist_pattern = self.values[i:i+pattern_length]
            hist_norm = normalize_pattern(hist_pattern)
            
            # Tính correlation
            if np.std(hist_norm) > 0 and np.std(current_norm) > 0:
                correlation = np.corrcoef(current_norm, hist_norm)[0, 1]
            else:
                correlation = 0
            
            if correlation >= min_similarity:
                # Lấy kết quả 5 ngày sau pattern đó
                future_values = self.values[i+pattern_length:i+pattern_length+5]
                if len(future_values) == 5:
                    start_val = self.values[i+pattern_length-1]
                    changes = [(v - start_val) / start_val * 100 if start_val != 0 else 0 for v in future_values]
                    
                    similar_patterns.append({
                        'index': i,
                        'correlation': correlation,
                        'future_changes': changes,
                        'outcome': 'UP' if changes[-1] > 2 else ('DOWN' if changes[-1] < -2 else 'FLAT')
                    })
        
        if not similar_patterns:
            result = {
                'method': 'Pattern Matching',
                'matches_found': 0,
                'forecasts': {f'T{i}': self.values[-1] for i in range(1, 6)},
                'confidence': 30,
                'explanation': f"Không tìm thấy mẫu tương tự với độ tương quan ≥{min_similarity:.0%} trong {self.n} ngày lịch sử."
            }
            self.method_results['pattern'] = result
            return result
        
        # Sắp xếp theo correlation
        similar_patterns.sort(key=lambda x: x['correlation'], reverse=True)
        top_matches = similar_patterns[:5]
        
        # Tính dự báo trung bình từ các match
        avg_changes = [0] * 5
        for match in top_matches:
            for i, change in enumerate(match['future_changes']):
                avg_changes[i] += change / len(top_matches)
        
        current_value = self.values[-1]
        forecasts = {}
        for i in range(5):
            pred = current_value * (1 + avg_changes[i] / 100)
            forecasts[f'T{i+1}'] = self.apply_bounds(pred)
        
        # Thống kê outcome
        outcomes = [m['outcome'] for m in top_matches]
        up_count = outcomes.count('UP')
        down_count = outcomes.count('DOWN')
        
        if up_count > down_count:
            prediction = 'TĂNG'
            prob = up_count / len(outcomes) * 100
        elif down_count > up_count:
            prediction = 'GIẢM'
            prob = down_count / len(outcomes) * 100
        else:
            prediction = 'KHÔNG RÕ'
            prob = 50
        
        result = {
            'method': 'Pattern Matching',
            'matches_found': len(similar_patterns),
            'top_matches': len(top_matches),
            'avg_correlation': round(np.mean([m['correlation'] for m in top_matches]), 3),
            'prediction': prediction,
            'prediction_probability': round(prob, 1),
            'avg_changes': [round(c, 2) for c in avg_changes],
            'forecasts': forecasts,
            'confidence': min(prob, 80),
            'explanation': f"Tìm thấy {len(similar_patterns)} mẫu tương tự (correlation ≥{min_similarity:.0%}). "
                          f"Top {len(top_matches)} mẫu: {up_count} tăng, {down_count} giảm. "
                          f"Dự đoán: {prediction} (xác suất {prob:.0f}%)."
        }
        
        self.method_results['pattern'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 8: FOURIER TRANSFORM
    # ============================================================
    
    def method_fourier_transform(self):
        """
        Phân tích chu kỳ bằng biến đổi Fourier
        Công dụng: Phát hiện chu kỳ lặp lại, dự đoán đỉnh/đáy theo chu kỳ
        """
        if self.n < 30:
            return None
        
        # Loại bỏ trend trước khi FFT
        detrended = self.values - np.linspace(self.values[0], self.values[-1], self.n)
        
        # FFT
        fft_values = fft(detrended)
        frequencies = fftfreq(self.n)
        
        # Lấy magnitude
        magnitudes = np.abs(fft_values)
        
        # Tìm các tần số chính (bỏ qua DC component)
        # Chỉ xét nửa đầu (symmetric)
        half_n = self.n // 2
        mag_half = magnitudes[1:half_n]
        freq_half = frequencies[1:half_n]
        
        # Tìm top 3 peaks
        if len(mag_half) > 3:
            top_indices = np.argsort(mag_half)[-3:][::-1]
            dominant_freqs = freq_half[top_indices]
            dominant_mags = mag_half[top_indices]
            
            # Chuyển frequency sang period (ngày)
            periods = [1 / abs(f) if f != 0 else self.n for f in dominant_freqs]
        else:
            periods = [self.n]
            dominant_mags = [0]
        
        # Tìm chu kỳ chính (period có magnitude lớn nhất và hợp lý)
        valid_periods = [(p, m) for p, m in zip(periods, dominant_mags) if 5 <= p <= self.n / 2]
        
        if valid_periods:
            main_period = valid_periods[0][0]
            main_magnitude = valid_periods[0][1]
            
            # Xác định vị trí trong chu kỳ
            cycle_position = self.n % main_period
            cycle_phase = cycle_position / main_period * 360  # Độ
            
            # Dự đoán đỉnh/đáy
            if 0 <= cycle_phase < 90:
                cycle_stage = 'ĐANG TĂNG (đầu chu kỳ)'
            elif 90 <= cycle_phase < 180:
                cycle_stage = 'GẦN ĐỈNH (giữa chu kỳ)'
            elif 180 <= cycle_phase < 270:
                cycle_stage = 'ĐANG GIẢM (sau đỉnh)'
            else:
                cycle_stage = 'GẦN ĐÁY (cuối chu kỳ)'
            
            # Dự đoán ngày đến đỉnh/đáy tiếp theo
            days_to_peak = (180 - cycle_phase) / 360 * main_period if cycle_phase < 180 else (540 - cycle_phase) / 360 * main_period
            days_to_trough = (360 - cycle_phase) / 360 * main_period if cycle_phase < 360 else main_period - (cycle_phase - 360) / 360 * main_period
            
            cycle_strength = main_magnitude / np.mean(magnitudes) if np.mean(magnitudes) > 0 else 0
            confidence = min(40 + cycle_strength * 10, 75)
        else:
            main_period = None
            cycle_stage = 'KHÔNG RÕ CHU KỲ'
            days_to_peak = None
            days_to_trough = None
            confidence = 30
            cycle_strength = 0
        
        # Dự báo đơn giản
        current_value = self.values[-1]
        forecasts = {}
        
        if main_period and main_period > 0:
            # Dùng sine wave để dự đoán
            amplitude = np.std(detrended)
            trend_slope = (self.values[-1] - self.values[0]) / self.n
            
            for i in range(1, 6):
                # Trend + Cycle
                trend_component = trend_slope * i
                cycle_component = amplitude * np.sin(2 * np.pi * (cycle_position + i) / main_period)
                pred = current_value + trend_component + cycle_component * 0.5
                forecasts[f'T{i}'] = self.apply_bounds(pred)
        else:
            for i in range(1, 6):
                forecasts[f'T{i}'] = current_value
        
        result = {
            'method': 'Fourier Transform',
            'main_period': round(main_period, 1) if main_period else None,
            'all_periods': [round(p, 1) for p in periods[:3]],
            'cycle_stage': cycle_stage,
            'cycle_strength': round(cycle_strength, 2),
            'days_to_peak': round(days_to_peak, 0) if days_to_peak else None,
            'days_to_trough': round(days_to_trough, 0) if days_to_trough else None,
            'forecasts': forecasts,
            'confidence': confidence,
            'explanation': f"Phân tích Fourier: "
                          f"{'Chu kỳ chính ' + str(round(main_period, 0)) + ' ngày' if main_period else 'Không phát hiện chu kỳ rõ ràng'}. "
                          f"{cycle_stage}. "
                          f"{'Đỉnh trong ~' + str(round(days_to_peak, 0)) + ' ngày' if days_to_peak and days_to_peak < 10 else ''} "
                          f"{'Đáy trong ~' + str(round(days_to_trough, 0)) + ' ngày' if days_to_trough and days_to_trough < 10 else ''}."
        }
        
        self.method_results['fourier'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 9: PROBABILITY & STATISTICS
    # ============================================================
    
    def method_probability_statistics(self):
        """
        Phân tích xác suất và thống kê
        Công dụng: Tính xác suất giá trị nằm trong vùng nào, mean reversion
        """
        if self.n < 20:
            return None
        
        current_value = self.values[-1]
        
        # Thống kê cơ bản
        mean = np.mean(self.values)
        std = np.std(self.values)
        median = np.median(self.values)
        min_val = np.min(self.values)
        max_val = np.max(self.values)
        
        # Z-score
        z_score = (current_value - mean) / std if std > 0 else 0
        
        # Percentile
        percentile = stats.percentileofscore(self.values, current_value)
        
        # Xác suất quay về mean (mean reversion)
        if z_score > 2:
            mean_reversion_prob = 85
            direction = 'GIẢM VỀ MEAN'
        elif z_score > 1:
            mean_reversion_prob = 70
            direction = 'CÓ THỂ GIẢM'
        elif z_score < -2:
            mean_reversion_prob = 85
            direction = 'TĂNG VỀ MEAN'
        elif z_score < -1:
            mean_reversion_prob = 70
            direction = 'CÓ THỂ TĂNG'
        else:
            mean_reversion_prob = 50
            direction = 'GẦN MEAN'
        
        # Bollinger-like bands
        upper_2std = mean + 2 * std
        lower_2std = mean - 2 * std
        upper_1std = mean + std
        lower_1std = mean - std
        
        # Vị trí trong bands
        if current_value > upper_2std:
            band_position = 'TRÊN 2σ (cực kỳ cao)'
        elif current_value > upper_1std:
            band_position = 'TRÊN 1σ (cao)'
        elif current_value < lower_2std:
            band_position = 'DƯỚI 2σ (cực kỳ thấp)'
        elif current_value < lower_1std:
            band_position = 'DƯỚI 1σ (thấp)'
        else:
            band_position = 'TRONG 1σ (bình thường)'
        
        # Dự báo (mean reversion)
        forecasts = {}
        reversion_speed = 0.2  # 20% về mean mỗi ngày
        
        for i in range(1, 6):
            # Dự báo dựa trên mean reversion
            distance_to_mean = current_value - mean
            pred = current_value - distance_to_mean * (1 - (1 - reversion_speed) ** i)
            forecasts[f'T{i}'] = self.apply_bounds(pred)
        
        result = {
            'method': 'Probability & Statistics',
            'mean': round(mean, 2),
            'std': round(std, 2),
            'median': round(median, 2),
            'min': round(min_val, 2),
            'max': round(max_val, 2),
            'current_value': round(current_value, 2),
            'z_score': round(z_score, 2),
            'percentile': round(percentile, 1),
            'band_position': band_position,
            'mean_reversion_prob': mean_reversion_prob,
            'direction': direction,
            'bands': {
                'upper_2std': round(upper_2std, 2),
                'upper_1std': round(upper_1std, 2),
                'mean': round(mean, 2),
                'lower_1std': round(lower_1std, 2),
                'lower_2std': round(lower_2std, 2)
            },
            'forecasts': forecasts,
            'confidence': mean_reversion_prob,
            'explanation': f"Thống kê: Mean={mean:.2f}, Std={std:.2f}. "
                          f"Hiện tại={current_value:.2f} (Z={z_score:.2f}, Percentile={percentile:.0f}%). "
                          f"{band_position}. Dự đoán: {direction} với xác suất {mean_reversion_prob}%."
        }
        
        self.method_results['statistics'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 10: FIBONACCI LEVELS
    # ============================================================
    
    def method_fibonacci_levels(self):
        """
        Tính các mức Fibonacci retracement
        Công dụng: Xác định vùng hỗ trợ/kháng cự tự nhiên
        """
        if self.n < 20:
            return None
        
        # Tìm high/low trong period
        high = np.max(self.values)
        low = np.min(self.values)
        diff = high - low
        
        current_value = self.values[-1]
        
        # Tính các mức Fibonacci
        fib_levels = {}
        for level in FIBONACCI_LEVELS:
            fib_levels[f'{level*100:.1f}%'] = low + diff * (1 - level)
        
        # Xác định vị trí hiện tại
        current_fib = (high - current_value) / diff * 100 if diff > 0 else 50
        
        # Tìm mức gần nhất
        nearest_support = None
        nearest_resistance = None
        
        for level_name, level_value in fib_levels.items():
            if level_value < current_value:
                if nearest_support is None or level_value > nearest_support[1]:
                    nearest_support = (level_name, level_value)
            elif level_value > current_value:
                if nearest_resistance is None or level_value < nearest_resistance[1]:
                    nearest_resistance = (level_name, level_value)
        
        # Đánh giá vùng
        if current_fib > 78.6:
            zone = 'GẦN ĐÁY (dưới 23.6%)'
            bias = 'BULLISH'
        elif current_fib > 61.8:
            zone = 'VÙNG HỖ TRỢ MẠNH (23.6%-38.2%)'
            bias = 'BULLISH'
        elif current_fib > 50:
            zone = 'VÙNG GOLDEN RATIO (38.2%-50%)'
            bias = 'NEUTRAL'
        elif current_fib > 38.2:
            zone = 'VÙNG KHÁNG CỰ (50%-61.8%)'
            bias = 'BEARISH'
        else:
            zone = 'GẦN ĐỈNH (trên 61.8%)'
            bias = 'BEARISH'
        
        # Dự báo dựa trên Fibonacci
        forecasts = {}
        
        if bias == 'BULLISH' and nearest_resistance:
            target = nearest_resistance[1]
            for i in range(1, 6):
                pred = current_value + (target - current_value) * i / 7
                forecasts[f'T{i}'] = self.apply_bounds(pred)
        elif bias == 'BEARISH' and nearest_support:
            target = nearest_support[1]
            for i in range(1, 6):
                pred = current_value - (current_value - target) * i / 7
                forecasts[f'T{i}'] = self.apply_bounds(pred)
        else:
            for i in range(1, 6):
                forecasts[f'T{i}'] = current_value
        
        result = {
            'method': 'Fibonacci Levels',
            'high': round(high, 2),
            'low': round(low, 2),
            'current_value': round(current_value, 2),
            'current_fib_pct': round(current_fib, 1),
            'fib_levels': {k: round(v, 2) for k, v in fib_levels.items()},
            'nearest_support': (nearest_support[0], round(nearest_support[1], 2)) if nearest_support else None,
            'nearest_resistance': (nearest_resistance[0], round(nearest_resistance[1], 2)) if nearest_resistance else None,
            'zone': zone,
            'bias': bias,
            'forecasts': forecasts,
            'confidence': 65,
            'explanation': f"Fibonacci: High={high:.2f}, Low={low:.2f}. "
                          f"Hiện tại ở mức {current_fib:.1f}% ({zone}). "
                          f"Hỗ trợ gần: {nearest_support[0] if nearest_support else 'N/A'}, "
                          f"Kháng cự gần: {nearest_resistance[0] if nearest_resistance else 'N/A'}. "
                          f"Xu hướng: {bias}."
        }
        
        self.method_results['fibonacci'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 11: LOGICAL RULES ENGINE
    # ============================================================
    
    def method_logical_rules(self):
        """
        Áp dụng các quy tắc logic dựa trên đặc tính chỉ báo
        Công dụng: Kết hợp nhiều điều kiện để đưa ra quyết định
        """
        if self.n < 10:
            return None
        
        current_value = self.values[-1]
        config = INDICATOR_FORECAST_CONFIG.get(self.name, {})
        
        rules_triggered = []
        buy_score = 0
        sell_score = 0
        
        # ===== RULE 1: Oversold/Overbought =====
        oversold = config.get('oversold')
        overbought = config.get('overbought')
        
        if oversold and current_value <= oversold:
            rules_triggered.append(f"✅ QUÁ BÁN: {self.name}={current_value:.1f} ≤ {oversold}")
            buy_score += 3
        elif overbought and current_value >= overbought:
            rules_triggered.append(f"⚠️ QUÁ MUA: {self.name}={current_value:.1f} ≥ {overbought}")
            sell_score += 3
        
        # ===== RULE 2: Trend ngắn hạn (3 ngày) =====
        if self.n >= 3:
            trend_3d = self.values[-1] - self.values[-3]
            if trend_3d > 0:
                rules_triggered.append(f"📈 Trend 3d: +{trend_3d:.2f}")
                buy_score += 1
            else:
                rules_triggered.append(f"📉 Trend 3d: {trend_3d:.2f}")
                sell_score += 1
        
        # ===== RULE 3: Trend trung hạn (10 ngày) =====
        if self.n >= 10:
            trend_10d = self.values[-1] - self.values[-10]
            if trend_10d > 0:
                rules_triggered.append(f"📈 Trend 10d: +{trend_10d:.2f}")
                buy_score += 1
            else:
                rules_triggered.append(f"📉 Trend 10d: {trend_10d:.2f}")
                sell_score += 1
        
        # ===== RULE 4: Vị trí so với trung bình =====
        mean_20d = np.mean(self.values[-20:]) if self.n >= 20 else np.mean(self.values)
        if current_value < mean_20d * 0.95:
            rules_triggered.append(f"📉 Dưới MA20 ({current_value:.1f} < {mean_20d:.1f})")
            buy_score += 1  # Có thể là cơ hội mua
        elif current_value > mean_20d * 1.05:
            rules_triggered.append(f"📈 Trên MA20 ({current_value:.1f} > {mean_20d:.1f})")
            sell_score += 1  # Có thể là cơ hội bán
        
        # ===== RULE 5: Velocity (tốc độ thay đổi) =====
        if self.n >= 5:
            velocity = np.mean(np.diff(self.values[-5:]))
            if velocity > 0:
                rules_triggered.append(f"⚡ Velocity dương: +{velocity:.3f}/ngày")
                buy_score += 1
            else:
                rules_triggered.append(f"⚡ Velocity âm: {velocity:.3f}/ngày")
                sell_score += 1
        
        # ===== RULE 6: Đảo chiều gần đây =====
        if self.n >= 5:
            recent = self.values[-5:]
            if recent[-1] > recent[-2] and recent[-2] < recent[-3]:
                rules_triggered.append("🔄 Đảo chiều TĂNG (V-bottom)")
                buy_score += 2
            elif recent[-1] < recent[-2] and recent[-2] > recent[-3]:
                rules_triggered.append("🔄 Đảo chiều GIẢM (Λ-top)")
                sell_score += 2
        
        # ===== TỔNG HỢP =====
        total_score = buy_score - sell_score
        
        if total_score >= 4:
            signal = 'MUA MẠNH'
            confidence = 85
        elif total_score >= 2:
            signal = 'MUA'
            confidence = 70
        elif total_score <= -4:
            signal = 'BÁN MẠNH'
            confidence = 85
        elif total_score <= -2:
            signal = 'BÁN'
            confidence = 70
        else:
            signal = 'TRUNG LẬP'
            confidence = 50
        
        # Dự báo
        forecasts = {}
        bias = total_score / 10  # Normalize to -1 to 1
        
        for i in range(1, 6):
            change_pct = bias * 2 * i  # Max ±10% in 5 days
            pred = current_value * (1 + change_pct / 100)
            forecasts[f'T{i}'] = self.apply_bounds(pred)
        
        result = {
            'method': 'Logical Rules Engine',
            'rules_triggered': rules_triggered,
            'buy_score': buy_score,
            'sell_score': sell_score,
            'total_score': total_score,
            'signal': signal,
            'forecasts': forecasts,
            'confidence': confidence,
            'explanation': f"Logic Rules: Buy={buy_score}, Sell={sell_score}, Net={total_score}. "
                          f"Kết luận: {signal}. "
                          f"Các quy tắc: {'; '.join(rules_triggered[:3])}..."
        }
        
        self.method_results['logical'] = result
        return result
    
    # ============================================================
    # PHƯƠNG PHÁP 12: ML ENSEMBLE (Simplified)
    # ============================================================
    
    def method_ml_ensemble(self):
        """
        Kết hợp nhiều mô hình đơn giản bằng voting
        Công dụng: Tăng độ chính xác bằng cách kết hợp nhiều phương pháp
        """
        # Chạy tất cả các phương pháp khác trước
        if not self.method_results:
            self.run_all_methods()
        
        if len(self.method_results) < 3:
            return None
        
        # Thu thập forecasts từ các phương pháp
        all_forecasts = {f'T{i}': [] for i in range(1, 6)}
        confidences = []
        trends = []
        
        for method_name, result in self.method_results.items():
            if method_name == 'ensemble':
                continue
            
            if result and 'forecasts' in result:
                for key, value in result['forecasts'].items():
                    if key in all_forecasts:
                        all_forecasts[key].append(value)
                
                confidences.append(result.get('confidence', 50))
                
                # Thu thập xu hướng
                if 'trend' in result:
                    trends.append(result['trend'])
                elif 'signal' in result:
                    trends.append(result['signal'])
                elif 'direction' in result:
                    trends.append(result['direction'])
        
        # Tính forecast tổng hợp (weighted average by confidence)
        ensemble_forecasts = {}
        
        for key, values in all_forecasts.items():
            if values:
                # Simple average
                ensemble_forecasts[key] = self.apply_bounds(np.mean(values))
        
        # Voting cho xu hướng
        up_votes = sum(1 for t in trends if t and any(x in str(t).upper() for x in ['UP', 'TĂNG', 'MUA', 'BULLISH']))
        down_votes = sum(1 for t in trends if t and any(x in str(t).upper() for x in ['DOWN', 'GIẢM', 'BÁN', 'BEARISH']))
        total_votes = len(trends)
        
        if total_votes > 0:
            up_pct = up_votes / total_votes * 100
            down_pct = down_votes / total_votes * 100
        else:
            up_pct = down_pct = 50
        
        if up_pct > 60:
            ensemble_trend = 'TĂNG'
            ensemble_confidence = min(up_pct, 90)
        elif down_pct > 60:
            ensemble_trend = 'GIẢM'
            ensemble_confidence = min(down_pct, 90)
        else:
            ensemble_trend = 'TRUNG LẬP'
            ensemble_confidence = 50
        
        result = {
            'method': 'ML Ensemble',
            'methods_used': len(self.method_results) - 1,
            'avg_confidence': round(np.mean(confidences), 1) if confidences else 50,
            'voting': {
                'up_votes': up_votes,
                'down_votes': down_votes,
                'total_votes': total_votes,
                'up_pct': round(up_pct, 1),
                'down_pct': round(down_pct, 1)
            },
            'ensemble_trend': ensemble_trend,
            'forecasts': ensemble_forecasts,
            'confidence': ensemble_confidence,
            'explanation': f"Ensemble {len(self.method_results)-1} phương pháp: "
                          f"Voting {up_votes} tăng / {down_votes} giảm ({up_pct:.0f}% / {down_pct:.0f}%). "
                          f"Kết luận: {ensemble_trend} (confidence {ensemble_confidence:.0f}%)."
        }
        
        self.method_results['ensemble'] = result
        return result
    
    # ============================================================
    # CHẠY TẤT CẢ PHƯƠNG PHÁP
    # ============================================================
    
    def run_all_methods(self):
        """Chạy tất cả 12 phương pháp và tổng hợp kết quả"""
        
        # 1-6: Các phương pháp cơ bản
        self.method_linear_regression()
        self.method_polynomial_regression()
        self.method_first_derivative()
        self.method_second_derivative()
        self.method_peak_trough_detection()
        self.method_multi_timeframe()
        
        # 7-10: Các phương pháp nâng cao
        self.method_pattern_matching()
        self.method_fourier_transform()
        self.method_probability_statistics()
        self.method_fibonacci_levels()
        
        # 11: Logic rules
        self.method_logical_rules()
        
        # 12: Ensemble (phải chạy cuối cùng)
        self.method_ml_ensemble()
        
        return self.method_results
    
    # ============================================================
    # TỔNG HỢP KẾT QUẢ
    # ============================================================
    
    def get_final_forecast(self):
        """Lấy kết quả dự báo cuối cùng từ ensemble"""
        if 'ensemble' not in self.method_results:
            self.run_all_methods()
        
        ensemble = self.method_results.get('ensemble', {})
        
        return {
            'indicator': self.name,
            'current_value': round(self.values[-1], 2) if len(self.values) > 0 else None,
            'forecasts': ensemble.get('forecasts', {}),
            'trend': ensemble.get('ensemble_trend', 'N/A'),
            'confidence': ensemble.get('confidence', 50),
            'methods_used': ensemble.get('methods_used', 0),
            'voting': ensemble.get('voting', {}),
            'weight': self.weight
        }
    
    def get_detailed_report(self):
        """Tạo báo cáo chi tiết với giải thích từng phương pháp"""
        if not self.method_results:
            self.run_all_methods()
        
        report = {
            'indicator': self.name,
            'current_value': round(self.values[-1], 2) if len(self.values) > 0 else None,
            'data_points': self.n,
            'weight': self.weight,
            'methods': {}
        }
        
        for method_name, result in self.method_results.items():
            if result:
                report['methods'][method_name] = {
                    'confidence': result.get('confidence', 0),
                    'explanation': result.get('explanation', ''),
                    'forecasts': result.get('forecasts', {})
                }
        
        # Thêm final forecast
        report['final'] = self.get_final_forecast()
        
        return report


# ============================================================
# HÀM TỔNG HỢP 26 BỘ DỰ BÁO
# ============================================================

def get_indicator_type(indicator_name):
    """Xác định loại chỉ báo để xử lý phù hợp"""
    
    # Oscillators (0-100 hoặc có bounds cố định)
    oscillators = ['RSI', 'Stoch_K', 'Stoch_D', 'StochRSI', 'MFI', 'Williams_R', 'CCI', 'ADX']
    
    # Giá trị nhỏ, có thể âm/dương
    small_values = ['MACD', 'MACD_Hist', 'MACD_Signal', 'ROC', 'CMF']
    
    # Giá trị lớn (theo giá hoặc khối lượng)
    large_values = ['OBV', 'AD', 'FI', 'Momentum', 'ATR', 'SMA_5', 'SMA_10', 'SMA_20', 
                    'SMA_50', 'SMA_100', 'SMA_200', 'EMA_12', 'EMA_26', 'EMA_50',
                    'BB_Upper', 'BB_Middle', 'BB_Lower', 'VWAP', 'WMA_10', 'WMA_20',
                    'TEMA_20', 'DEMA_20', 'SAR']
    
    if indicator_name in oscillators:
        return 'oscillator'
    elif indicator_name in small_values:
        return 'small'
    elif indicator_name in large_values:
        return 'large'
    else:
        return 'unknown'


def normalize_indicator_forecast(indicator_name, current_value, forecast_value, indicator_type=None):
    """
    Chuẩn hóa dự báo của chỉ báo thành điểm số 0-100
    Để có thể so sánh và tổng hợp các chỉ báo khác nhau
    """
    
    if indicator_type is None:
        indicator_type = get_indicator_type(indicator_name)
    
    if current_value is None or forecast_value is None:
        return 50  # Trung lập
    
    try:
        current_value = float(current_value)
        forecast_value = float(forecast_value)
    except (ValueError, TypeError):
        return 50
    
    # Tính % thay đổi
    if current_value != 0:
        change_pct = (forecast_value - current_value) / abs(current_value) * 100
    else:
        change_pct = 0
    
    # Với oscillators, dùng giá trị trực tiếp (đã 0-100)
    if indicator_type == 'oscillator':
        config = INDICATOR_FORECAST_CONFIG.get(indicator_name, {})
        oversold = config.get('oversold', 30)
        overbought = config.get('overbought', 70)
        
        # Nếu dự báo từ oversold lên → bullish
        # Nếu dự báo từ overbought xuống → bearish
        if current_value <= oversold:
            if forecast_value > current_value:
                return min(50 + (forecast_value - current_value) * 2, 90)
            else:
                return max(50 - (current_value - forecast_value), 20)
        elif current_value >= overbought:
            if forecast_value < current_value:
                return min(50 + (current_value - forecast_value) * 2, 90)
            else:
                return max(50 - (forecast_value - current_value), 20)
        else:
            # Vùng trung lập
            return 50 + change_pct
    
    # Với các loại khác, dùng % thay đổi
    else:
        # Giới hạn điểm trong 0-100
        score = 50 + change_pct * 5  # Mỗi 1% thay đổi = 5 điểm
        return max(0, min(100, score))


def forecast_all_26_indicators_v2(df, selected_indicators=None):
    """
    PHIÊN BẢN 2.0: Dự báo cho tất cả chỉ báo có trong DataFrame
    Tự động dò tìm các cột chỉ báo, không mapping cứng
    
    Args:
        df: DataFrame chứa dữ liệu với các cột chỉ báo đã tính
        selected_indicators: Danh sách chỉ báo muốn dự báo (None = tất cả có sẵn)
    
    Returns:
        Dict chứa kết quả dự báo cho từng chỉ báo và tổng hợp
    """
    
    # Danh sách tất cả các tên cột chỉ báo có thể có
    all_possible_indicators = [
        # Oscillators
        'RSI', 'Stoch_K', 'Stoch_D', 'StochRSI', 'MFI', 'Williams_R', 'CCI', 'ADX',
        'Plus_DI', 'Minus_DI',
        # MACD
        'MACD', 'MACD_Hist', 'MACD_Signal',
        # Momentum
        'ROC', 'Momentum',
        # Moving Averages
        'SMA_5', 'SMA_10', 'SMA_20', 'SMA_50', 'SMA_100', 'SMA_200',
        'EMA_12', 'EMA_26', 'EMA_50',
        'WMA_10', 'WMA_20', 'TEMA_20', 'DEMA_20',
        # Volume
        'OBV', 'CMF', 'FI', 'AD', 'VWAP',
        # Volatility & Bands
        'ATR', 'BB_Upper', 'BB_Middle', 'BB_Lower',
        # Other
        'SAR'
    ]
    
    results = {
        'individual_forecasts': {},
        'normalized_scores': {},  # Điểm chuẩn hóa 0-100
        'combined_forecast': {},
        'summary': {},
        'explanations': [],
        'indicators_found': [],
        'indicators_analyzed': 0
    }
    
    # TỰ ĐỘNG DÒ TÌM các chỉ báo có trong DataFrame
    available_indicators = []
    for ind in all_possible_indicators:
        if ind in df.columns:
            # Kiểm tra có dữ liệu không (không toàn NaN)
            if df[ind].notna().sum() >= 10:  # Cần ít nhất 10 điểm dữ liệu
                available_indicators.append(ind)
    
    results['indicators_found'] = available_indicators
    
    # Nếu có selected_indicators, chỉ lấy những cái có trong available
    if selected_indicators:
        indicators_to_analyze = [ind for ind in available_indicators if ind in selected_indicators]
    else:
        indicators_to_analyze = available_indicators
    
    if not indicators_to_analyze:
        results['summary'] = {
            'overall_trend': 'KHÔNG XÁC ĐỊNH',
            'overall_confidence': 0,
            'indicators_analyzed': 0,
            'error': 'Không tìm thấy chỉ báo nào có đủ dữ liệu'
        }
        return results
    
    # Dự báo cho từng chỉ báo
    all_normalized_scores = {f'T{i}': [] for i in range(1, 6)}
    all_weights = {f'T{i}': [] for i in range(1, 6)}
    trend_votes = {'UP': 0, 'DOWN': 0, 'NEUTRAL': 0}
    total_weight = 0
    
    for ind_name in indicators_to_analyze:
        try:
            # Lấy dữ liệu chỉ báo
            values = df[ind_name].dropna().values
            
            if len(values) < 10:
                continue
            
            # Lấy cấu hình
            config = INDICATOR_FORECAST_CONFIG.get(ind_name, {})
            weight = config.get('weight', 1)
            ind_type = get_indicator_type(ind_name)
            
            # Tạo forecaster và chạy
            forecaster = IndicatorForecaster(ind_name, values, config)
            forecaster.run_all_methods()
            
            final = forecaster.get_final_forecast()
            detailed = forecaster.get_detailed_report()
            
            # Lưu kết quả gốc
            results['individual_forecasts'][ind_name] = {
                'final': final,
                'detailed': detailed,
                'indicator_type': ind_type
            }
            
            # CHUẨN HÓA điểm số
            current_value = final.get('current_value')
            forecasts = final.get('forecasts', {})
            
            normalized = {}
            for key in ['T1', 'T2', 'T3', 'T4', 'T5']:
                if key in forecasts and forecasts[key] is not None:
                    norm_score = normalize_indicator_forecast(
                        ind_name, current_value, forecasts[key], ind_type
                    )
                    normalized[key] = norm_score
                    
                    # Thu thập cho tổng hợp
                    all_normalized_scores[key].append(norm_score)
                    all_weights[key].append(weight)
            
            results['normalized_scores'][ind_name] = normalized
            
            # Voting cho xu hướng
            trend = final.get('trend', 'NEUTRAL')
            if any(x in str(trend).upper() for x in ['TĂNG', 'UP', 'MUA', 'BULLISH']):
                trend_votes['UP'] += weight
            elif any(x in str(trend).upper() for x in ['GIẢM', 'DOWN', 'BÁN', 'BEARISH']):
                trend_votes['DOWN'] += weight
            else:
                trend_votes['NEUTRAL'] += weight
            
            total_weight += weight
            
            # Thêm explanation
            results['explanations'].append({
                'indicator': ind_name,
                'type': ind_type,
                'current_value': current_value,
                'trend': trend,
                'confidence': final.get('confidence', 50),
                'normalized_T5': normalized.get('T5', 50)
            })
            
        except Exception as e:
            print(f"Lỗi dự báo {ind_name}: {e}")
            continue
    
    results['indicators_analyzed'] = len(results['individual_forecasts'])
    
    # TÍNH ĐIỂM TỔNG HỢP CÓ TRỌNG SỐ (từ điểm chuẩn hóa)
    combined_scores = {}
    for key in ['T1', 'T2', 'T3', 'T4', 'T5']:
        if all_normalized_scores[key] and all_weights[key]:
            total_w = sum(all_weights[key])
            if total_w > 0:
                weighted_sum = sum(s * w for s, w in zip(all_normalized_scores[key], all_weights[key]))
                combined_scores[key] = round(weighted_sum / total_w, 1)
    
    results['combined_forecast'] = {
        'scores': combined_scores,  # Điểm 0-100
        'trend_votes': trend_votes,
        'total_weight': total_weight
    }
    
    # Xác định xu hướng chung
    if total_weight > 0:
        up_pct = trend_votes['UP'] / total_weight * 100
        down_pct = trend_votes['DOWN'] / total_weight * 100
        
        if up_pct > 55:
            overall_trend = 'TĂNG'
            overall_confidence = min(up_pct, 90)
        elif down_pct > 55:
            overall_trend = 'GIẢM'
            overall_confidence = min(down_pct, 90)
        else:
            overall_trend = 'TRUNG LẬP'
            overall_confidence = 50
    else:
        overall_trend = 'KHÔNG XÁC ĐỊNH'
        overall_confidence = 0
        up_pct = down_pct = 0
    
    results['summary'] = {
        'indicators_analyzed': results['indicators_analyzed'],
        'indicators_found': len(available_indicators),
        'overall_trend': overall_trend,
        'overall_confidence': round(overall_confidence, 1),
        'trend_breakdown': {
            'up_pct': round(up_pct, 1),
            'down_pct': round(down_pct, 1),
            'neutral_pct': round(100 - up_pct - down_pct, 1)
        }
    }
    
    return results
forecast_all_26_indicators = forecast_all_26_indicators_v2

def calculate_daily_composite_score_v2(forecast_results, t0_score, current_price, atr_value=None):
    """
    PHIÊN BẢN 2.0: Tính điểm tổng hợp và DỰ BÁO GIÁ cho từng ngày
    
    Args:
        forecast_results: Kết quả từ forecast_all_26_indicators_v2()
        t0_score: Điểm T0 hiện tại (từ hệ thống cũ)
        current_price: Giá hiện tại
        atr_value: Giá trị ATR để tính biên độ giá
    
    Returns:
        Dict chứa điểm và giá dự báo từng ngày
    """
    
    if atr_value is None or atr_value <= 0:
        atr_value = current_price * 0.02  # Mặc định 2%
    
    combined = forecast_results.get('combined_forecast', {})
    scores = combined.get('scores', {})
    
    daily_results = {
        'T0': {
            'score': t0_score,
            'price': current_price,
            'change_score': 0,
            'change_price': 0,
            'change_price_pct': 0,
            'direction': 'BASE',
            'reversal': False,
            'reversal_type': None
        }
    }
    
    prev_score = t0_score
    prev_direction = None
    prev_price = current_price
    
    # Tính cho T1-T5
    for i in range(1, 6):
        key = f'T{i}'
        
        if key in scores:
            score = scores[key]
        else:
            # Nếu không có, ước tính từ xu hướng
            trend = forecast_results.get('summary', {}).get('overall_trend', 'TRUNG LẬP')
            if 'TĂNG' in trend:
                score = t0_score + 2 * i
            elif 'GIẢM' in trend:
                score = t0_score - 2 * i
            else:
                score = t0_score
            score = max(0, min(100, score))
        
        # Tính thay đổi điểm
        change_score = score - t0_score
        
        # TÍNH GIÁ DỰ BÁO từ điểm số
        # Điểm > 50: giá tăng, Điểm < 50: giá giảm
        # Mỗi 10 điểm chênh lệch = 1 ATR
        score_bias = (score - 50) / 50  # -1 to 1
        price_change = atr_value * score_bias * i * 0.5  # Điều chỉnh theo ngày
        predicted_price = current_price + price_change
        change_price_pct = (predicted_price - current_price) / current_price * 100
        
        # Xác định hướng
        if change_score > 3:
            direction = 'UP'
        elif change_score < -3:
            direction = 'DOWN'
        else:
            direction = 'FLAT'
        
        # Phát hiện đảo chiều
        reversal = False
        reversal_type = None
        if prev_direction and direction != 'FLAT' and prev_direction != 'FLAT':
            if prev_direction == 'DOWN' and direction == 'UP':
                reversal = True
                reversal_type = 'ĐẢO CHIỀU TĂNG'
            elif prev_direction == 'UP' and direction == 'DOWN':
                reversal = True
                reversal_type = 'ĐẢO CHIỀU GIẢM'
        
        daily_results[key] = {
            'score': round(score, 1),
            'price': round(predicted_price, 2),
            'change_score': round(change_score, 1),
            'change_price': round(price_change, 2),
            'change_price_pct': round(change_price_pct, 2),
            'direction': direction,
            'reversal': reversal,
            'reversal_type': reversal_type
        }
        
        prev_score = score
        if direction != 'FLAT':
            prev_direction = direction
        prev_price = predicted_price
    
    # Tính cho W1-W4 (dựa trên xu hướng T5)
    t5_data = daily_results.get('T5', {})
    t5_score = t5_data.get('score', t0_score)
    t5_price = t5_data.get('price', current_price)
    
    weekly_score_change = (t5_score - t0_score) / 5  # Thay đổi trung bình mỗi ngày
    weekly_price_change = (t5_price - current_price) / 5
    
    for i in range(1, 5):
        key = f'W{i}'
        # Mỗi tuần = 5 ngày, nhưng giảm dần momentum
        decay = 0.8 ** i
        
        score = t0_score + weekly_score_change * 5 * i * decay
        score = max(0, min(100, score))
        
        price = current_price + weekly_price_change * 5 * i * decay
        change_pct = (price - current_price) / current_price * 100
        
        daily_results[key] = {
            'score': round(score, 1),
            'price': round(price, 2),
            'change_score': round(score - t0_score, 1),
            'change_price': round(price - current_price, 2),
            'change_price_pct': round(change_pct, 2),
            'direction': 'UP' if score > t0_score + 3 else ('DOWN' if score < t0_score - 3 else 'FLAT'),
            'reversal': False,
            'reversal_type': None,
            'note': 'Ước tính từ T5'
        }
    
    # Tính cho M1-M3
    w4_data = daily_results.get('W4', {})
    w4_score = w4_data.get('score', t0_score)
    w4_price = w4_data.get('price', current_price)
    
    monthly_score_change = (w4_score - t0_score) / 20  # 4 tuần = 20 ngày
    monthly_price_change = (w4_price - current_price) / 20
    
    for i in range(1, 4):
        key = f'M{i}'
        # Mỗi tháng = 22 ngày, giảm dần
        decay = 0.6 ** i
        
        score = t0_score + monthly_score_change * 22 * i * decay
        score = max(0, min(100, score))
        
        price = current_price + monthly_price_change * 22 * i * decay
        change_pct = (price - current_price) / current_price * 100
        
        daily_results[key] = {
            'score': round(score, 1),
            'price': round(price, 2),
            'change_score': round(score - t0_score, 1),
            'change_price': round(price - current_price, 2),
            'change_price_pct': round(change_pct, 2),
            'direction': 'UP' if score > t0_score + 3 else ('DOWN' if score < t0_score - 3 else 'FLAT'),
            'reversal': False,
            'reversal_type': None,
            'note': 'Ước tính dài hạn'
        }
    
    # Phân tích tổng thể
    reversals = [k for k, v in daily_results.items() if v.get('reversal')]
    
    return {
        'daily_results': daily_results,
        'reversals_detected': reversals,
        'overall_direction': forecast_results.get('summary', {}).get('overall_trend', 'N/A'),
        'confidence': forecast_results.get('summary', {}).get('overall_confidence', 50)
    }

# =============================================================================
# SỬA LỖI 1: TÍNH ĐỘ TIN CẬY THỰC SỰ TỪ 12 PHƯƠNG PHÁP
# =============================================================================

def calculate_real_confidence(method_results):
    """
    Tính độ tin cậy thực sự dựa trên sự đồng thuận của 12 phương pháp
    
    Returns:
        tuple: (confidence_score, explanation_text)
    """
    if not method_results or not isinstance(method_results, dict):
        return 0.5, "Không có dữ liệu phương pháp"
    
    # Trọng số cho từng phương pháp
    method_weights = {
        'linear': 1.5,
        'polynomial': 1.5,
        'fourier': 1.2,
        'pattern': 1.3,
        'fibonacci': 1.0,
        'velocity': 1.0,
        'acceleration': 1.0,
        'peak_trough': 1.4,
        'multi_timeframe': 1.2,
        'statistics': 1.1,
        'logical': 0.8,
        'ensemble': 1.3
    }
    
    total_weight = 0
    weighted_votes = {'up': 0, 'down': 0, 'neutral': 0}
    active_methods = 0
    confidences = []
    
    for method_key, method_data in method_results.items():
        if method_data is None or not isinstance(method_data, dict):
            continue
        
        active_methods += 1
        weight = method_weights.get(method_key, 1.0)
        total_weight += weight
        
        # Lấy confidence của phương pháp
        method_conf = method_data.get('confidence', 50)
        if isinstance(method_conf, (int, float)):
            confidences.append(method_conf)
        
        # Xác định hướng từ nhiều nguồn
        direction = None
        
        # Từ trường 'trend'
        trend = method_data.get('trend', '')
        if isinstance(trend, str):
            trend_upper = trend.upper()
            if any(x in trend_upper for x in ['TĂNG', 'UP', 'BULLISH', 'MUA']):
                direction = 'up'
            elif any(x in trend_upper for x in ['GIẢM', 'DOWN', 'BEARISH', 'BÁN']):
                direction = 'down'
        
        # Từ trường 'signal'
        if direction is None:
            signal = method_data.get('signal', '')
            if isinstance(signal, str):
                signal_upper = signal.upper()
                if any(x in signal_upper for x in ['TĂNG', 'UP', 'BULLISH', 'MUA', 'BUY']):
                    direction = 'up'
                elif any(x in signal_upper for x in ['GIẢM', 'DOWN', 'BEARISH', 'BÁN', 'SELL']):
                    direction = 'down'
        
        # Từ forecasts (so sánh T5 vs T1)
        if direction is None:
            forecasts = method_data.get('forecasts', {})
            if isinstance(forecasts, dict):
                t5 = forecasts.get('T5')
                t1 = forecasts.get('T1')
                if t5 is not None and t1 is not None:
                    try:
                        t5_val = float(t5)
                        t1_val = float(t1)
                        if t5_val > t1_val * 1.02:
                            direction = 'up'
                        elif t5_val < t1_val * 0.98:
                            direction = 'down'
                    except (ValueError, TypeError):
                        pass
        
        # Vote
        if direction == 'up':
            weighted_votes['up'] += weight
        elif direction == 'down':
            weighted_votes['down'] += weight
        else:
            weighted_votes['neutral'] += weight
    
    if total_weight == 0 or active_methods == 0:
        return 0.5, "Không đủ dữ liệu"
    
    # Tính tỷ lệ đồng thuận
    max_direction = max(weighted_votes, key=weighted_votes.get)
    max_votes = weighted_votes[max_direction]
    
    # Confidence = tỷ lệ phương pháp đồng ý / tổng
    agreement_ratio = max_votes / total_weight
    
    # Kết hợp với confidence trung bình của các phương pháp
    avg_method_confidence = np.mean(confidences) / 100 if confidences else 0.5
    
    # Confidence cuối = 60% từ agreement + 40% từ avg confidence
    confidence = agreement_ratio * 0.6 + avg_method_confidence * 0.4
    
    # Điều chỉnh: nếu quá ít phương pháp hoạt động, giảm confidence
    if active_methods < 6:
        confidence *= (active_methods / 12)
    
    # Bonus nếu có sự đồng thuận cao
    if agreement_ratio > 0.7:
        confidence = min(confidence * 1.1, 0.95)
    
    # Giới hạn trong khoảng 0.3 - 0.95
    confidence = max(0.3, min(0.95, confidence))
    
    # Tạo giải thích
    direction_vn = {'up': 'TĂNG', 'down': 'GIẢM', 'neutral': 'ĐI NGANG'}
    up_pct = weighted_votes['up'] / total_weight * 100
    down_pct = weighted_votes['down'] / total_weight * 100
    
    explanation = f"{active_methods}/12 PP hoạt động | {up_pct:.0f}% TĂNG, {down_pct:.0f}% GIẢM | Đồng thuận: {direction_vn[max_direction]}"
    
    return round(confidence, 2), explanation


# =============================================================================
# SỬA LỖI 2: PHÁT HIỆN ĐẢO CHIỀU (TURNING POINTS) - CẬP NHẬT
# =============================================================================

def detect_turning_points(df, indicator_results):
    """
    Phát hiện các điểm đảo chiều từ dữ liệu lịch sử
    Trả về dict với đầy đủ thông tin
    """
    turning_points = {
        'detected': False,
        'type': None,
        'confidence': 0,
        'days_from_turn': None,
        'historical_similar': [],
        'warning': None,
        'details': [],  # THÊM: Chi tiết các phát hiện
        'signals': []   # THÊM: Các tín hiệu cụ thể
    }
    
    if df is None or len(df) < 30:
        return turning_points
    
    close = df['close'].values if 'close' in df.columns else None
    if close is None or len(close) < 30:
        return turning_points
    
    current_price = close[-1]
    current_idx = len(close) - 1
    
    # 1. Phát hiện đỉnh/đáy gần đây (trong 10 ngày)
    try:
        from scipy.signal import argrelextrema
        
        # Tìm đỉnh cục bộ
        peaks_idx = argrelextrema(close, np.greater, order=5)[0]
        troughs_idx = argrelextrema(close, np.less, order=5)[0]
        
        # Kiểm tra đỉnh gần nhất
        if len(peaks_idx) > 0:
            last_peak_idx = peaks_idx[-1]
            days_from_peak = current_idx - last_peak_idx
            
            if days_from_peak <= 10 and days_from_peak > 0:
                peak_price = close[last_peak_idx]
                drop_pct = (current_price - peak_price) / peak_price * 100
                
                if drop_pct < -3:  # Giảm > 3% từ đỉnh
                    turning_points['detected'] = True
                    turning_points['type'] = 'peak_reversal'
                    turning_points['days_from_turn'] = days_from_peak
                    turning_points['confidence'] = min(0.9, abs(drop_pct) / 15)
                    
                    detail = f"📉 VỪA TẠO ĐỈNH {days_from_peak} ngày trước tại {peak_price:,.0f}, đã giảm {abs(drop_pct):.1f}%"
                    turning_points['details'].append(detail)
                    turning_points['warning'] = detail
                    turning_points['signals'].append({
                        'type': 'PEAK_FORMED',
                        'price': peak_price,
                        'days_ago': days_from_peak,
                        'change_pct': drop_pct
                    })
        
        # Kiểm tra đáy gần nhất
        if len(troughs_idx) > 0:
            last_trough_idx = troughs_idx[-1]
            days_from_trough = current_idx - last_trough_idx
            
            if days_from_trough <= 10 and days_from_trough > 0:
                trough_price = close[last_trough_idx]
                rise_pct = (current_price - trough_price) / trough_price * 100
                
                if rise_pct > 3:  # Tăng > 3% từ đáy
                    turning_points['detected'] = True
                    turning_points['type'] = 'trough_reversal'
                    turning_points['days_from_turn'] = days_from_trough
                    turning_points['confidence'] = min(0.9, rise_pct / 15)
                    
                    detail = f"📈 VỪA TẠO ĐÁY {days_from_trough} ngày trước tại {trough_price:,.0f}, đã tăng {rise_pct:.1f}%"
                    turning_points['details'].append(detail)
                    if not turning_points['warning']:
                        turning_points['warning'] = detail
                    else:
                        turning_points['warning'] += f"\n{detail}"
                    turning_points['signals'].append({
                        'type': 'TROUGH_FORMED',
                        'price': trough_price,
                        'days_ago': days_from_trough,
                        'change_pct': rise_pct
                    })
    except Exception as e:
        print(f"Lỗi phát hiện đỉnh/đáy: {e}")
    
    # 2. Kiểm tra RSI divergence (phân kỳ)
    if 'RSI' in df.columns:
        try:
            rsi = df['RSI'].values
            
            if len(close) >= 20 and len(rsi) >= 20:
                # Lấy dữ liệu 20 ngày gần nhất
                close_20 = close[-20:]
                rsi_20 = rsi[-20:]
                
                # Loại bỏ NaN
                valid_mask = ~np.isnan(rsi_20)
                if np.sum(valid_mask) >= 15:
                    price_trend = (close_20[-1] - close_20[0]) / close_20[0] * 100
                    
                    # Tính RSI trend (bỏ qua NaN)
                    rsi_valid = rsi_20[valid_mask]
                    rsi_trend = rsi_valid[-1] - rsi_valid[0] if len(rsi_valid) > 1 else 0
                    
                    # Phân kỳ âm: giá tăng nhưng RSI giảm
                    if price_trend > 5 and rsi_trend < -10:
                        turning_points['detected'] = True
                        turning_points['type'] = 'bearish_divergence'
                        turning_points['confidence'] = max(turning_points['confidence'], 0.75)
                        
                        detail = f"⚠️ PHÂN KỲ ÂM: Giá +{price_trend:.1f}% nhưng RSI {rsi_trend:.1f} (20 ngày)"
                        turning_points['details'].append(detail)
                        if turning_points['warning']:
                            turning_points['warning'] += f"\n{detail}"
                        else:
                            turning_points['warning'] = detail
                        turning_points['signals'].append({
                            'type': 'BEARISH_DIVERGENCE',
                            'price_change': price_trend,
                            'rsi_change': rsi_trend
                        })
                    
                    # Phân kỳ dương: giá giảm nhưng RSI tăng
                    elif price_trend < -5 and rsi_trend > 10:
                        turning_points['detected'] = True
                        turning_points['type'] = 'bullish_divergence'
                        turning_points['confidence'] = max(turning_points['confidence'], 0.75)
                        
                        detail = f"🔄 PHÂN KỲ DƯƠNG: Giá {price_trend:.1f}% nhưng RSI +{rsi_trend:.1f} (20 ngày)"
                        turning_points['details'].append(detail)
                        if turning_points['warning']:
                            turning_points['warning'] += f"\n{detail}"
                        else:
                            turning_points['warning'] = detail
                        turning_points['signals'].append({
                            'type': 'BULLISH_DIVERGENCE',
                            'price_change': price_trend,
                            'rsi_change': rsi_trend
                        })
        except Exception as e:
            print(f"Lỗi phân tích RSI divergence: {e}")
    
    # 3. Kiểm tra MACD crossover gần đây
    if 'MACD_Hist' in df.columns:
        try:
            macd_hist = df['MACD_Hist'].values
            
            # Tìm crossover trong 5 ngày gần nhất
            for i in range(1, min(6, len(macd_hist))):
                if i >= len(macd_hist):
                    break
                    
                prev_hist = macd_hist[-(i+1)]
                curr_hist = macd_hist[-i]
                
                if pd.notna(prev_hist) and pd.notna(curr_hist):
                    if prev_hist < 0 and curr_hist > 0:
                        detail = f"🔼 MACD Cross Up cách đây {i} ngày"
                        turning_points['details'].append(detail)
                        turning_points['signals'].append({
                            'type': 'MACD_CROSS_UP',
                            'days_ago': i
                        })
                        if not turning_points['detected']:
                            turning_points['detected'] = True
                            turning_points['type'] = 'macd_bullish_cross'
                            turning_points['confidence'] = 0.65
                        break
                    elif prev_hist > 0 and curr_hist < 0:
                        detail = f"🔽 MACD Cross Down cách đây {i} ngày"
                        turning_points['details'].append(detail)
                        turning_points['signals'].append({
                            'type': 'MACD_CROSS_DOWN',
                            'days_ago': i
                        })
                        if not turning_points['detected']:
                            turning_points['detected'] = True
                            turning_points['type'] = 'macd_bearish_cross'
                            turning_points['confidence'] = 0.65
                        break
        except Exception as e:
            print(f"Lỗi phân tích MACD: {e}")
    
    # 4. Kiểm tra xu hướng thay đổi (momentum shift)
    if len(close) >= 10:
        try:
            # Tính momentum ngắn hạn vs trung hạn
            mom_5d = (close[-1] - close[-5]) / close[-5] * 100 if len(close) >= 5 else 0
            mom_10d = (close[-1] - close[-10]) / close[-10] * 100 if len(close) >= 10 else 0
            
            # Momentum đang đảo chiều
            if mom_10d < -5 and mom_5d > 2:
                detail = f"📊 Momentum đảo chiều TĂNG: 10d={mom_10d:.1f}%, 5d={mom_5d:+.1f}%"
                turning_points['details'].append(detail)
                if not turning_points['detected']:
                    turning_points['detected'] = True
                    turning_points['type'] = 'momentum_shift_up'
                    turning_points['confidence'] = 0.6
            elif mom_10d > 5 and mom_5d < -2:
                detail = f"📊 Momentum đảo chiều GIẢM: 10d=+{mom_10d:.1f}%, 5d={mom_5d:.1f}%"
                turning_points['details'].append(detail)
                if not turning_points['detected']:
                    turning_points['detected'] = True
                    turning_points['type'] = 'momentum_shift_down'
                    turning_points['confidence'] = 0.6
        except Exception as e:
            print(f"Lỗi phân tích momentum: {e}")
    
    # 5. So sánh với lịch sử (nếu có đủ dữ liệu)
    if len(close) >= 252:  # Ít nhất 1 năm
        try:
            current_pattern = close[-20:] / close[-20] * 100
            
            similar_patterns = []
            for i in range(252, len(close) - 25, 20):
                hist_pattern = close[i-20:i] / close[i-20] * 100
                
                if len(hist_pattern) == len(current_pattern):
                    corr = np.corrcoef(current_pattern, hist_pattern)[0, 1]
                    
                    if corr > 0.85:
                        future_return = (close[i+5] - close[i]) / close[i] * 100
                        similar_patterns.append({
                            'date_idx': i,
                            'correlation': corr,
                            'future_5d_return': future_return
                        })
            
            if len(similar_patterns) >= 3:
                avg_future = np.mean([p['future_5d_return'] for p in similar_patterns])
                turning_points['historical_similar'] = similar_patterns[:5]
                
                if avg_future < -3:
                    detail = f"📉 Lịch sử: {len(similar_patterns)} mẫu tương tự → TB giảm {abs(avg_future):.1f}% trong 5 ngày"
                    turning_points['details'].append(detail)
                elif avg_future > 3:
                    detail = f"📈 Lịch sử: {len(similar_patterns)} mẫu tương tự → TB tăng {avg_future:.1f}% trong 5 ngày"
                    turning_points['details'].append(detail)
        except Exception as e:
            print(f"Lỗi so sánh lịch sử: {e}")
    
    return turning_points


# =============================================================================
# SỬA LỖI 3: DỰ BÁO CÓ DAO ĐỘNG (KHÔNG CÒN 1 CHIỀU)
# =============================================================================

def calculate_realistic_forecast(indicator_results, current_price, atr_value, turning_points):
    """
    Tính dự báo thực tế hơn, có dao động, không phải đường thẳng
    """
    forecasts = {}
    
    if current_price is None or current_price <= 0:
        for day in ['T1', 'T2', 'T3', 'T4', 'T5']:
            forecasts[day] = {
                'score': 50,
                'price': 0,
                'change_pct': 0,
                'direction': 'neutral',
                'volatility_range': (0, 0)
            }
        return forecasts
    
    # Xử lý ATR
    if atr_value is None or atr_value <= 0:
        atr_value = current_price * 0.02
    
    volatility = atr_value / current_price
    
    # Lấy điểm tổng hợp từ indicator_results
    base_scores = {}
    overall_bias = 0
    
    if indicator_results and isinstance(indicator_results, dict):
        combined = indicator_results.get('combined_forecast', {})
        
        # Xử lý scores
        if isinstance(combined, dict):
            scores = combined.get('scores', {})
            if isinstance(scores, dict):
                for key in ['T1', 'T2', 'T3', 'T4', 'T5']:
                    if key in scores:
                        val = scores[key]
                        if isinstance(val, (int, float)):
                            base_scores[key] = float(val)
        
        # Lấy xu hướng tổng
        summary = indicator_results.get('summary', {})
        if isinstance(summary, dict):
            trend = summary.get('overall_trend', '')
            if isinstance(trend, str):
                if 'TĂNG' in trend.upper():
                    overall_bias = 0.6
                elif 'GIẢM' in trend.upper():
                    overall_bias = -0.6
    
    # Điều chỉnh theo turning points
    turn_adjustment = 0
    turn_volatility_mult = 1.0
    
    if turning_points and isinstance(turning_points, dict) and turning_points.get('detected'):
        turn_type = turning_points.get('type', '')
        turn_conf = turning_points.get('confidence', 0.5)
        
        try:
            turn_conf = float(turn_conf)
        except:
            turn_conf = 0.5
        
        if isinstance(turn_type, str):
            if 'bearish' in turn_type.lower() or 'peak' in turn_type.lower():
                turn_adjustment = -0.3 * turn_conf
                turn_volatility_mult = 1.2
            elif 'bullish' in turn_type.lower() or 'trough' in turn_type.lower():
                turn_adjustment = 0.3 * turn_conf
                turn_volatility_mult = 1.2
    
    # Tính dự báo cho từng ngày với DAO ĐỘNG
    for i, day in enumerate(['T1', 'T2', 'T3', 'T4', 'T5'], 1):
        # Điểm cơ bản
        base_score = base_scores.get(day, 50)
        
        # Thêm dao động theo ngày (sine wave)
        wave = np.sin(i * 0.7) * 5  # Dao động ±5 điểm
        
        # Thêm xu hướng
        trend_component = overall_bias * i * 3
        
        # Thêm turning point adjustment
        turn_component = turn_adjustment * 10 * (1 - i * 0.1)  # Giảm dần theo thời gian
        
        # Điểm cuối cùng
        final_score = base_score + wave + trend_component + turn_component
        final_score = max(20, min(80, final_score))
        
        # Tính % thay đổi từ score
        score_bias = (final_score - 50) / 50  # -1 to 1
        
        # Base change từ score
        base_change = score_bias * volatility * i * 0.8
        
        # Thêm wave component cho giá
        price_wave = np.sin(i * 0.9) * volatility * 0.3
        
        # Turning point adjustment cho giá
        price_turn_adj = turn_adjustment * volatility * i * 0.5
        
        # Tổng % thay đổi
        total_change = base_change + price_wave + price_turn_adj
        
        # Tính giá
        forecast_price = current_price * (1 + total_change)
        
        # Tính range (uncertainty tăng theo thời gian)
        uncertainty = volatility * i * 0.5 * turn_volatility_mult
        price_low = forecast_price * (1 - uncertainty)
        price_high = forecast_price * (1 + uncertainty)
        
        # Xác định hướng
        if total_change > 0.005:
            direction = 'up'
        elif total_change < -0.005:
            direction = 'down'
        else:
            direction = 'neutral'
        
        forecasts[day] = {
            'score': round(final_score, 1),
            'price': round(forecast_price, 2),
            'change_pct': round(total_change * 100, 2),
            'direction': direction,
            'volatility_range': (round(price_low, 2), round(price_high, 2)),
            'confidence': round(max(0.3, 0.8 - i * 0.1), 2)  # Confidence giảm theo thời gian
        }
    
    return forecasts

def run_12_methods_forecast_v2(df, symbol, weighted_scores, final_score, current_price=None, atr_value=None):
    """
    PHIÊN BẢN ĐÃ SỬA
    """
    
    result = {
        'indicator_forecasts': None,
        'daily_composite': None,
        'markdown_report': '',
        'word_report': {},
        'success': False,
        'confidence': 0.5,
        'confidence_explanation': '',
        'turning_points': {},
        'warnings': [],
        'realistic_forecasts': {}
    }
    
    try:
        # Lấy giá hiện tại
        if current_price is None:
            close_col = 'close' if 'close' in df.columns else 'Close'
            if close_col in df.columns:
                current_price = float(df[close_col].iloc[-1])
            else:
                current_price = 0
        
        # Lấy ATR
        if atr_value is None:
            atr_col = 'ATR' if 'ATR' in df.columns else 'atr'
            if atr_col in df.columns:
                atr_value = float(df[atr_col].iloc[-1])
            else:
                atr_value = current_price * 0.02
        
        # 1. Chạy dự báo 12 phương pháp
        print(f"\n[1/5] Đang chạy 12 phương pháp cho {symbol}...")
        indicator_forecast_results = forecast_all_26_indicators_v2(df, None)
        result['indicator_forecasts'] = indicator_forecast_results
        
        # 2. TÍNH ĐỘ TIN CẬY - SỬA: Lấy đúng method_results
        print(f"[2/5] Đang tính độ tin cậy...")
        sample_method_results = {}
        
        if indicator_forecast_results:
            individual = indicator_forecast_results.get('individual_forecasts', {})
            for ind_name, ind_data in individual.items():
                if ind_data and isinstance(ind_data, dict):
                    detailed = ind_data.get('detailed')
                    if detailed and isinstance(detailed, dict):
                        methods = detailed.get('methods')
                        if methods and isinstance(methods, dict) and len(methods) > 0:
                            sample_method_results = methods
                            print(f"   → Lấy methods từ {ind_name}: {len(methods)} phương pháp")
                            break
        
        confidence, conf_explanation = calculate_real_confidence(sample_method_results)
        result['confidence'] = confidence
        result['confidence_explanation'] = conf_explanation
        print(f"   → Độ tin cậy: {confidence*100:.0f}% ({conf_explanation})")
        
        # 3. PHÁT HIỆN ĐẢO CHIỀU
        print(f"[3/5] Đang phát hiện đảo chiều...")
        turning_points = detect_turning_points(df, indicator_forecast_results)
        result['turning_points'] = turning_points
        
        if turning_points.get('detected'):
            print(f"   → PHÁT HIỆN: {turning_points.get('type')}")
            for detail in turning_points.get('details', []):
                print(f"      {detail}")
                result['warnings'].append(detail)
        else:
            print(f"   → Không phát hiện đảo chiều")
        
        # 4. TÍNH DỰ BÁO THỰC TẾ (có dao động)
        print(f"[4/5] Đang tính dự báo thực tế...")
        realistic_forecasts = calculate_realistic_forecast(
            indicator_forecast_results,
            current_price,
            atr_value,
            turning_points
        )
        result['realistic_forecasts'] = realistic_forecasts
        
        # 5. Tính điểm tổng hợp
        print(f"[5/5] Đang tính điểm tổng hợp...")
        daily_composite = calculate_daily_composite_score_v2(
            indicator_forecast_results, 
            final_score,
            current_price,
            atr_value
        )
        
        # Cập nhật với realistic forecasts
        if daily_composite and realistic_forecasts:
            daily_composite['confidence'] = confidence
            daily_composite['confidence_explanation'] = conf_explanation
            daily_composite['turning_points'] = turning_points
            daily_composite['realistic_forecasts'] = realistic_forecasts
            
            # Merge giá từ realistic_forecasts vào daily_results
            if 'daily_results' in daily_composite:
                for day, forecast_data in realistic_forecasts.items():
                    if day in daily_composite['daily_results']:
                        daily_composite['daily_results'][day].update({
                            'price': forecast_data['price'],
                            'change_price_pct': forecast_data['change_pct'],
                            'direction': forecast_data['direction'].upper(),
                            'volatility_range': forecast_data.get('volatility_range'),
                            'forecast_confidence': forecast_data.get('confidence', 0.5)
                        })
        
        result['daily_composite'] = daily_composite
        
        # Tạo báo cáo
        markdown_report = generate_forecast_report_markdown_v2(
            symbol,
            indicator_forecast_results,
            daily_composite,
            current_price
        )
        
        # Thêm cảnh báo đảo chiều vào báo cáo
        if turning_points.get('detected') and turning_points.get('details'):
            warning_section = f"\n\n## ⚠️ CẢNH BÁO ĐẢO CHIỀU\n\n"
            warning_section += f"**Loại:** {turning_points.get('type', 'N/A')}\n\n"
            warning_section += f"**Độ tin cậy:** {turning_points.get('confidence', 0)*100:.0f}%\n\n"
            warning_section += f"**Chi tiết:**\n"
            for detail in turning_points.get('details', []):
                warning_section += f"- {detail}\n"
            
            # Chèn vào đầu báo cáo
            if markdown_report.startswith("#"):
                lines = markdown_report.split("\n", 1)
                markdown_report = lines[0] + "\n" + warning_section + "\n" + (lines[1] if len(lines) > 1 else "")
            else:
                markdown_report = warning_section + markdown_report
        
        result['markdown_report'] = markdown_report
        
        # Word report
        word_report = generate_forecast_report_for_word(
            symbol,
            indicator_forecast_results,
            daily_composite,
            current_price
        )
        word_report['confidence'] = confidence
        word_report['turning_points'] = turning_points
        result['word_report'] = word_report
        
        result['success'] = True
        
        # In kết quả
        print(f"\n{'='*60}")
        print(f"KẾT QUẢ DỰ BÁO 12 PHƯƠNG PHÁP CHO {symbol}")
        print(f"{'='*60}")
        print(f"Giá hiện tại: {current_price:,.0f}")
        print(f"Chỉ báo phân tích: {indicator_forecast_results.get('summary', {}).get('indicators_analyzed', 0) if indicator_forecast_results else 0}")
        print(f"Độ tin cậy: {confidence*100:.0f}% ({conf_explanation})")
        
        if turning_points.get('detected'):
            print(f"\n⚠️ CẢNH BÁO ĐẢO CHIỀU:")
            for detail in turning_points.get('details', []):
                print(f"   {detail}")
        
        print(f"\nDỰ BÁO GIÁ (có dao động):")
        for day in ['T1', 'T2', 'T3', 'T4', 'T5']:
            if day in realistic_forecasts:
                f = realistic_forecasts[day]
                dir_icon = '📈' if f['direction'] == 'up' else ('📉' if f['direction'] == 'down' else '➡️')
                vol_range = f.get('volatility_range', (f['price'], f['price']))
                print(f"   {day}: {f['price']:,.0f} ({f['change_pct']:+.1f}%) {dir_icon}")
                print(f"       Range: {vol_range[0]:,.0f} - {vol_range[1]:,.0f} | Tin cậy: {f.get('confidence', 0.5)*100:.0f}%")
        
        return result
        
    except Exception as e:
        print(f"Lỗi run_12_methods_forecast_v2 cho {symbol}: {e}")
        import traceback
        traceback.print_exc()
        result['markdown_report'] = f"## {symbol}\n\n❌ Lỗi: {str(e)}"
        return result
        
    except Exception as e:
        print(f"Lỗi run_12_methods_forecast_v2 cho {symbol}: {e}")
        import traceback
        traceback.print_exc()
        result['markdown_report'] = f"## {symbol}\n\n❌ Lỗi: {str(e)}"
        result['word_report'] = {'error': str(e)}
        return result


# Alias để tương thích ngược
run_12_methods_forecast = run_12_methods_forecast_v2

def generate_forecast_report_markdown_v2(symbol, forecast_results, daily_composite, current_price):
    """
    PHIÊN BẢN 2.0: Tạo báo cáo Markdown với giá dự báo
    """
    
    if not forecast_results or not daily_composite:
        return f"## {symbol}\n\n❌ Không có dữ liệu dự báo"
    
    summary = forecast_results.get('summary', {})
    individual = forecast_results.get('individual_forecasts', {})
    normalized = forecast_results.get('normalized_scores', {})
    daily_results = daily_composite.get('daily_results', {})
    reversals = daily_composite.get('reversals_detected', [])
    
    report = []
    
    # ===== HEADER =====
    report.append(f"# 📊 BÁO CÁO DỰ BÁO CHI TIẾT: {symbol}")
    report.append(f"\n**Thời gian:** {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    report.append(f"\n**Giá hiện tại:** {current_price:,.0f}")
    report.append(f"\n**Số chỉ báo tìm thấy:** {summary.get('indicators_found', 0)}")
    report.append(f"\n**Số chỉ báo phân tích:** {summary.get('indicators_analyzed', 0)}")
    report.append("\n---\n")
    
    # ===== I. TÓM TẮT =====
    report.append("## I. TÓM TẮT DỰ BÁO\n")
    
    overall_trend = summary.get('overall_trend', 'N/A')
    confidence = summary.get('overall_confidence', 0)
    trend_breakdown = summary.get('trend_breakdown', {})
    
    if 'TĂNG' in overall_trend:
        trend_emoji = '🟢'
    elif 'GIẢM' in overall_trend:
        trend_emoji = '🔴'
    else:
        trend_emoji = '🟡'
    
    report.append(f"### {trend_emoji} XU HƯỚNG TỔNG HỢP: **{overall_trend}**\n")
    report.append(f"- **Độ tin cậy:** {confidence:.1f}%")
    report.append(f"- **Tỷ lệ tăng:** {trend_breakdown.get('up_pct', 0):.1f}%")
    report.append(f"- **Tỷ lệ giảm:** {trend_breakdown.get('down_pct', 0):.1f}%")
    report.append(f"- **Tỷ lệ trung lập:** {trend_breakdown.get('neutral_pct', 0):.1f}%")
    report.append("")
    
    # ===== II. DỰ BÁO GIÁ THEO NGÀY =====
    report.append("## II. DỰ BÁO GIÁ THEO THỜI GIAN\n")
    
    # Bảng T0-T5
    report.append("### A. Ngắn hạn (T0-T5)\n")
    report.append("| Ngày | Điểm | Giá dự báo | Thay đổi | Hướng | Đảo chiều |")
    report.append("|:----:|:----:|:----------:|:--------:|:-----:|:---------:|")
    
    for key in ['T0', 'T1', 'T2', 'T3', 'T4', 'T5']:
        data = daily_results.get(key, {})
        score = data.get('score', '-')
        price = data.get('price', '-')
        change_pct = data.get('change_price_pct', 0)
        direction = data.get('direction', '-')
        reversal = data.get('reversal_type', '')
        
        score_str = f"{score:.1f}" if isinstance(score, (int, float)) else str(score)
        price_str = f"{price:,.0f}" if isinstance(price, (int, float)) else str(price)
        pct_str = f"{change_pct:+.1f}%" if isinstance(change_pct, (int, float)) else "-"
        
        dir_emoji = '📈' if direction == 'UP' else ('📉' if direction == 'DOWN' else '➡️')
        reversal_str = f"⚠️ {reversal}" if reversal else "-"
        
        report.append(f"| {key} | {score_str} | {price_str} | {pct_str} | {dir_emoji} {direction} | {reversal_str} |")
    
    report.append("")
    
    # Bảng W1-W4
    report.append("### B. Trung hạn (W1-W4)\n")
    report.append("| Tuần | Điểm | Giá dự báo | Thay đổi | Hướng |")
    report.append("|:----:|:----:|:----------:|:--------:|:-----:|")
    
    for key in ['W1', 'W2', 'W3', 'W4']:
        data = daily_results.get(key, {})
        score = data.get('score', '-')
        price = data.get('price', '-')
        change_pct = data.get('change_price_pct', 0)
        direction = data.get('direction', '-')
        
        score_str = f"{score:.1f}" if isinstance(score, (int, float)) else str(score)
        price_str = f"{price:,.0f}" if isinstance(price, (int, float)) else str(price)
        pct_str = f"{change_pct:+.1f}%" if isinstance(change_pct, (int, float)) else "-"
        dir_emoji = '📈' if direction == 'UP' else ('📉' if direction == 'DOWN' else '➡️')
        
        report.append(f"| {key} | {score_str} | {price_str} | {pct_str} | {dir_emoji} {direction} |")
    
    report.append("")
    
    # Bảng M1-M3
    report.append("### C. Dài hạn (M1-M3)\n")
    report.append("| Tháng | Điểm | Giá dự báo | Thay đổi | Hướng |")
    report.append("|:-----:|:----:|:----------:|:--------:|:-----:|")
    
    for key in ['M1', 'M2', 'M3']:
        data = daily_results.get(key, {})
        score = data.get('score', '-')
        price = data.get('price', '-')
        change_pct = data.get('change_price_pct', 0)
        direction = data.get('direction', '-')
        
        score_str = f"{score:.1f}" if isinstance(score, (int, float)) else str(score)
        price_str = f"{price:,.0f}" if isinstance(price, (int, float)) else str(price)
        pct_str = f"{change_pct:+.1f}%" if isinstance(change_pct, (int, float)) else "-"
        dir_emoji = '📈' if direction == 'UP' else ('📉' if direction == 'DOWN' else '➡️')
        
        report.append(f"| {key} | {score_str} | {price_str} | {pct_str} | {dir_emoji} {direction} |")
    
    report.append("")
    
    # ===== III. PHÁT HIỆN ĐẢO CHIỀU =====
    report.append("## III. PHÁT HIỆN ĐẢO CHIỀU\n")
    
    if reversals:
        report.append(f"⚠️ **Phát hiện {len(reversals)} điểm đảo chiều tiềm năng:**\n")
        for rev_key in reversals:
            rev_data = daily_results.get(rev_key, {})
            rev_type = rev_data.get('reversal_type', 'N/A')
            rev_price = rev_data.get('price', 0)
            report.append(f"- **{rev_key}:** {rev_type} (Giá dự báo: {rev_price:,.0f})")
        report.append("")
    else:
        report.append("✅ Không phát hiện điểm đảo chiều trong kỳ dự báo.\n")
    
    # ===== IV. PHÂN TÍCH CHI TIẾT =====
    report.append("## IV. PHÂN TÍCH CHI TIẾT TỪNG CHỈ BÁO\n")
    
    # Sắp xếp theo confidence
    sorted_indicators = sorted(
        individual.items(),
        key=lambda x: x[1].get('final', {}).get('confidence', 0),
        reverse=True
    )
    
    for ind_name, ind_data in sorted_indicators[:10]:
        final = ind_data.get('final', {})
        detailed = ind_data.get('detailed', {})
        ind_type = ind_data.get('indicator_type', 'unknown')
        norm_scores = normalized.get(ind_name, {})
        
        current_val = final.get('current_value', 'N/A')
        trend = final.get('trend', 'N/A')
        confidence = final.get('confidence', 0)
        weight = final.get('weight', 1)
        
        if 'TĂNG' in str(trend) or 'UP' in str(trend):
            ind_emoji = '🟢'
        elif 'GIẢM' in str(trend) or 'DOWN' in str(trend):
            ind_emoji = '🔴'
        else:
            ind_emoji = '🟡'
        
        report.append(f"### {ind_emoji} {ind_name} ({ind_type})\n")
        report.append(f"- **Giá trị hiện tại:** {current_val}")
        report.append(f"- **Xu hướng:** {trend}")
        report.append(f"- **Độ tin cậy:** {confidence:.1f}%")
        report.append(f"- **Trọng số:** {weight}")
        
        # Điểm chuẩn hóa
        if norm_scores:
            norm_str = " → ".join([f"T{i}: {norm_scores.get(f'T{i}', '-'):.0f}" for i in range(1, 6) if f'T{i}' in norm_scores])
            report.append(f"- **Điểm chuẩn hóa:** {norm_str}")
        
        report.append("")
        report.append("---\n")
    
    # ===== V. KẾT LUẬN =====
    report.append("## V. KẾT LUẬN VÀ KHUYẾN NGHỊ\n")
    
    # Tính tín hiệu từ các chỉ báo quan trọng
    important_indicators = ['RSI', 'MACD_Hist', 'Stoch_K', 'MFI', 'ADX', 'OBV', 'CCI']
    signals = {'buy': 0, 'sell': 0, 'hold': 0}
    
    for ind_name in important_indicators:
        if ind_name in individual:
            trend = individual[ind_name].get('final', {}).get('trend', '')
            if any(x in str(trend).upper() for x in ['TĂNG', 'UP', 'MUA', 'BULLISH']):
                signals['buy'] += 1
            elif any(x in str(trend).upper() for x in ['GIẢM', 'DOWN', 'BÁN', 'BEARISH']):
                signals['sell'] += 1
            else:
                signals['hold'] += 1
    
    total_signals = signals['buy'] + signals['sell'] + signals['hold']
    
    report.append(f"### Tín hiệu từ {total_signals} chỉ báo chính:")
    report.append(f"- 🟢 **MUA:** {signals['buy']}/{total_signals}")
    report.append(f"- 🔴 **BÁN:** {signals['sell']}/{total_signals}")
    report.append(f"- 🟡 **GIỮ:** {signals['hold']}/{total_signals}")
    report.append("")
    
    # Khuyến nghị
    if signals['buy'] >= total_signals * 0.6:
        recommendation = "🟢 **MUA** - Đa số chỉ báo đồng thuận tích cực"
    elif signals['sell'] >= total_signals * 0.6:
        recommendation = "🔴 **BÁN** - Đa số chỉ báo đồng thuận tiêu cực"
    elif signals['buy'] > signals['sell']:
        recommendation = "🟢 **CÂN NHẮC MUA** - Tín hiệu tích cực chiếm ưu thế"
    elif signals['sell'] > signals['buy']:
        recommendation = "🔴 **CÂN NHẮC BÁN** - Tín hiệu tiêu cực chiếm ưu thế"
    else:
        recommendation = "🟡 **THEO DÕI** - Tín hiệu hỗn hợp, chờ xác nhận"
    
    report.append(f"### Khuyến nghị: {recommendation}\n")
    
    # Dự báo giá
    t5_data = daily_results.get('T5', {})
    w4_data = daily_results.get('W4', {})
    m3_data = daily_results.get('M3', {})
    
    report.append("### Dự báo giá:")
    report.append(f"- **T5 (5 ngày):** {t5_data.get('price', '-'):,.0f} ({t5_data.get('change_price_pct', 0):+.1f}%)")
    report.append(f"- **W4 (4 tuần):** {w4_data.get('price', '-'):,.0f} ({w4_data.get('change_price_pct', 0):+.1f}%)")
    report.append(f"- **M3 (3 tháng):** {m3_data.get('price', '-'):,.0f} ({m3_data.get('change_price_pct', 0):+.1f}%)")
    report.append("")
    
    # Lưu ý
    report.append("### ⚠️ Lưu ý quan trọng:")
    report.append("- Dự báo dựa trên phân tích kỹ thuật, không phải lời khuyên đầu tư")
    report.append("- Độ tin cậy giảm dần theo thời gian (T1 > T5 > W4 > M3)")
    report.append("- Luôn kết hợp với phân tích cơ bản và quản lý rủi ro")
    report.append("- Đặt stop-loss để bảo vệ vốn")
    report.append("")
    
    return "\n".join(report)


# ============================================================
# HÀM TÍNH ĐIỂM TỔNG HỢP THEO NGÀY
# ============================================================

def calculate_daily_composite_score(forecast_results, t0_score):
    """
    Tính điểm tổng hợp có trọng số cho từng ngày (T1-T5, W1-W4, M1-M3)
    và so sánh với T0
    
    Args:
        forecast_results: Kết quả từ forecast_all_26_indicators()
        t0_score: Điểm T0 hiện tại (từ hệ thống cũ)
    
    Returns:
        Dict chứa điểm từng ngày và phân tích đảo chiều
    """
    
    combined = forecast_results.get('combined_forecast', {})
    forecasts = combined.get('forecasts', {})
    
    daily_scores = {
        'T0': {
            'score': t0_score,
            'change': 0,
            'change_pct': 0,
            'direction': 'BASE',
            'reversal': False
        }
    }
    
    prev_score = t0_score
    prev_direction = None
    
    # Tính cho T1-T5
    for i in range(1, 6):
        key = f'T{i}'
        if key in forecasts:
            score = forecasts[key]
            change = score - t0_score
            change_pct = (change / t0_score * 100) if t0_score != 0 else 0
            
            # Xác định hướng
            if change > 1:
                direction = 'UP'
            elif change < -1:
                direction = 'DOWN'
            else:
                direction = 'FLAT'
            
            # Phát hiện đảo chiều
            reversal = False
            if prev_direction and direction != 'FLAT':
                if prev_direction == 'DOWN' and direction == 'UP':
                    reversal = True
                    reversal_type = 'ĐẢO CHIỀU TĂNG'
                elif prev_direction == 'UP' and direction == 'DOWN':
                    reversal = True
                    reversal_type = 'ĐẢO CHIỀU GIẢM'
                else:
                    reversal_type = None
            else:
                reversal_type = None
            
            daily_scores[key] = {
                'score': round(score, 2),
                'change': round(change, 2),
                'change_pct': round(change_pct, 2),
                'direction': direction,
                'reversal': reversal,
                'reversal_type': reversal_type
            }
            
            prev_score = score
            if direction != 'FLAT':
                prev_direction = direction
    
    # Tính cho W1-W4 (ước tính)
    t5_score = daily_scores.get('T5', {}).get('score', t0_score)
    weekly_growth = (t5_score - t0_score) / 5 if t0_score != 0 else 0
    
    for i in range(1, 5):
        key = f'W{i}'
        # Ước tính: mỗi tuần = 5 ngày giao dịch
        estimated_score = t0_score + weekly_growth * 5 * i * 0.8  # Giảm dần
        change = estimated_score - t0_score
        change_pct = (change / t0_score * 100) if t0_score != 0 else 0
        
        daily_scores[key] = {
            'score': round(estimated_score, 2),
            'change': round(change, 2),
            'change_pct': round(change_pct, 2),
            'direction': 'UP' if change > 0 else ('DOWN' if change < 0 else 'FLAT'),
            'reversal': False,
            'reversal_type': None,
            'note': 'Ước tính'
        }
    
    # Tính cho M1-M3 (ước tính)
    for i in range(1, 4):
        key = f'M{i}'
        # Ước tính: mỗi tháng = 22 ngày giao dịch
        estimated_score = t0_score + weekly_growth * 22 * i * 0.5  # Giảm mạnh hơn
        change = estimated_score - t0_score
        change_pct = (change / t0_score * 100) if t0_score != 0 else 0
        
        daily_scores[key] = {
            'score': round(estimated_score, 2),
            'change': round(change, 2),
            'change_pct': round(change_pct, 2),
            'direction': 'UP' if change > 0 else ('DOWN' if change < 0 else 'FLAT'),
            'reversal': False,
            'reversal_type': None,
            'note': 'Ước tính dài hạn'
        }
    
    # Phân tích tổng thể
    reversals = [k for k, v in daily_scores.items() if v.get('reversal')]
    
    analysis = {
        'daily_scores': daily_scores,
        'reversals_detected': reversals,
        'overall_direction': forecast_results.get('summary', {}).get('overall_trend', 'N/A'),
        'confidence': forecast_results.get('summary', {}).get('overall_confidence', 50)
    }
    
    return analysis

# ============================================================
# PHẦN 4: XUẤT BÁO CÁO CHI TIẾT VỚI GIẢI THÍCH
# ============================================================

def generate_forecast_report_markdown(symbol, forecast_results, daily_composite, current_price):
    """
    Tạo báo cáo Markdown chi tiết với giải thích từng phương pháp
    
    Args:
        symbol: Mã cổ phiếu
        forecast_results: Kết quả từ forecast_all_26_indicators()
        daily_composite: Kết quả từ calculate_daily_composite_score()
        current_price: Giá hiện tại
    
    Returns:
        String Markdown
    """
    
    if not forecast_results or not daily_composite:
        return f"## {symbol}\n\n❌ Không có dữ liệu dự báo"
    
    summary = forecast_results.get('summary', {})
    individual = forecast_results.get('individual_forecasts', {})
    explanations = forecast_results.get('explanations', [])
    daily_scores = daily_composite.get('daily_scores', {})
    reversals = daily_composite.get('reversals_detected', [])
    
    report = []
    
    # ===== HEADER =====
    report.append(f"# 📊 BÁO CÁO DỰ BÁO CHI TIẾT: {symbol}")
    report.append(f"\n**Thời gian:** {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    report.append(f"\n**Giá hiện tại:** {current_price:,.0f}")
    report.append(f"\n**Số chỉ báo phân tích:** {summary.get('indicators_analyzed', 0)}")
    report.append("\n---\n")
    
    # ===== I. TÓM TẮT DỰ BÁO =====
    report.append("## I. TÓM TẮT DỰ BÁO\n")
    
    overall_trend = summary.get('overall_trend', 'N/A')
    confidence = summary.get('overall_confidence', 0)
    trend_breakdown = summary.get('trend_breakdown', {})
    
    # Emoji theo xu hướng
    if 'TĂNG' in overall_trend:
        trend_emoji = '🟢'
    elif 'GIẢM' in overall_trend:
        trend_emoji = '🔴'
    else:
        trend_emoji = '🟡'
    
    report.append(f"### {trend_emoji} XU HƯỚNG TỔNG HỢP: **{overall_trend}**\n")
    report.append(f"- **Độ tin cậy:** {confidence:.1f}%")
    report.append(f"- **Tỷ lệ tăng:** {trend_breakdown.get('up_pct', 0):.1f}%")
    report.append(f"- **Tỷ lệ giảm:** {trend_breakdown.get('down_pct', 0):.1f}%")
    report.append(f"- **Tỷ lệ trung lập:** {trend_breakdown.get('neutral_pct', 0):.1f}%")
    report.append("")
    
    # ===== II. BẢNG DỰ BÁO THEO NGÀY =====
    report.append("## II. DỰ BÁO THEO THỜI GIAN\n")
    
    # Bảng T0-T5
    report.append("### A. Ngắn hạn (T0-T5)\n")
    report.append("| Ngày | Điểm | Thay đổi | % | Hướng | Đảo chiều |")
    report.append("|:----:|:----:|:--------:|:-:|:-----:|:---------:|")
    
    for key in ['T0', 'T1', 'T2', 'T3', 'T4', 'T5']:
        data = daily_scores.get(key, {})
        score = data.get('score', '-')
        change = data.get('change', 0)
        change_pct = data.get('change_pct', 0)
        direction = data.get('direction', '-')
        reversal = data.get('reversal_type', '')
        
        # Format
        if isinstance(score, (int, float)):
            score_str = f"{score:.1f}"
        else:
            score_str = str(score)
        
        change_str = f"{change:+.2f}" if isinstance(change, (int, float)) else "-"
        pct_str = f"{change_pct:+.1f}%" if isinstance(change_pct, (int, float)) else "-"
        
        dir_emoji = '📈' if direction == 'UP' else ('📉' if direction == 'DOWN' else '➡️')
        reversal_str = f"⚠️ {reversal}" if reversal else "-"
        
        report.append(f"| {key} | {score_str} | {change_str} | {pct_str} | {dir_emoji} {direction} | {reversal_str} |")
    
    report.append("")
    
    # Bảng W1-W4
    report.append("### B. Trung hạn (W1-W4)\n")
    report.append("| Tuần | Điểm | Thay đổi | % | Hướng | Ghi chú |")
    report.append("|:----:|:----:|:--------:|:-:|:-----:|:-------:|")
    
    for key in ['W1', 'W2', 'W3', 'W4']:
        data = daily_scores.get(key, {})
        score = data.get('score', '-')
        change = data.get('change', 0)
        change_pct = data.get('change_pct', 0)
        direction = data.get('direction', '-')
        note = data.get('note', '-')
        
        score_str = f"{score:.1f}" if isinstance(score, (int, float)) else str(score)
        change_str = f"{change:+.2f}" if isinstance(change, (int, float)) else "-"
        pct_str = f"{change_pct:+.1f}%" if isinstance(change_pct, (int, float)) else "-"
        dir_emoji = '📈' if direction == 'UP' else ('📉' if direction == 'DOWN' else '➡️')
        
        report.append(f"| {key} | {score_str} | {change_str} | {pct_str} | {dir_emoji} {direction} | {note} |")
    
    report.append("")
    
    # Bảng M1-M3
    report.append("### C. Dài hạn (M1-M3)\n")
    report.append("| Tháng | Điểm | Thay đổi | % | Hướng | Ghi chú |")
    report.append("|:-----:|:----:|:--------:|:-:|:-----:|:-------:|")
    
    for key in ['M1', 'M2', 'M3']:
        data = daily_scores.get(key, {})
        score = data.get('score', '-')
        change = data.get('change', 0)
        change_pct = data.get('change_pct', 0)
        direction = data.get('direction', '-')
        note = data.get('note', '-')
        
        score_str = f"{score:.1f}" if isinstance(score, (int, float)) else str(score)
        change_str = f"{change:+.2f}" if isinstance(change, (int, float)) else "-"
        pct_str = f"{change_pct:+.1f}%" if isinstance(change_pct, (int, float)) else "-"
        dir_emoji = '📈' if direction == 'UP' else ('📉' if direction == 'DOWN' else '➡️')
        
        report.append(f"| {key} | {score_str} | {change_str} | {pct_str} | {dir_emoji} {direction} | {note} |")
    
    report.append("")
    
    # ===== III. PHÁT HIỆN ĐẢO CHIỀU =====
    report.append("## III. PHÁT HIỆN ĐẢO CHIỀU\n")
    
    if reversals:
        report.append(f"⚠️ **Phát hiện {len(reversals)} điểm đảo chiều tiềm năng:**\n")
        for rev_key in reversals:
            rev_data = daily_scores.get(rev_key, {})
            rev_type = rev_data.get('reversal_type', 'N/A')
            report.append(f"- **{rev_key}:** {rev_type}")
        report.append("")
    else:
        report.append("✅ Không phát hiện điểm đảo chiều trong kỳ dự báo.\n")
    
    # ===== IV. PHÂN TÍCH CHI TIẾT TỪNG CHỈ BÁO =====
    report.append("## IV. PHÂN TÍCH CHI TIẾT TỪNG CHỈ BÁO\n")
    
    # Sắp xếp theo confidence
    sorted_indicators = sorted(
        individual.items(),
        key=lambda x: x[1].get('final', {}).get('confidence', 0),
        reverse=True
    )
    
    for ind_name, ind_data in sorted_indicators[:10]:  # Top 10 chỉ báo
        final = ind_data.get('final', {})
        detailed = ind_data.get('detailed', {})
        
        current_val = final.get('current_value', 'N/A')
        trend = final.get('trend', 'N/A')
        confidence = final.get('confidence', 0)
        weight = final.get('weight', 1)
        
        # Emoji theo xu hướng
        if 'TĂNG' in str(trend) or 'UP' in str(trend):
            ind_emoji = '🟢'
        elif 'GIẢM' in str(trend) or 'DOWN' in str(trend):
            ind_emoji = '🔴'
        else:
            ind_emoji = '🟡'
        
        report.append(f"### {ind_emoji} {ind_name}\n")
        report.append(f"- **Giá trị hiện tại:** {current_val}")
        report.append(f"- **Xu hướng:** {trend}")
        report.append(f"- **Độ tin cậy:** {confidence:.1f}%")
        report.append(f"- **Trọng số:** {weight}")
        report.append("")
        
        # Dự báo T1-T5
        forecasts = final.get('forecasts', {})
        if forecasts:
            forecast_str = " → ".join([f"T{i}: {forecasts.get(f'T{i}', '-')}" for i in range(1, 6)])
            report.append(f"- **Dự báo:** {forecast_str}")
            report.append("")
        
        # Chi tiết các phương pháp
        methods = detailed.get('methods', {})
        if methods:
            report.append("**Phân tích từ 12 phương pháp:**\n")
            
            for method_name, method_data in methods.items():
                explanation = method_data.get('explanation', '')
                method_confidence = method_data.get('confidence', 0)
                
                if explanation:
                    # Truncate nếu quá dài
                    if len(explanation) > 150:
                        explanation = explanation[:147] + "..."
                    
                    conf_emoji = '🔹' if method_confidence >= 70 else ('🔸' if method_confidence >= 50 else '⚪')
                    report.append(f"  {conf_emoji} **{method_name}** ({method_confidence:.0f}%): {explanation}")
            
            report.append("")
        
        report.append("---\n")
    
    # ===== V. GIẢI THÍCH 12 PHƯƠNG PHÁP =====
    report.append("## V. GIẢI THÍCH 12 PHƯƠNG PHÁP DỰ BÁO\n")
    
    method_explanations = [
        ("1. Linear Regression", "Tìm đường thẳng khớp nhất với dữ liệu lịch sử", "Xác định xu hướng chính (tăng/giảm/đi ngang)", "Slope > 0: Tăng, Slope < 0: Giảm, R² cao: Khớp tốt"),
        ("2. Polynomial Regression", "Tìm đường cong bậc 2 (parabola) khớp với dữ liệu", "Phát hiện điểm uốn, đỉnh/đáy, đảo chiều", "a > 0: Lõm (đáy), a < 0: Lồi (đỉnh)"),
        ("3. First Derivative", "Tính đạo hàm bậc 1 (tốc độ thay đổi)", "Biết chỉ báo đang tăng/giảm nhanh cỡ nào", "Velocity > 0: Đang tăng, < 0: Đang giảm"),
        ("4. Second Derivative", "Tính đạo hàm bậc 2 (gia tốc)", "Biết xu hướng đang tăng tốc hay chậm lại", "Acceleration đổi dấu: Sắp đảo chiều"),
        ("5. Peak/Trough Detection", "Tìm các đỉnh và đáy trong lịch sử", "Xác định pha hiện tại, vùng đảo chiều", "Gần đáy lịch sử: Có thể tăng, Gần đỉnh: Có thể giảm"),
        ("6. Multi-Timeframe", "Phân tích xu hướng 10d, 30d, toàn bộ", "So sánh xu hướng ngắn/trung/dài hạn", "Đồng thuận: Mạnh, Phân kỳ: Cẩn thận"),
        ("7. Pattern Matching", "So sánh mô hình hiện tại với lịch sử", "Tìm tình huống tương tự để dự đoán", "Correlation > 90%: Rất tương tự"),
        ("8. Fourier Transform", "Phân tích chu kỳ bằng biến đổi Fourier", "Phát hiện chu kỳ lặp lại, dự đoán đỉnh/đáy", "Chu kỳ rõ: Dự đoán tốt hơn"),
        ("9. Probability & Statistics", "Phân tích xác suất, Z-score, percentile", "Tính xác suất mean reversion", "Z > 2: Quá cao, Z < -2: Quá thấp"),
        ("10. Fibonacci Levels", "Tính các mức Fibonacci retracement", "Xác định vùng hỗ trợ/kháng cự tự nhiên", "38.2%, 50%, 61.8% là các mức quan trọng"),
        ("11. Logical Rules", "Áp dụng quy tắc logic AND/OR", "Kết hợp nhiều điều kiện để quyết định", "Oversold + Trend up: Mua, Overbought + Trend down: Bán"),
        ("12. ML Ensemble", "Kết hợp 11 phương pháp bằng voting", "Tăng độ chính xác bằng đa số", "Majority vote > 60%: Tin cậy cao"),
    ]
    
    report.append("| # | Phương pháp | Mô tả | Công dụng | Cách đọc |")
    report.append("|:-:|:------------|:------|:----------|:---------|")
    
    for method in method_explanations:
        report.append(f"| {method[0][:2]} | {method[0][3:]} | {method[1]} | {method[2]} | {method[3]} |")
    
    report.append("")
    
    # ===== VI. KẾT LUẬN VÀ KHUYẾN NGHỊ =====
    report.append("## VI. KẾT LUẬN VÀ KHUYẾN NGHỊ\n")
    
    # Tổng hợp từ các chỉ báo quan trọng
    important_indicators = ['RSI', 'MACD_Hist', 'Stoch_K', 'MFI', 'ADX']
    signals = {'buy': 0, 'sell': 0, 'hold': 0}
    
    for ind_name in important_indicators:
        if ind_name in individual:
            trend = individual[ind_name].get('final', {}).get('trend', '')
            if 'TĂNG' in str(trend) or 'UP' in str(trend) or 'MUA' in str(trend):
                signals['buy'] += 1
            elif 'GIẢM' in str(trend) or 'DOWN' in str(trend) or 'BÁN' in str(trend):
                signals['sell'] += 1
            else:
                signals['hold'] += 1
    
    report.append(f"### Tín hiệu từ 5 chỉ báo chính:")
    report.append(f"- 🟢 **MUA:** {signals['buy']}/5")
    report.append(f"- 🔴 **BÁN:** {signals['sell']}/5")
    report.append(f"- 🟡 **GIỮ:** {signals['hold']}/5")
    report.append("")
    
    # Khuyến nghị
    if signals['buy'] >= 4:
        recommendation = "🟢 **MUA MẠNH** - Đa số chỉ báo đồng thuận tích cực"
    elif signals['buy'] >= 3:
        recommendation = "🟢 **MUA** - Tín hiệu tích cực chiếm ưu thế"
    elif signals['sell'] >= 4:
        recommendation = "🔴 **BÁN MẠNH** - Đa số chỉ báo đồng thuận tiêu cực"
    elif signals['sell'] >= 3:
        recommendation = "🔴 **BÁN** - Tín hiệu tiêu cực chiếm ưu thế"
    else:
        recommendation = "🟡 **THEO DÕI** - Tín hiệu hỗn hợp, chờ xác nhận"
    
    report.append(f"### Khuyến nghị: {recommendation}\n")
    
    # Lưu ý
    report.append("### ⚠️ Lưu ý quan trọng:")
    report.append("- Dự báo dựa trên phân tích kỹ thuật, không phải lời khuyên đầu tư")
    report.append("- Độ tin cậy giảm dần theo thời gian (T1 > T5 > W4 > M3)")
    report.append("- Luôn kết hợp với phân tích cơ bản và quản lý rủi ro")
    report.append("- Đặt stop-loss để bảo vệ vốn")
    report.append("")
    
    return "\n".join(report)


def generate_forecast_report_for_word(symbol, forecast_results, daily_composite, current_price):
    """
    Tạo nội dung báo cáo cho file Word
    Trả về dict để dễ format trong python-docx
    """
    
    if not forecast_results or not daily_composite:
        return {'error': 'Không có dữ liệu'}
    
    summary = forecast_results.get('summary', {})
    individual = forecast_results.get('individual_forecasts', {})
    daily_scores = daily_composite.get('daily_scores', {})
    reversals = daily_composite.get('reversals_detected', [])
    
    report = {
        'symbol': symbol,
        'timestamp': datetime.now().strftime('%d/%m/%Y %H:%M'),
        'current_price': current_price,
        'summary': {
            'overall_trend': summary.get('overall_trend', 'N/A'),
            'confidence': summary.get('overall_confidence', 0),
            'indicators_analyzed': summary.get('indicators_analyzed', 0),
            'trend_breakdown': summary.get('trend_breakdown', {})
        },
        'daily_forecasts': {
            'short_term': {key: daily_scores.get(key, {}) for key in ['T0', 'T1', 'T2', 'T3', 'T4', 'T5']},
            'mid_term': {key: daily_scores.get(key, {}) for key in ['W1', 'W2', 'W3', 'W4']},
            'long_term': {key: daily_scores.get(key, {}) for key in ['M1', 'M2', 'M3']}
        },
        'reversals': reversals,
        'top_indicators': [],
        'method_summary': []
    }
    
    # Top indicators
    sorted_indicators = sorted(
        individual.items(),
        key=lambda x: x[1].get('final', {}).get('confidence', 0),
        reverse=True
    )
    
    for ind_name, ind_data in sorted_indicators[:5]:
        final = ind_data.get('final', {})
        report['top_indicators'].append({
            'name': ind_name,
            'current_value': final.get('current_value'),
            'trend': final.get('trend'),
            'confidence': final.get('confidence'),
            'forecasts': final.get('forecasts', {})
        })
    
    return report


def export_forecast_to_excel_sheet(writer, symbol, forecast_results, daily_composite):
    """
    Xuất kết quả dự báo vào sheet Excel
    
    Args:
        writer: pd.ExcelWriter object
        symbol: Mã cổ phiếu
        forecast_results: Kết quả dự báo
        daily_composite: Điểm tổng hợp theo ngày
    """
    
    if not forecast_results or not daily_composite:
        return
    
    daily_scores = daily_composite.get('daily_scores', {})
    individual = forecast_results.get('individual_forecasts', {})
    
    # Sheet 1: Dự báo theo ngày
    daily_data = []
    for key in ['T0', 'T1', 'T2', 'T3', 'T4', 'T5', 'W1', 'W2', 'W3', 'W4', 'M1', 'M2', 'M3']:
        data = daily_scores.get(key, {})
        daily_data.append({
            'Thời gian': key,
            'Điểm': data.get('score', '-'),
            'Thay đổi': data.get('change', '-'),
            '% Thay đổi': data.get('change_pct', '-'),
            'Hướng': data.get('direction', '-'),
            'Đảo chiều': data.get('reversal_type', '-'),
            'Ghi chú': data.get('note', '-')
        })
    
    df_daily = pd.DataFrame(daily_data)
    df_daily.to_excel(writer, sheet_name=f'{symbol}_DuBao', index=False)
    
    # Sheet 2: Chi tiết chỉ báo
    indicator_data = []
    for ind_name, ind_data in individual.items():
        final = ind_data.get('final', {})
        indicator_data.append({
            'Chỉ báo': ind_name,
            'Giá trị hiện tại': final.get('current_value', '-'),
            'Xu hướng': final.get('trend', '-'),
            'Độ tin cậy': final.get('confidence', '-'),
            'Trọng số': final.get('weight', '-'),
            'T1': final.get('forecasts', {}).get('T1', '-'),
            'T2': final.get('forecasts', {}).get('T2', '-'),
            'T3': final.get('forecasts', {}).get('T3', '-'),
            'T4': final.get('forecasts', {}).get('T4', '-'),
            'T5': final.get('forecasts', {}).get('T5', '-'),
        })
    
    df_indicators = pd.DataFrame(indicator_data)
    df_indicators.to_excel(writer, sheet_name=f'{symbol}_ChiBao', index=False)


# ============================================================
# CẤU HÌNH
# ============================================================

DEFAULT_SYMBOLS = ["VCK", "DGW", "VNM", "FPT", "VIC", "HPG", "MWG", "TCB", "VCB", "ACB", "VPB"]

BASE_OUTPUT_DIR = "./output"

# 26 Chỉ báo chia theo nhóm
INDICATOR_GROUPS = {
    '📈 XU HƯỚNG': {
        'SMA': {'name': 'SMA (5,10,20,50,100,200)', 'default': True},
        'EMA': {'name': 'EMA (12,26,50)', 'default': True},
        'WMA': {'name': 'WMA (10,20)', 'default': False},
        'TEMA': {'name': 'TEMA (20)', 'default': False},
        'DEMA': {'name': 'DEMA (20)', 'default': False},
        'MACD': {'name': 'MACD (12,26,9) + Crossover', 'default': True},
        'SAR': {'name': 'Parabolic SAR', 'default': False},
    },
    '⚡ ĐỘNG LƯỢNG': {
        'RSI': {'name': 'RSI (14)', 'default': True},
        'STOCH': {'name': 'Stochastic %K/%D', 'default': True},
        'STOCHRSI': {'name': 'Stochastic RSI', 'default': False},
        'ROC': {'name': 'ROC (10)', 'default': False},
        'MOM': {'name': 'Momentum (10)', 'default': False},
    },
    '🔄 DAO ĐỘNG': {
        'CCI': {'name': 'CCI (20)', 'default': False},
        'WILLR': {'name': 'Williams %R', 'default': False},
        'ADX': {'name': 'ADX (14)', 'default': False},
        'ATR': {'name': 'ATR (14)', 'default': True},
        'BB': {'name': 'Bollinger Bands', 'default': True},
    },
    '📊 KHỐI LƯỢNG': {
        'OBV': {'name': 'OBV', 'default': True},
        'MFI': {'name': 'MFI (14)', 'default': True},
        'CMF': {'name': 'CMF (20)', 'default': False},
        'AD': {'name': 'A/D Line', 'default': False},
        'VWAP': {'name': 'VWAP', 'default': False},
        'FI': {'name': 'Force Index', 'default': False},
    },
}

# ============================================================
# CHỈ BÁO BẮT BUỘC (AUTO BẬT)
# ============================================================

REQUIRED_INDICATORS = ['RSI', 'MACD', 'STOCH', 'BB', 'SMA', 'ATR', 'MFI', 'OBV']

# ============================================================
# TRỌNG SỐ CHỈ BÁO (để tính điểm có trọng số)
# ============================================================

INDICATOR_WEIGHTS = {
    # Nhóm XU HƯỚNG (tổng ~26)
    'SMA': 5,
    'EMA': 4,
    'WMA': 2,
    'TEMA': 2,
    'DEMA': 2,
    'MACD': 8,
    'SAR': 3,
    
    # Nhóm ĐỘNG LƯỢNG (tổng ~25)
    'RSI': 8,
    'STOCH': 6,
    'STOCHRSI': 4,
    'ROC': 3,
    'MOM': 4,
    
    # Nhóm DAO ĐỘNG (tổng ~25)
    'CCI': 4,
    'WILLR': 4,
    'ADX': 6,
    'ATR': 5,
    'BB': 6,
    
    # Nhóm KHỐI LƯỢNG (tổng ~20)
    'OBV': 5,
    'MFI': 5,
    'CMF': 3,
    'AD': 3,
    'VWAP': 2,
    'FI': 2
}

# Ngưỡng hành động
ACTION_THRESHOLDS = {
    'mua_manh': 70,
    'mua': 55,
    'can_nhac_mua': 40,
    'trung_lap': 30,
    'can_nhac_ban': 20,
    'ban': 10,
    'ban_manh': 0
}

# ============================================================
# HÀM LẤY DỮ LIỆU
# ============================================================

def get_stock_data(symbol, start_date, end_date):
    """Lấy dữ liệu cổ phiếu từ vnstock"""
    try:
        from vnstock import Vnstock
        stock = Vnstock().stock(symbol=symbol, source='VCI')
        df = stock.quote.history(start=start_date, end=end_date, interval='1D')
        return df
    except Exception as e:
        print(f"Lỗi lấy {symbol}: {e}")
        return None


def get_all_symbols():
    """Lấy danh sách tất cả mã cổ phiếu"""
    try:
        from vnstock import Vnstock
        stock = Vnstock().stock(symbol='VNM', source='VCI')
        df_list = stock.listing.all_symbols()
        if df_list is not None and len(df_list) > 0:
            if 'symbol' in df_list.columns:
                return sorted(df_list['symbol'].tolist())
            elif 'ticker' in df_list.columns:
                return sorted(df_list['ticker'].tolist())
        return DEFAULT_SYMBOLS
    except:
        return DEFAULT_SYMBOLS


def get_market_indices():
    """Lấy chỉ số thị trường VN-Index, VN30, HNX30"""
    indices = {'VNINDEX': 'VN-Index', 'VN30': 'VN30', 'HNX30': 'HNX30'}
    data = []
    
    start_dt = (datetime.now() - timedelta(days=400)).strftime('%Y-%m-%d')
    end_dt = datetime.now().strftime('%Y-%m-%d')
    
    for code, name in indices.items():
        try:
            from vnstock import Vnstock
            stock = Vnstock().stock(symbol=code, source='VCI')
            df = stock.quote.history(start=start_dt, end=end_dt, interval='1D')
            
            if df is not None and len(df) > 0:
                df['time'] = pd.to_datetime(df['time'])
                df = df.sort_values('time', ascending=False)
                lat = df.iloc[0]
                cur = lat['close']
                
                def pct(d):
                    if len(df) > d and df.iloc[d]['close'] > 0:
                        return round((cur - df.iloc[d]['close']) / df.iloc[d]['close'] * 100, 2)
                    return None
                
                data.append({
                    'Chỉ số': name,
                    'Giá': round(cur, 2),
                    'Ngày': lat['time'].strftime('%d/%m/%Y'),
                    'D%': pct(1),
                    'W%': pct(5),
                    'M%': pct(22),
                    'Q%': pct(66),
                    'Y%': pct(252)
                })
            time.sleep(0.3)
        except Exception as e:
            print(f"Lỗi {code}: {e}")
    
    return pd.DataFrame(data)


def analyze_market(df_m):
    """Phân tích xu hướng thị trường"""
    if df_m is None or len(df_m) == 0:
        return "Không có dữ liệu", "neutral"
    
    vni = df_m[df_m['Chỉ số'] == 'VN-Index']
    if len(vni) == 0:
        return "Không có VN-Index", "neutral"
    
    d = vni.iloc[0].get('D%', 0) or 0
    w = vni.iloc[0].get('W%', 0) or 0
    
    if d > 1 and w > 3:
        return "🟢 TĂNG MẠNH", "bullish"
    elif d > 0 and w > 0:
        return "🟢 TĂNG", "slightly_bullish"
    elif d < -1 and w < -3:
        return "🔴 GIẢM MẠNH", "bearish"
    elif d < 0 and w < 0:
        return "🔴 GIẢM", "slightly_bearish"
    else:
        return "🟡 TÍCH LŨY", "neutral"


# ============================================================
# HÀM TÍNH 26 CHỈ BÁO KỸ THUẬT
# ============================================================

def calculate_indicators(df, selected):
    """Tính các chỉ báo đã chọn"""
    results = {}
    c = df['close'].astype(float)
    h = df['high'].astype(float)
    l = df['low'].astype(float)
    v = df['volume'].astype(float)
    n = len(df)
    
    # SMA
    if 'SMA' in selected:
        for p in [5, 10, 20, 50, 100, 200]:
            if n >= p:
                results[f'SMA_{p}'] = c.rolling(p).mean().round(2)
    
    # EMA
    if 'EMA' in selected:
        for p in [12, 26, 50]:
            if n >= p:
                results[f'EMA_{p}'] = c.ewm(span=p, adjust=False).mean().round(2)
    
    # WMA
    if 'WMA' in selected:
        for p in [10, 20]:
            if n >= p:
                weights = np.arange(1, p + 1)
                results[f'WMA_{p}'] = c.rolling(p).apply(lambda x: np.dot(x, weights) / weights.sum(), raw=True).round(2)
    
    # TEMA
    if 'TEMA' in selected and n >= 20:
        e1 = c.ewm(span=20, adjust=False).mean()
        e2 = e1.ewm(span=20, adjust=False).mean()
        e3 = e2.ewm(span=20, adjust=False).mean()
        results['TEMA_20'] = (3 * e1 - 3 * e2 + e3).round(2)
    
    # DEMA
    if 'DEMA' in selected and n >= 20:
        e1 = c.ewm(span=20, adjust=False).mean()
        e2 = e1.ewm(span=20, adjust=False).mean()
        results['DEMA_20'] = (2 * e1 - e2).round(2)
    
    # MACD (12, 26, 9)
    if 'MACD' in selected and n >= 26:
        e12 = c.ewm(span=12, adjust=False).mean()
        e26 = c.ewm(span=26, adjust=False).mean()
        macd = e12 - e26
        sig = macd.ewm(span=9, adjust=False).mean()
        hist = macd - sig
        
        results['MACD'] = macd.round(3)
        results['MACD_Signal'] = sig.round(3)
        results['MACD_Hist'] = hist.round(3)
        
        # Crossover
        h_prev = hist.shift(1)
        cross = pd.Series('', index=c.index)
        cross[(h_prev < 0) & (hist > 0)] = 'CẮT_LÊN'
        cross[(h_prev > 0) & (hist < 0)] = 'CẮT_XUỐNG'
        results['MACD_Cross'] = cross
    
    # Parabolic SAR
    if 'SAR' in selected and n >= 5:
        sar = pd.Series(index=c.index, dtype=float)
        sar.iloc[0] = l.iloc[0]
        af, af_max = 0.02, 0.2
        ep = h.iloc[0]
        trend = 1
        cur_af = af
        for i in range(1, n):
            if trend == 1:
                sar.iloc[i] = sar.iloc[i-1] + cur_af * (ep - sar.iloc[i-1])
                if h.iloc[i] > ep:
                    ep = h.iloc[i]
                    cur_af = min(cur_af + af, af_max)
                if l.iloc[i] < sar.iloc[i]:
                    trend = -1
                    sar.iloc[i] = ep
                    ep = l.iloc[i]
                    cur_af = af
            else:
                sar.iloc[i] = sar.iloc[i-1] + cur_af * (ep - sar.iloc[i-1])
                if l.iloc[i] < ep:
                    ep = l.iloc[i]
                    cur_af = min(cur_af + af, af_max)
                if h.iloc[i] > sar.iloc[i]:
                    trend = 1
                    sar.iloc[i] = ep
                    ep = h.iloc[i]
                    cur_af = af
        results['SAR'] = sar.round(2)
    
    # RSI (14)
    if 'RSI' in selected and n >= 14:
        delta = c.diff()
        gain = (delta.where(delta > 0, 0)).rolling(14).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(14).mean()
        results['RSI'] = (100 - 100 / (1 + gain / loss)).round(2)
    
    # Stochastic
    if 'STOCH' in selected and n >= 14:
        lo = l.rolling(14).min()
        hi = h.rolling(14).max()
        results['Stoch_K'] = (100 * (c - lo) / (hi - lo)).round(2)
        results['Stoch_D'] = results['Stoch_K'].rolling(3).mean().round(2)
    
    # Stochastic RSI
    if 'STOCHRSI' in selected and n >= 28:
        delta = c.diff()
        gain = (delta.where(delta > 0, 0)).rolling(14).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(14).mean()
        rsi = 100 - 100 / (1 + gain / loss)
        lo_rsi = rsi.rolling(14).min()
        hi_rsi = rsi.rolling(14).max()
        denom = (hi_rsi - lo_rsi).replace(0, np.nan)
        results['StochRSI'] = ((rsi - lo_rsi) / denom * 100).round(2)
    
    # ROC (10)
    if 'ROC' in selected and n >= 10:
        results['ROC'] = ((c - c.shift(10)) / c.shift(10) * 100).round(2)
    
    # Momentum (10)
    if 'MOM' in selected and n >= 10:
        results['Momentum'] = (c - c.shift(10)).round(2)
    
    # CCI (20)
    if 'CCI' in selected and n >= 20:
        tp = (h + l + c) / 3
        sma_tp = tp.rolling(20).mean()
        mad = tp.rolling(20).apply(lambda x: np.abs(x - x.mean()).mean(), raw=True)
        results['CCI'] = ((tp - sma_tp) / (0.015 * mad)).round(2)
    
    # Williams %R (14)
    if 'WILLR' in selected and n >= 14:
        hi = h.rolling(14).max()
        lo = l.rolling(14).min()
        results['Williams_R'] = (-100 * (hi - c) / (hi - lo)).round(2)
    
    # ADX (14)
    if 'ADX' in selected and n >= 28:
        pc = c.shift(1)
        tr = pd.concat([h - l, abs(h - pc), abs(l - pc)], axis=1).max(axis=1)
        pdm = h.diff().clip(lower=0)
        mdm = (-l.diff()).clip(lower=0)
        pdm = pdm.where(pdm > mdm, 0)
        mdm = mdm.where(mdm > pdm, 0)
        atr = tr.ewm(span=14, adjust=False).mean()
        pdi = 100 * pdm.ewm(span=14, adjust=False).mean() / atr
        mdi = 100 * mdm.ewm(span=14, adjust=False).mean() / atr
        dx = 100 * abs(pdi - mdi) / (pdi + mdi)
        results['ADX'] = dx.ewm(span=14, adjust=False).mean().round(2)
        results['Plus_DI'] = pdi.round(2)
        results['Minus_DI'] = mdi.round(2)
    
    # ATR (14)
    if 'ATR' in selected and n >= 14:
        pc = c.shift(1)
        tr = pd.concat([h - l, abs(h - pc), abs(l - pc)], axis=1).max(axis=1)
        results['ATR'] = tr.ewm(span=14, adjust=False).mean().round(2)
    
    # Bollinger Bands (20, 2)
    if 'BB' in selected and n >= 20:
        sma = c.rolling(20).mean()
        std = c.rolling(20).std()
        results['BB_Upper'] = (sma + 2 * std).round(2)
        results['BB_Middle'] = sma.round(2)
        results['BB_Lower'] = (sma - 2 * std).round(2)
    
    # OBV
    if 'OBV' in selected:
        obv = pd.Series(index=c.index, dtype=float)
        obv.iloc[0] = v.iloc[0]
        for i in range(1, n):
            if c.iloc[i] > c.iloc[i - 1]:
                obv.iloc[i] = obv.iloc[i - 1] + v.iloc[i]
            elif c.iloc[i] < c.iloc[i - 1]:
                obv.iloc[i] = obv.iloc[i - 1] - v.iloc[i]
            else:
                obv.iloc[i] = obv.iloc[i - 1]
        results['OBV'] = obv.round(0)
    
    # MFI (14)
    if 'MFI' in selected and n >= 14:
        tp = (h + l + c) / 3
        mf = tp * v
        pmf = mf.where(tp > tp.shift(1), 0).rolling(14).sum()
        nmf = mf.where(tp < tp.shift(1), 0).rolling(14).sum()
        results['MFI'] = (100 - 100 / (1 + pmf / nmf.replace(0, np.nan))).round(2)
    
    # CMF (20)
    if 'CMF' in selected and n >= 20:
        denom = (h - l).replace(0, np.nan)
        mfv = ((c - l) - (h - c)) / denom * v
        results['CMF'] = (mfv.rolling(20).sum() / v.rolling(20).sum()).round(4)
    
    # A/D Line
    if 'AD' in selected:
        denom = (h - l).replace(0, np.nan)
        clv = ((c - l) - (h - c)) / denom
        results['AD'] = (clv.fillna(0) * v).cumsum().round(0)
    
    # VWAP
    if 'VWAP' in selected:
        tp = (h + l + c) / 3
        results['VWAP'] = ((tp * v).cumsum() / v.cumsum()).round(2)
    
    # Force Index (13)
    if 'FI' in selected and n >= 13:
        fi = (c - c.shift(1)) * v
        results['FI'] = fi.ewm(span=13, adjust=False).mean().round(0)
    
    return results


# ============================================================
# HÀM XÁC ĐỊNH PHA THỊ TRƯỜNG (GIỮ LẠI HÀM RIÊNG)
# ============================================================

def detect_market_phase(df, rsi, stoch_k, macd_hist, close, sma_20, sma_50):
    """Xác định pha: ĐÁY / TĂNG / ĐỈNH / GIẢM / TÍCH_LŨY"""
    
    # Vị trí giá trong range 20 ngày
    high_20d = df['high'].tail(20).max()
    low_20d = df['low'].tail(20).min()
    price_range = high_20d - low_20d
    price_position = (close - low_20d) / price_range * 100 if price_range > 0 else 50
    
    # Trend ngắn hạn
    closes = df['close'].values
    n = len(closes)
    trend_3d = (closes[-1] - closes[-3]) / closes[-3] * 100 if n >= 3 else 0
    trend_5d = (closes[-1] - closes[-5]) / closes[-5] * 100 if n >= 5 else 0
    
    # MACD trend
    macd_hist_trend = 0
    if 'MACD_Hist' in df.columns and len(df) >= 3:
        hists = df['MACD_Hist'].tail(3).values
        if all(pd.notna(hists)):
            macd_hist_trend = 1 if hists[-1] > hists[-2] else -1
    
    confidence = 50
    
    # Logic xác định pha
    if price_position < 25 and rsi < 35 and stoch_k < 30:
        if trend_3d > -0.5 or macd_hist_trend > 0:
            phase = 'ĐÁY'
            confidence = min(90, 50 + (35 - rsi) + (30 - stoch_k) / 2)
        else:
            phase = 'GIẢM'
            confidence = 60
    elif price_position > 75 and rsi > 65 and stoch_k > 70:
        if trend_3d < 0.5 or macd_hist_trend < 0:
            phase = 'ĐỈNH'
            confidence = min(90, 50 + (rsi - 65) + (stoch_k - 70) / 2)
        else:
            phase = 'TĂNG'
            confidence = 60
    elif trend_3d > 1.5 and trend_5d > 2 and 45 < rsi < 70:
        phase = 'TĂNG'
        confidence = min(80, 50 + trend_3d * 3)
    elif trend_3d < -1.5 and trend_5d < -2 and 30 < rsi < 55:
        phase = 'GIẢM'
        confidence = min(80, 50 + abs(trend_3d) * 3)
    else:
        phase = 'TÍCH_LŨY'
        confidence = 50
    
    return phase, round(confidence, 1)

# ============================================================
# HÀM PHÂN TÍCH CHI TIẾT TÍN HIỆU VÀ HÀNH ĐỘNG
# ============================================================

def analyze_indicator_signals(ind, available, close, df):
    """
    Phân tích chi tiết tín hiệu từ từng chỉ báo
    Trả về: signals_table, buy_group, sell_group, hold_group, conflicts, recommendations
    """
    
    signals_table = []  # Bảng tín hiệu từng chỉ báo
    buy_group = []      # Nhóm tín hiệu MUA
    sell_group = []     # Nhóm tín hiệu BÁN
    hold_group = []     # Nhóm tín hiệu GIỮ
    conflicts = []      # Xung đột phát hiện
    
    buy_points = 0
    sell_points = 0
    hold_points = 0
    
    # ===== 1. RSI =====
    if available.get('RSI'):
        rsi = ind['RSI']
        if rsi < 30:
            signals_table.append({'indicator': 'RSI', 'value': f'{rsi:.0f}', 'signal': '📉 QUÁ BÁN', 'signal_type': 'oversold', 'interpretation': 'GIÁ RẺ / CƠ HỘI MUA'})
            buy_group.append(f"RSI={rsi:.0f} quá bán (+2)")
            buy_points += 2
        elif rsi < 40:
            signals_table.append({'indicator': 'RSI', 'value': f'{rsi:.0f}', 'signal': '📉 VÙNG THẤP', 'signal_type': 'low', 'interpretation': 'Có thể tích lũy'})
            buy_group.append(f"RSI={rsi:.0f} vùng thấp (+1)")
            buy_points += 1
        elif rsi > 70:
            signals_table.append({'indicator': 'RSI', 'value': f'{rsi:.0f}', 'signal': '📈 QUÁ MUA', 'signal_type': 'overbought', 'interpretation': 'GIÁ CAO / CẨN THẬN'})
            sell_group.append(f"RSI={rsi:.0f} quá mua (+2)")
            sell_points += 2
        elif rsi > 60:
            signals_table.append({'indicator': 'RSI', 'value': f'{rsi:.0f}', 'signal': '📈 VÙNG CAO', 'signal_type': 'high', 'interpretation': 'Cẩn thận chốt lời'})
            sell_group.append(f"RSI={rsi:.0f} vùng cao (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'RSI', 'value': f'{rsi:.0f}', 'signal': '⚪ TRUNG LẬP', 'signal_type': 'neutral', 'interpretation': 'Chờ tín hiệu rõ hơn'})
            hold_group.append(f"RSI={rsi:.0f} trung lập (+1)")
            hold_points += 1
    
    # ===== 2. Stochastic =====
    if available.get('STOCH'):
        stoch_k = ind['Stoch_K']
        stoch_d = ind.get('Stoch_D', stoch_k)
        
        if stoch_k < 20:
            signals_table.append({'indicator': 'Stochastic', 'value': f'{stoch_k:.0f}', 'signal': '📉 QUÁ BÁN', 'signal_type': 'oversold', 'interpretation': 'GIÁ RẺ / CƠ HỘI MUA'})
            buy_group.append(f"Stochastic={stoch_k:.0f} quá bán (+2)")
            buy_points += 2
        elif stoch_k < 40:
            signals_table.append({'indicator': 'Stochastic', 'value': f'{stoch_k:.0f}', 'signal': '📉 VÙNG THẤP', 'signal_type': 'low', 'interpretation': 'Có thể tích lũy'})
            buy_group.append(f"Stochastic={stoch_k:.0f} vùng thấp (+1)")
            buy_points += 1
        elif stoch_k > 80:
            signals_table.append({'indicator': 'Stochastic', 'value': f'{stoch_k:.0f}', 'signal': '📈 QUÁ MUA', 'signal_type': 'overbought', 'interpretation': 'GIÁ CAO / CẨN THẬN'})
            sell_group.append(f"Stochastic={stoch_k:.0f} quá mua (+2)")
            sell_points += 2
        elif stoch_k > 60:
            signals_table.append({'indicator': 'Stochastic', 'value': f'{stoch_k:.0f}', 'signal': '📈 VÙNG CAO', 'signal_type': 'high', 'interpretation': 'Cẩn thận chốt lời'})
            sell_group.append(f"Stochastic={stoch_k:.0f} vùng cao (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'Stochastic', 'value': f'{stoch_k:.0f}', 'signal': '⚪ TRUNG LẬP', 'signal_type': 'neutral', 'interpretation': 'Chờ tín hiệu rõ hơn'})
            hold_group.append(f"Stochastic={stoch_k:.0f} trung lập (+1)")
            hold_points += 1
        
        # K/D Crossover
        if stoch_k > stoch_d:
            signals_table.append({'indicator': 'Stoch K/D', 'value': f'K>{stoch_d:.0f}', 'signal': '🔼 K CẮT LÊN D', 'signal_type': 'bullish', 'interpretation': 'Tín hiệu tăng ngắn hạn'})
        elif stoch_k < stoch_d:
            signals_table.append({'indicator': 'Stoch K/D', 'value': f'K<{stoch_d:.0f}', 'signal': '🔽 K CẮT XUỐNG D', 'signal_type': 'bearish', 'interpretation': 'Tín hiệu giảm ngắn hạn'})
    
    # ===== 3. StochRSI =====
    if available.get('STOCHRSI'):
        stoch_rsi = ind['StochRSI']
        if stoch_rsi < 20:
            signals_table.append({'indicator': 'StochRSI', 'value': f'{stoch_rsi:.0f}', 'signal': '📉 QUÁ BÁN', 'signal_type': 'oversold', 'interpretation': 'GIÁ RẺ / CƠ HỘI MUA'})
            buy_group.append(f"StochRSI={stoch_rsi:.0f} quá bán (+2)")
            buy_points += 2
        elif stoch_rsi > 80:
            signals_table.append({'indicator': 'StochRSI', 'value': f'{stoch_rsi:.0f}', 'signal': '📈 QUÁ MUA', 'signal_type': 'overbought', 'interpretation': 'GIÁ CAO / CẨN THẬN'})
            sell_group.append(f"StochRSI={stoch_rsi:.0f} quá mua (+2)")
            sell_points += 2
        else:
            signals_table.append({'indicator': 'StochRSI', 'value': f'{stoch_rsi:.0f}', 'signal': '⚪ TRUNG LẬP', 'signal_type': 'neutral', 'interpretation': 'Chờ tín hiệu'})
            hold_points += 1
    
    # ===== 4. MACD =====
    if available.get('MACD'):
        macd = ind['MACD']
        macd_signal = ind.get('MACD_Signal', 0)
        macd_hist = ind.get('MACD_Hist', 0)
        
        # MACD Histogram
        if macd_hist > 0:
            signals_table.append({'indicator': 'MACD Hist', 'value': f'{macd_hist:.3f}', 'signal': '📈 BULLISH', 'signal_type': 'bullish', 'interpretation': 'Động lượng tăng'})
            buy_group.append(f"MACD Histogram > 0 (+1)")
            buy_points += 1
        else:
            signals_table.append({'indicator': 'MACD Hist', 'value': f'{macd_hist:.3f}', 'signal': '📉 BEARISH', 'signal_type': 'bearish', 'interpretation': 'Động lượng giảm'})
            sell_group.append(f"MACD Histogram < 0 (+1)")
            sell_points += 1
        
        # MACD Crossover
        if len(df) > 1 and 'MACD_Hist' in df.columns:
            prev_hist = df['MACD_Hist'].iloc[-2] if pd.notna(df['MACD_Hist'].iloc[-2]) else 0
            if prev_hist < 0 and macd_hist > 0:
                signals_table.append({'indicator': 'MACD Cross', 'value': 'Cross Up', 'signal': '🔼 CẮT LÊN', 'signal_type': 'bullish_cross', 'interpretation': '🚀 TÍN HIỆU MUA MẠNH'})
                buy_group.append(f"🔼 MACD Cross Up (+2)")
                buy_points += 2
            elif prev_hist > 0 and macd_hist < 0:
                signals_table.append({'indicator': 'MACD Cross', 'value': 'Cross Down', 'signal': '🔽 CẮT XUỐNG', 'signal_type': 'bearish_cross', 'interpretation': '⚠️ TÍN HIỆU BÁN'})
                sell_group.append(f"🔽 MACD Cross Down (+2)")
                sell_points += 2
            else:
                signals_table.append({'indicator': 'MACD Cross', 'value': '-', 'signal': '⚪ KHÔNG CÓ', 'signal_type': 'none', 'interpretation': 'Chờ crossover'})
    
    # ===== 5. Bollinger Bands =====
    if available.get('BB'):
        bb_upper = ind.get('BB_Upper', close * 1.05)
        bb_lower = ind.get('BB_Lower', close * 0.95)
        bb_middle = ind.get('BB_Middle', close)
        bb_range = bb_upper - bb_lower
        bb_pos = (close - bb_lower) / bb_range * 100 if bb_range > 0 else 50
        
        if close < bb_lower:
            signals_table.append({'indicator': 'BB Position', 'value': f'DƯỚI Lower', 'signal': '📉 DƯỚI BAND', 'signal_type': 'oversold', 'interpretation': 'GIÁ RẺ / CƠ HỘI MUA'})
            buy_group.append(f"Giá dưới BB Lower (+2)")
            buy_points += 2
        elif bb_pos < 20:
            signals_table.append({'indicator': 'BB Position', 'value': f'{bb_pos:.0f}%', 'signal': '📉 GẦN LOWER', 'signal_type': 'low', 'interpretation': 'Gần vùng hỗ trợ'})
            buy_group.append(f"BB vị trí {bb_pos:.0f}% gần Lower (+1)")
            buy_points += 1
        elif close > bb_upper:
            signals_table.append({'indicator': 'BB Position', 'value': f'TRÊN Upper', 'signal': '📈 TRÊN BAND', 'signal_type': 'overbought', 'interpretation': 'GIÁ CAO / CẨN THẬN'})
            sell_group.append(f"Giá trên BB Upper (+2)")
            sell_points += 2
        elif bb_pos > 80:
            signals_table.append({'indicator': 'BB Position', 'value': f'{bb_pos:.0f}%', 'signal': '📈 GẦN UPPER', 'signal_type': 'high', 'interpretation': 'Gần vùng kháng cự'})
            sell_group.append(f"BB vị trí {bb_pos:.0f}% gần Upper (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'BB Position', 'value': f'{bb_pos:.0f}%', 'signal': '⚪ GIỮA BAND', 'signal_type': 'neutral', 'interpretation': 'Vùng trung lập'})
            hold_points += 1
    
    # ===== 6. MFI =====
    if available.get('MFI'):
        mfi = ind['MFI']
        if mfi < 20:
            signals_table.append({'indicator': 'MFI', 'value': f'{mfi:.0f}', 'signal': '📉 DÒNG TIỀN QUÁ BÁN', 'signal_type': 'oversold', 'interpretation': 'Dòng tiền yếu - Cơ hội tích lũy'})
            buy_group.append(f"MFI={mfi:.0f} dòng tiền quá bán (+2)")
            buy_points += 2
        elif mfi < 40:
            signals_table.append({'indicator': 'MFI', 'value': f'{mfi:.0f}', 'signal': '📉 DÒNG TIỀN YẾU', 'signal_type': 'low', 'interpretation': 'Dòng tiền đang yếu'})
            buy_group.append(f"MFI={mfi:.0f} dòng tiền yếu (+1)")
            buy_points += 1
        elif mfi > 80:
            signals_table.append({'indicator': 'MFI', 'value': f'{mfi:.0f}', 'signal': '📈 DÒNG TIỀN QUÁ MUA', 'signal_type': 'overbought', 'interpretation': 'Dòng tiền mạnh - Cẩn thận'})
            sell_group.append(f"MFI={mfi:.0f} dòng tiền quá mua (+2)")
            sell_points += 2
        elif mfi > 60:
            signals_table.append({'indicator': 'MFI', 'value': f'{mfi:.0f}', 'signal': '📈 DÒNG TIỀN MẠNH', 'signal_type': 'high', 'interpretation': 'Dòng tiền đang mạnh'})
            sell_group.append(f"MFI={mfi:.0f} dòng tiền mạnh (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'MFI', 'value': f'{mfi:.0f}', 'signal': '⚪ TRUNG LẬP', 'signal_type': 'neutral', 'interpretation': 'Dòng tiền cân bằng'})
            hold_points += 1
    
    # ===== 7. ADX =====
    if available.get('ADX'):
        adx = ind['ADX']
        plus_di = ind.get('Plus_DI', 25)
        minus_di = ind.get('Minus_DI', 25)
        
        if adx < 20:
            signals_table.append({'indicator': 'ADX', 'value': f'{adx:.0f}', 'signal': '⚪ KHÔNG TREND', 'signal_type': 'no_trend', 'interpretation': 'Thị trường sideway'})
            hold_group.append(f"ADX={adx:.0f} không trend (+1)")
            hold_points += 1
        elif adx < 25:
            signals_table.append({'indicator': 'ADX', 'value': f'{adx:.0f}', 'signal': '⚪ TREND YẾU', 'signal_type': 'weak_trend', 'interpretation': 'Xu hướng yếu'})
            hold_points += 1
        else:
            if plus_di > minus_di:
                signals_table.append({'indicator': 'ADX', 'value': f'{adx:.0f}', 'signal': '📈 UPTREND MẠNH', 'signal_type': 'strong_up', 'interpretation': f'+DI({plus_di:.0f}) > -DI({minus_di:.0f})'})
                buy_group.append(f"ADX={adx:.0f} uptrend +DI>-DI (+1)")
                buy_points += 1
            else:
                signals_table.append({'indicator': 'ADX', 'value': f'{adx:.0f}', 'signal': '📉 DOWNTREND MẠNH', 'signal_type': 'strong_down', 'interpretation': f'-DI({minus_di:.0f}) > +DI({plus_di:.0f})'})
                sell_group.append(f"ADX={adx:.0f} downtrend -DI>+DI (+1)")
                sell_points += 1
        
        # +DI vs -DI
        signals_table.append({'indicator': '+DI vs -DI', 'value': f'{plus_di:.0f} vs {minus_di:.0f}', 'signal': '📈 UPTREND' if plus_di > minus_di else '📉 DOWNTREND', 'signal_type': 'bullish' if plus_di > minus_di else 'bearish', 'interpretation': 'Hướng xu hướng'})
    
    # ===== 8. CCI =====
    if available.get('CCI'):
        cci = ind['CCI']
        if cci < -200:
            signals_table.append({'indicator': 'CCI', 'value': f'{cci:.0f}', 'signal': '📉 CỰC KỲ QUÁ BÁN', 'signal_type': 'extreme_oversold', 'interpretation': 'GIÁ RẤT RẺ / CƠ HỘI LỚN'})
            buy_group.append(f"CCI={cci:.0f} cực kỳ quá bán (+2)")
            buy_points += 2
        elif cci < -100:
            signals_table.append({'indicator': 'CCI', 'value': f'{cci:.0f}', 'signal': '📉 QUÁ BÁN', 'signal_type': 'oversold', 'interpretation': 'GIÁ RẺ / CƠ HỘI MUA'})
            buy_group.append(f"CCI={cci:.0f} quá bán (+1)")
            buy_points += 1
        elif cci > 200:
            signals_table.append({'indicator': 'CCI', 'value': f'{cci:.0f}', 'signal': '📈 CỰC KỲ QUÁ MUA', 'signal_type': 'extreme_overbought', 'interpretation': 'GIÁ RẤT CAO / RỦI RO'})
            sell_group.append(f"CCI={cci:.0f} cực kỳ quá mua (+2)")
            sell_points += 2
        elif cci > 100:
            signals_table.append({'indicator': 'CCI', 'value': f'{cci:.0f}', 'signal': '📈 QUÁ MUA', 'signal_type': 'overbought', 'interpretation': 'GIÁ CAO / CẨN THẬN'})
            sell_group.append(f"CCI={cci:.0f} quá mua (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'CCI', 'value': f'{cci:.0f}', 'signal': '⚪ TRUNG LẬP', 'signal_type': 'neutral', 'interpretation': 'Vùng cân bằng'})
            hold_points += 1
    
    # ===== 9. Williams %R =====
    if available.get('WILLR'):
        willr = ind['Williams_R']
        if willr < -80:
            signals_table.append({'indicator': 'Williams %R', 'value': f'{willr:.0f}', 'signal': '📉 QUÁ BÁN', 'signal_type': 'oversold', 'interpretation': 'GIÁ RẺ / CƠ HỘI MUA'})
            buy_group.append(f"Williams %R={willr:.0f} quá bán (+1)")
            buy_points += 1
        elif willr > -20:
            signals_table.append({'indicator': 'Williams %R', 'value': f'{willr:.0f}', 'signal': '📈 QUÁ MUA', 'signal_type': 'overbought', 'interpretation': 'GIÁ CAO / CẨN THẬN'})
            sell_group.append(f"Williams %R={willr:.0f} quá mua (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'Williams %R', 'value': f'{willr:.0f}', 'signal': '⚪ TRUNG LẬP', 'signal_type': 'neutral', 'interpretation': 'Vùng cân bằng'})
            hold_points += 1
    
    # ===== 10. SAR =====
    if available.get('SAR'):
        sar = ind['SAR']
        if close > sar:
            signals_table.append({'indicator': 'SAR', 'value': f'{sar:,.0f}', 'signal': '📈 UPTREND', 'signal_type': 'bullish', 'interpretation': 'Giá trên SAR - Xu hướng tăng'})
            buy_group.append(f"SAR uptrend (Giá > SAR) (+1)")
            buy_points += 1
        else:
            signals_table.append({'indicator': 'SAR', 'value': f'{sar:,.0f}', 'signal': '📉 DOWNTREND', 'signal_type': 'bearish', 'interpretation': 'Giá dưới SAR - Xu hướng giảm'})
            sell_group.append(f"SAR downtrend (Giá < SAR) (+1)")
            sell_points += 1
        
        # SAR Reversal
        if len(df) > 1 and 'SAR' in df.columns:
            prev_sar = df['SAR'].iloc[-2] if pd.notna(df['SAR'].iloc[-2]) else sar
            prev_close = df['close'].iloc[-2]
            if prev_close < prev_sar and close > sar:
                signals_table.append({'indicator': 'SAR Reversal', 'value': 'Đảo chiều', 'signal': '🔼 ĐẢO CHIỀU LÊN', 'signal_type': 'bullish_reversal', 'interpretation': '🚀 TÍN HIỆU ĐẢO CHIỀU TĂNG'})
                buy_group.append(f"🔄 SAR đảo chiều LÊN (+2)")
                buy_points += 2
            elif prev_close > prev_sar and close < sar:
                signals_table.append({'indicator': 'SAR Reversal', 'value': 'Đảo chiều', 'signal': '🔽 ĐẢO CHIỀU XUỐNG', 'signal_type': 'bearish_reversal', 'interpretation': '⚠️ TÍN HIỆU ĐẢO CHIỀU GIẢM'})
                sell_group.append(f"🔄 SAR đảo chiều XUỐNG (+2)")
                sell_points += 2
    
    # ===== 11. OBV =====
    if available.get('OBV') and len(df) > 5:
        obv_now = ind['OBV']
        obv_5d = df['OBV'].iloc[-6] if pd.notna(df['OBV'].iloc[-6]) else obv_now
        obv_change = (obv_now - obv_5d) / abs(obv_5d) * 100 if obv_5d != 0 else 0
        
        if obv_change > 10:
            signals_table.append({'indicator': 'OBV', 'value': f'+{obv_change:.0f}%', 'signal': '📈 TÍCH LŨY MẠNH', 'signal_type': 'accumulation', 'interpretation': 'Smart money đang MUA'})
            buy_group.append(f"OBV +{obv_change:.0f}% tích lũy mạnh (+1)")
            buy_points += 1
        elif obv_now > obv_5d:
            signals_table.append({'indicator': 'OBV', 'value': f'+{obv_change:.0f}%', 'signal': '📈 TÍCH LŨY', 'signal_type': 'slight_accumulation', 'interpretation': 'Dòng tiền đang vào'})
            buy_group.append(f"OBV tăng (+1)")
            buy_points += 1
        elif obv_change < -10:
            signals_table.append({'indicator': 'OBV', 'value': f'{obv_change:.0f}%', 'signal': '📉 PHÂN PHỐI MẠNH', 'signal_type': 'distribution', 'interpretation': 'Smart money đang BÁN'})
            sell_group.append(f"OBV {obv_change:.0f}% phân phối mạnh (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'OBV', 'value': f'{obv_change:.0f}%', 'signal': '📉 PHÂN PHỐI', 'signal_type': 'slight_distribution', 'interpretation': 'Dòng tiền đang ra'})
            sell_group.append(f"OBV giảm (+1)")
            sell_points += 1
    
    # ===== 12. CMF =====
    if available.get('CMF'):
        cmf = ind['CMF']
        if cmf > 0.1:
            signals_table.append({'indicator': 'CMF', 'value': f'{cmf:.3f}', 'signal': '📈 ÁP LỰC MUA MẠNH', 'signal_type': 'strong_buying', 'interpretation': 'Dòng tiền vào mạnh'})
            buy_group.append(f"CMF={cmf:.2f} áp lực mua mạnh (+1)")
            buy_points += 1
        elif cmf > 0.05:
            signals_table.append({'indicator': 'CMF', 'value': f'{cmf:.3f}', 'signal': '📈 ÁP LỰC MUA', 'signal_type': 'buying', 'interpretation': 'Dòng tiền đang vào'})
            buy_group.append(f"CMF={cmf:.2f} áp lực mua (+1)")
            buy_points += 1
        elif cmf < -0.1:
            signals_table.append({'indicator': 'CMF', 'value': f'{cmf:.3f}', 'signal': '📉 ÁP LỰC BÁN MẠNH', 'signal_type': 'strong_selling', 'interpretation': 'Dòng tiền ra mạnh'})
            sell_group.append(f"CMF={cmf:.2f} áp lực bán mạnh (+1)")
            sell_points += 1
        elif cmf < -0.05:
            signals_table.append({'indicator': 'CMF', 'value': f'{cmf:.3f}', 'signal': '📉 ÁP LỰC BÁN', 'signal_type': 'selling', 'interpretation': 'Dòng tiền đang ra'})
            sell_group.append(f"CMF={cmf:.2f} áp lực bán (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'CMF', 'value': f'{cmf:.3f}', 'signal': '⚪ CÂN BẰNG', 'signal_type': 'neutral', 'interpretation': 'Dòng tiền cân bằng'})
            hold_points += 1
    
    # ===== 13. VWAP =====
    if available.get('VWAP'):
        vwap = ind['VWAP']
        vwap_diff = (close - vwap) / vwap * 100 if vwap > 0 else 0
        
        if close > vwap:
            signals_table.append({'indicator': 'VWAP', 'value': f'+{vwap_diff:.1f}%', 'signal': '📈 GIÁ > VWAP', 'signal_type': 'bullish', 'interpretation': 'Giá trên giá trị hợp lý'})
            buy_group.append(f"Giá > VWAP +{vwap_diff:.1f}% (+1)")
            buy_points += 1
        else:
            signals_table.append({'indicator': 'VWAP', 'value': f'{vwap_diff:.1f}%', 'signal': '📉 GIÁ < VWAP', 'signal_type': 'bearish', 'interpretation': 'Giá dưới giá trị hợp lý'})
            sell_group.append(f"Giá < VWAP {vwap_diff:.1f}% (+1)")
            sell_points += 1
    
    # ===== 14. Force Index =====
    if available.get('FI'):
        fi = ind['FI']
        if fi > 0:
            signals_table.append({'indicator': 'Force Index', 'value': f'{fi:,.0f}', 'signal': '📈 LỰC MUA', 'signal_type': 'buying_force', 'interpretation': 'Lực mua đang chiếm ưu thế'})
            buy_group.append(f"Force Index={fi:,.0f} lực mua (+1)")
            buy_points += 1
        else:
            signals_table.append({'indicator': 'Force Index', 'value': f'{fi:,.0f}', 'signal': '📉 LỰC BÁN', 'signal_type': 'selling_force', 'interpretation': 'Lực bán đang chiếm ưu thế'})
            sell_group.append(f"Force Index={fi:,.0f} lực bán (+1)")
            sell_points += 1
    
    # ===== 15-19. SMA/EMA/WMA/TEMA/DEMA =====
    # SMA
    if available.get('SMA'):
        sma_signals = []
        if available.get('SMA_20') and close > ind['SMA_20']:
            sma_signals.append('SMA20↑')
        elif available.get('SMA_20'):
            sma_signals.append('SMA20↓')
        if available.get('SMA_50') and close > ind['SMA_50']:
            sma_signals.append('SMA50↑')
        elif available.get('SMA_50'):
            sma_signals.append('SMA50↓')
        if available.get('SMA_200') and close > ind['SMA_200']:
            sma_signals.append('SMA200↑')
        elif available.get('SMA_200'):
            sma_signals.append('SMA200↓')
        
        up_count = sum(1 for s in sma_signals if '↑' in s)
        down_count = sum(1 for s in sma_signals if '↓' in s)
        
        if up_count > down_count:
            signals_table.append({'indicator': 'SMA', 'value': ', '.join(sma_signals), 'signal': '📈 BULLISH', 'signal_type': 'bullish', 'interpretation': 'Giá trên đường trung bình'})
            buy_group.append(f"SMA bullish ({up_count}/{len(sma_signals)}) (+1)")
            buy_points += 1
        elif down_count > up_count:
            signals_table.append({'indicator': 'SMA', 'value': ', '.join(sma_signals), 'signal': '📉 BEARISH', 'signal_type': 'bearish', 'interpretation': 'Giá dưới đường trung bình'})
            sell_group.append(f"SMA bearish ({down_count}/{len(sma_signals)}) (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'SMA', 'value': ', '.join(sma_signals), 'signal': '⚪ HỖN HỢP', 'signal_type': 'mixed', 'interpretation': 'Tín hiệu lẫn lộn'})
            hold_points += 1
        
        # Golden/Death Cross
        if available.get('SMA_50') and available.get('SMA_200') and len(df) > 1:
            prev_sma50 = df['SMA_50'].iloc[-2] if pd.notna(df['SMA_50'].iloc[-2]) else ind['SMA_50']
            prev_sma200 = df['SMA_200'].iloc[-2] if pd.notna(df['SMA_200'].iloc[-2]) else ind['SMA_200']
            
            if prev_sma50 < prev_sma200 and ind['SMA_50'] > ind['SMA_200']:
                signals_table.append({'indicator': 'SMA 50/200', 'value': 'Golden Cross', 'signal': '🌟 GOLDEN CROSS', 'signal_type': 'golden_cross', 'interpretation': '🚀 TÍN HIỆU MUA RẤT MẠNH'})
                buy_group.append(f"🌟 Golden Cross (+3)")
                buy_points += 3
            elif prev_sma50 > prev_sma200 and ind['SMA_50'] < ind['SMA_200']:
                signals_table.append({'indicator': 'SMA 50/200', 'value': 'Death Cross', 'signal': '💀 DEATH CROSS', 'signal_type': 'death_cross', 'interpretation': '⚠️ TÍN HIỆU BÁN RẤT MẠNH'})
                sell_group.append(f"💀 Death Cross (+3)")
                sell_points += 3
            elif ind['SMA_50'] > ind['SMA_200']:
                signals_table.append({'indicator': 'SMA 50/200', 'value': f"50>{ind['SMA_200']:,.0f}", 'signal': '📈 BULLISH', 'signal_type': 'bullish', 'interpretation': 'SMA50 trên SMA200'})
            else:
                signals_table.append({'indicator': 'SMA 50/200', 'value': f"50<{ind['SMA_200']:,.0f}", 'signal': '📉 BEARISH', 'signal_type': 'bearish', 'interpretation': 'SMA50 dưới SMA200'})
    
    # EMA
    if available.get('EMA'):
        if available.get('EMA_12') and available.get('EMA_26'):
            if ind['EMA_12'] > ind['EMA_26']:
                signals_table.append({'indicator': 'EMA 12/26', 'value': f"12>{ind['EMA_26']:,.0f}", 'signal': '📈 BULLISH', 'signal_type': 'bullish', 'interpretation': 'EMA ngắn hạn > dài hạn'})
                buy_group.append(f"EMA12 > EMA26 (+1)")
                buy_points += 1
            else:
                signals_table.append({'indicator': 'EMA 12/26', 'value': f"12<{ind['EMA_26']:,.0f}", 'signal': '📉 BEARISH', 'signal_type': 'bearish', 'interpretation': 'EMA ngắn hạn < dài hạn'})
                sell_group.append(f"EMA12 < EMA26 (+1)")
                sell_points += 1
    
    # ===== 20. ROC =====
    if available.get('ROC'):
        roc = ind['ROC']
        if roc > 5:
            signals_table.append({'indicator': 'ROC', 'value': f'{roc:.1f}%', 'signal': '📈 MOMENTUM MẠNH', 'signal_type': 'strong_momentum', 'interpretation': 'Đà tăng mạnh'})
            buy_group.append(f"ROC={roc:.1f}% momentum mạnh (+1)")
            buy_points += 1
        elif roc > 0:
            signals_table.append({'indicator': 'ROC', 'value': f'{roc:.1f}%', 'signal': '📈 TĂNG', 'signal_type': 'positive', 'interpretation': 'Đà tăng'})
            buy_points += 1
        elif roc < -5:
            signals_table.append({'indicator': 'ROC', 'value': f'{roc:.1f}%', 'signal': '📉 MOMENTUM YẾU', 'signal_type': 'weak_momentum', 'interpretation': 'Đà giảm mạnh'})
            sell_group.append(f"ROC={roc:.1f}% momentum yếu (+1)")
            sell_points += 1
        else:
            signals_table.append({'indicator': 'ROC', 'value': f'{roc:.1f}%', 'signal': '📉 GIẢM', 'signal_type': 'negative', 'interpretation': 'Đà giảm'})
            sell_points += 1
    
    # ===== 21. Momentum =====
    if available.get('MOM'):
        mom = ind['MOM']
        if mom > 0:
            signals_table.append({'indicator': 'Momentum', 'value': f'{mom:,.0f}', 'signal': '📈 DƯƠNG', 'signal_type': 'positive', 'interpretation': 'Động lượng tăng'})
            buy_group.append(f"Momentum={mom:,.0f} dương (+1)")
            buy_points += 1
        else:
            signals_table.append({'indicator': 'Momentum', 'value': f'{mom:,.0f}', 'signal': '📉 ÂM', 'signal_type': 'negative', 'interpretation': 'Động lượng giảm'})
            sell_group.append(f"Momentum={mom:,.0f} âm (+1)")
            sell_points += 1
    
    # ===== 22. ATR (Chỉ cảnh báo, không tạo tín hiệu mua/bán) =====
    if available.get('ATR'):
        atr = ind['ATR']
        atr_pct = (atr / close) * 100 if close > 0 else 2
        if atr_pct > 5:
            signals_table.append({'indicator': 'ATR', 'value': f'{atr_pct:.1f}%', 'signal': '⚠️ BIẾN ĐỘNG CAO', 'signal_type': 'high_volatility', 'interpretation': 'Rủi ro cao - Cẩn thận position size'})
        elif atr_pct < 1.5:
            signals_table.append({'indicator': 'ATR', 'value': f'{atr_pct:.1f}%', 'signal': '✅ BIẾN ĐỘNG THẤP', 'signal_type': 'low_volatility', 'interpretation': 'Rủi ro thấp'})
        else:
            signals_table.append({'indicator': 'ATR', 'value': f'{atr_pct:.1f}%', 'signal': '⚪ BIẾN ĐỘNG TB', 'signal_type': 'normal_volatility', 'interpretation': 'Biến động bình thường'})
    
    # ============================================================
    # PHÁT HIỆN XUNG ĐỘT
    # ============================================================
    
    # Xung đột 1: Động lượng vs Xu hướng
    momentum_oversold = any('quá bán' in s.lower() for s in buy_group)
    trend_bearish = any('bearish' in s.lower() or 'downtrend' in s.lower() for s in sell_group)
    if momentum_oversold and trend_bearish:
        conflicts.append({
            'type': 'ĐỘNG LƯỢNG vs XU HƯỚNG',
            'description': 'Động lượng QUÁ BÁN nhưng Xu hướng vẫn BEARISH',
            'interpretation': 'Có thể đang tạo đáy trong xu hướng giảm',
            'suggestion': 'Mua từng phần, chờ xác nhận đảo chiều'
        })
    
    # Xung đột 2: Động lượng vs Xu hướng (ngược lại)
    momentum_overbought = any('quá mua' in s.lower() for s in sell_group)
    trend_bullish = any('bullish' in s.lower() or 'uptrend' in s.lower() for s in buy_group)
    if momentum_overbought and trend_bullish:
        conflicts.append({
            'type': 'ĐỘNG LƯỢNG vs XU HƯỚNG',
            'description': 'Động lượng QUÁ MUA nhưng Xu hướng vẫn BULLISH',
            'interpretation': 'Có thể đang tạo đỉnh trong xu hướng tăng',
            'suggestion': 'Cân nhắc chốt lời từng phần'
        })
    
    # Xung đột 3: Dòng tiền vs Giá
    obv_accumulation = any('tích lũy' in s.lower() for s in buy_group)
    price_down = any('giá < vwap' in s.lower() or 'sma bearish' in s.lower() for s in sell_group)
    if obv_accumulation and price_down:
        conflicts.append({
            'type': 'PHÂN KỲ DƯƠNG',
            'description': 'Giá giảm nhưng dòng tiền đang vào',
            'interpretation': 'Smart money có thể đang tích lũy đáy',
            'suggestion': 'Theo dõi sát, chuẩn bị mua khi có xác nhận'
        })
    
    obv_distribution = any('phân phối' in s.lower() for s in sell_group)
    price_up = any('giá > vwap' in s.lower() or 'sma bullish' in s.lower() for s in buy_group)
    if obv_distribution and price_up:
        conflicts.append({
            'type': 'PHÂN KỲ ÂM',
            'description': 'Giá tăng nhưng dòng tiền đang ra',
            'interpretation': 'Smart money có thể đang phân phối đỉnh',
            'suggestion': 'Cẩn thận, cân nhắc chốt lời'
        })
    
    return {
        'signals_table': signals_table,
        'buy_group': buy_group,
        'sell_group': sell_group,
        'hold_group': hold_group,
        'buy_points': buy_points,
        'sell_points': sell_points,
        'hold_points': hold_points,
        'conflicts': conflicts
    }


def generate_action_analysis(phase, buy_points, sell_points, hold_points, conflicts, ind, close):
    """
    Tạo phân tích hành động chi tiết dựa trên tín hiệu
    """
    
    net_signal = buy_points - sell_points
    total_signals = buy_points + sell_points + hold_points
    
    analysis = {
        'summary': '',
        'factors': [],
        'warnings': [],
        'action_chua_co': '',
        'action_chua_co_display': '',
        'action_dang_giu': '',
        'action_dang_giu_display': '',
        'action_reason_chua_co': '',
        'action_reason_dang_giu': '',
        'strategy_chua_co': '',
        'strategy_dang_giu': ''
    }
    
    # ============================================================
    # XÁC ĐỊNH HÀNH ĐỘNG CHO NGƯỜI CHƯA CÓ
    # ============================================================
    
    if buy_points >= sell_points + 5:
        analysis['action_chua_co'] = 'MUA_MANH'
        analysis['action_chua_co_display'] = '🟢 MUA MẠNH'
        analysis['action_reason_chua_co'] = 'Đa số chỉ báo đồng thuận tích cực, nên vào lệnh'
        analysis['strategy_chua_co'] = 'Mua 70-100% vị thế dự kiến'
    elif buy_points >= sell_points + 3:
        analysis['action_chua_co'] = 'MUA'
        analysis['action_chua_co_display'] = '🟢 MUA'
        analysis['action_reason_chua_co'] = 'Tín hiệu tích cực, có thể mua'
        analysis['strategy_chua_co'] = 'Mua 50-70% vị thế dự kiến'
    elif buy_points > sell_points:
        analysis['action_chua_co'] = 'CAN_NHAC_MUA'
        analysis['action_chua_co_display'] = '🟢 CÂN NHẮC MUA'
        analysis['action_reason_chua_co'] = 'Có tín hiệu tích cực nhưng chưa đủ mạnh'
        analysis['strategy_chua_co'] = 'Mua 30% vị thế, chờ xác nhận để mua thêm'
    elif sell_points > buy_points + 3:
        analysis['action_chua_co'] = 'TRANH'
        analysis['action_chua_co_display'] = '🔴 TRÁNH'
        analysis['action_reason_chua_co'] = 'Tín hiệu tiêu cực, không nên vào'
        analysis['strategy_chua_co'] = 'Không mua, chờ tín hiệu tốt hơn'
    elif sell_points > buy_points:
        analysis['action_chua_co'] = 'CHO'
        analysis['action_chua_co_display'] = '🟡 CHỜ TÍN HIỆU'
        analysis['action_reason_chua_co'] = 'Tín hiệu tiêu cực nhẹ'
        analysis['strategy_chua_co'] = 'Chờ đợi, theo dõi các chỉ báo'
    else:
        analysis['action_chua_co'] = 'THEO_DOI'
        analysis['action_chua_co_display'] = '⚪ THEO DÕI'
        analysis['action_reason_chua_co'] = 'Xung đột tín hiệu, chờ rõ hơn'
        analysis['strategy_chua_co'] = 'Quan sát, chưa hành động'
    
    # ============================================================
    # XÁC ĐỊNH HÀNH ĐỘNG CHO NGƯỜI ĐANG GIỮ
    # ============================================================
    
    if sell_points >= buy_points + 5:
        analysis['action_dang_giu'] = 'BAN_MANH'
        analysis['action_dang_giu_display'] = '🔴 BÁN MẠNH'
        analysis['action_reason_dang_giu'] = 'Nhiều tín hiệu tiêu cực, nên thoát'
        analysis['strategy_dang_giu'] = 'Bán 70-100% vị thế'
    elif sell_points >= buy_points + 3:
        analysis['action_dang_giu'] = 'BAN'
        analysis['action_dang_giu_display'] = '🔴 BÁN / CHỐT LỜI'
        analysis['action_reason_dang_giu'] = 'Tín hiệu tiêu cực, nên chốt lời'
        analysis['strategy_dang_giu'] = 'Bán 50-70% vị thế, giữ lại 30%'
    elif sell_points > buy_points:
        analysis['action_dang_giu'] = 'CAN_NHAC_BAN'
        analysis['action_dang_giu_display'] = '🟡 CÂN NHẮC BÁN'
        analysis['action_reason_dang_giu'] = 'Có tín hiệu tiêu cực nhẹ'
        analysis['strategy_dang_giu'] = 'Bán 30% vị thế, đặt trailing stop'
    elif buy_points >= sell_points + 4:
        analysis['action_dang_giu'] = 'MUA_THEM'
        analysis['action_dang_giu_display'] = '🟢 MUA THÊM'
        analysis['action_reason_dang_giu'] = 'Tín hiệu rất tích cực'
        analysis['strategy_dang_giu'] = 'Mua thêm 30-50% vị thế hiện tại'
    elif buy_points >= sell_points + 2:
        analysis['action_dang_giu'] = 'GIU_MUA_THEM'
        analysis['action_dang_giu_display'] = '🟢 GIỮ + MUA THÊM'
        analysis['action_reason_dang_giu'] = 'Tín hiệu tích cực, đang ở vùng hỗ trợ'
        analysis['strategy_dang_giu'] = 'Giữ nguyên, mua thêm 20% nếu có điều kiện'
    else:
        analysis['action_dang_giu'] = 'GIU'
        analysis['action_dang_giu_display'] = '⚪ GIỮ / THEO DÕI'
        analysis['action_reason_dang_giu'] = 'Tín hiệu hỗn hợp, giữ nguyên vị thế'
        analysis['strategy_dang_giu'] = 'Giữ nguyên, đặt stop loss bảo vệ'
    
    # ============================================================
    # YẾU TỐ CHÍNH
    # ============================================================
    
    rsi = ind.get('RSI', 50)
    stoch_k = ind.get('Stoch_K', 50)
    mfi = ind.get('MFI', 50)
    macd_hist = ind.get('MACD_Hist', 0)
    
    if rsi < 30:
        analysis['factors'].append(f"📉 RSI={rsi:.0f} < 30: CỔ PHIẾU QUÁ BÁN - Cơ hội mua tích lũy")
    elif rsi > 70:
        analysis['factors'].append(f"📈 RSI={rsi:.0f} > 70: CỔ PHIẾU QUÁ MUA - Cẩn thận chốt lời")
    
    if stoch_k < 20:
        analysis['factors'].append(f"📉 Stochastic={stoch_k:.0f} < 20: Vùng QUÁ BÁN - Tín hiệu tích lũy")
    elif stoch_k > 80:
        analysis['factors'].append(f"📈 Stochastic={stoch_k:.0f} > 80: Vùng QUÁ MUA - Cẩn thận")
    
    if mfi < 20:
        analysis['factors'].append(f"💰 MFI={mfi:.0f} < 20: Dòng tiền QUÁ BÁN - Cơ hội tích lũy")
    elif mfi > 80:
        analysis['factors'].append(f"💰 MFI={mfi:.0f} > 80: Dòng tiền QUÁ MUA - Cẩn thận")
    
    if macd_hist > 0:
        analysis['factors'].append("📊 MACD Histogram > 0: Động lượng TĂNG")
    else:
        analysis['factors'].append("📊 MACD Histogram < 0: Động lượng GIẢM")
    
    # ============================================================
    # PHÂN TÍCH HÀNH ĐỘNG
    # ============================================================
    
    action_analysis = []
    
    # Phân tích dựa trên pha
    if phase == 'ĐÁY':
        if rsi < 40 and stoch_k < 40:
            action_analysis.append("✅ RSI/Stochastic quá bán + Pha đáy → Cơ hội TÍCH LŨY tốt")
        if mfi < 40:
            action_analysis.append("✅ Dòng tiền đang tích lũy → Xác nhận đáy có smart money mua")
        action_analysis.append("⚠️ Xu hướng vẫn giảm → Nên chia nhỏ lệnh, không all-in")
    elif phase == 'TĂNG':
        if rsi > 50 and stoch_k > 50:
            action_analysis.append("✅ RSI/Stochastic vùng tăng + Pha tăng → Theo trend")
        if macd_hist > 0:
            action_analysis.append("✅ MACD dương → Động lượng đang mạnh")
        action_analysis.append("⚠️ Cẩn thận vùng kháng cự → Đặt trailing stop")
    elif phase == 'ĐỈNH':
        if rsi > 60:
            action_analysis.append("⚠️ RSI vùng cao + Pha đỉnh → Cẩn thận chốt lời")
        if stoch_k > 80:
            action_analysis.append("⚠️ Stochastic quá mua → Có thể điều chỉnh")
        action_analysis.append("✅ Nên bán dần, không bán hết cùng lúc")
    elif phase == 'GIẢM':
        action_analysis.append("⚠️ Pha giảm → Hạn chế mua mới")
        if rsi < 30:
            action_analysis.append("✅ RSI quá bán → Có thể tích lũy nhỏ cho dài hạn")
        action_analysis.append("⚠️ Chờ tín hiệu đảo chiều rõ ràng")
    else:  # TÍCH_LŨY
        action_analysis.append("⚪ Pha tích lũy → Chờ breakout")
        if buy_points > sell_points:
            action_analysis.append("✅ Tín hiệu thiên về TĂNG → Chuẩn bị mua khi breakout")
        elif sell_points > buy_points:
            action_analysis.append("⚠️ Tín hiệu thiên về GIẢM → Cẩn thận breakdown")
    
    # Kết luận
    if buy_points > sell_points:
        action_analysis.append("→ Nhiều chỉ báo thiên về TĂNG")
    elif sell_points > buy_points:
        action_analysis.append("→ Nhiều chỉ báo thiên về GIẢM")
    else:
        action_analysis.append("→ Tín hiệu XUNG ĐỘT, cần chờ rõ hơn")
    
    analysis['action_analysis'] = action_analysis
    
    # ============================================================
    # CẢNH BÁO
    # ============================================================
    
    for conflict in conflicts:
        analysis['warnings'].append(f"⚠️ {conflict['type']}: {conflict['description']} - {conflict['interpretation']}")
    
    # ============================================================
    # TÓM TẮT
    # ============================================================
    
    analysis['summary'] = f"Mua({buy_points}) - Bán({sell_points}) = {'+' if net_signal > 0 else ''}{net_signal}"
    
    return analysis

# ============================================================
# HÀM PHÂN TÍCH NÂNG CAO - KHAI THÁC TỐI ĐA DỮ LIỆU
# ============================================================

def advanced_indicator_analysis(df, ind, available, close):
    """
    Phân tích nâng cao: Phân kỳ, Xu hướng chỉ báo, Crossover timing, BB Squeeze
    """
    
    advanced_signals = {
        'divergences': [],      # Phân kỳ
        'indicator_trends': [], # Xu hướng chỉ báo
        'crossover_timing': [], # Thời điểm crossover
        'bb_analysis': {},      # Phân tích BB nâng cao
        'pattern_signals': [],  # Nhận diện mô hình
        'confluence_zones': [], # Vùng hội tụ
        'strength_score': 0,    # Điểm mạnh tổng hợp
    }
    
    n = len(df)
    if n < 10:
        return advanced_signals
    
    # ============================================================
    # 1. PHÂN TÍCH PHÂN KỲ (DIVERGENCE)
    # ============================================================
    
    # Tìm đỉnh/đáy giá trong 20 ngày
    price_data = df['close'].tail(20).values
    
    # RSI Divergence
    if available.get('RSI') and 'RSI' in df.columns:
        rsi_data = df['RSI'].tail(20).values
        
        # Bullish Divergence: Giá tạo đáy thấp hơn, RSI tạo đáy cao hơn
        if n >= 10:
            price_recent_low = df['close'].tail(5).min()
            price_prev_low = df['close'].tail(10).head(5).min()
            rsi_recent_low = df['RSI'].tail(5).min()
            rsi_prev_low = df['RSI'].tail(10).head(5).min()
            
            if price_recent_low < price_prev_low and rsi_recent_low > rsi_prev_low:
                advanced_signals['divergences'].append({
                    'type': 'BULLISH_DIVERGENCE',
                    'indicator': 'RSI',
                    'description': 'Giá tạo đáy thấp hơn nhưng RSI tạo đáy cao hơn',
                    'interpretation': '🟢 Lực bán đang yếu đi, có thể đảo chiều TĂNG',
                    'strength': 'MẠNH' if ind['RSI'] < 35 else 'TRUNG BÌNH',
                    'action': 'CÂN NHẮC MUA',
                    'score': 3 if ind['RSI'] < 35 else 2
                })
            
            # Bearish Divergence: Giá tạo đỉnh cao hơn, RSI tạo đỉnh thấp hơn
            price_recent_high = df['close'].tail(5).max()
            price_prev_high = df['close'].tail(10).head(5).max()
            rsi_recent_high = df['RSI'].tail(5).max()
            rsi_prev_high = df['RSI'].tail(10).head(5).max()
            
            if price_recent_high > price_prev_high and rsi_recent_high < rsi_prev_high:
                advanced_signals['divergences'].append({
                    'type': 'BEARISH_DIVERGENCE',
                    'indicator': 'RSI',
                    'description': 'Giá tạo đỉnh cao hơn nhưng RSI tạo đỉnh thấp hơn',
                    'interpretation': '🔴 Lực mua đang yếu đi, có thể đảo chiều GIẢM',
                    'strength': 'MẠNH' if ind['RSI'] > 65 else 'TRUNG BÌNH',
                    'action': 'CÂN NHẮC BÁN',
                    'score': -3 if ind['RSI'] > 65 else -2
                })
    
    # MACD Divergence
    if available.get('MACD') and 'MACD_Hist' in df.columns:
        if n >= 10:
            price_recent_low = df['close'].tail(5).min()
            price_prev_low = df['close'].tail(10).head(5).min()
            macd_recent_low = df['MACD_Hist'].tail(5).min()
            macd_prev_low = df['MACD_Hist'].tail(10).head(5).min()
            
            if price_recent_low < price_prev_low and macd_recent_low > macd_prev_low:
                advanced_signals['divergences'].append({
                    'type': 'BULLISH_DIVERGENCE',
                    'indicator': 'MACD',
                    'description': 'Giá tạo đáy thấp hơn nhưng MACD Histogram tạo đáy cao hơn',
                    'interpretation': '🟢 Động lượng giảm đang yếu đi, có thể đảo chiều',
                    'strength': 'MẠNH',
                    'action': 'CÂN NHẮC MUA',
                    'score': 3
                })
            
            price_recent_high = df['close'].tail(5).max()
            price_prev_high = df['close'].tail(10).head(5).max()
            macd_recent_high = df['MACD_Hist'].tail(5).max()
            macd_prev_high = df['MACD_Hist'].tail(10).head(5).max()
            
            if price_recent_high > price_prev_high and macd_recent_high < macd_prev_high:
                advanced_signals['divergences'].append({
                    'type': 'BEARISH_DIVERGENCE',
                    'indicator': 'MACD',
                    'description': 'Giá tạo đỉnh cao hơn nhưng MACD Histogram tạo đỉnh thấp hơn',
                    'interpretation': '🔴 Động lượng tăng đang yếu đi, có thể đảo chiều',
                    'strength': 'MẠNH',
                    'action': 'CÂN NHẮC BÁN',
                    'score': -3
                })
    
    # OBV Divergence
    if available.get('OBV') and 'OBV' in df.columns:
        if n >= 10:
            price_change_5d = (df['close'].iloc[-1] - df['close'].iloc[-6]) / df['close'].iloc[-6] * 100
            obv_change_5d = (df['OBV'].iloc[-1] - df['OBV'].iloc[-6]) / abs(df['OBV'].iloc[-6]) * 100 if df['OBV'].iloc[-6] != 0 else 0
            
            # Giá tăng nhưng OBV giảm → Bearish
            if price_change_5d > 2 and obv_change_5d < -5:
                advanced_signals['divergences'].append({
                    'type': 'BEARISH_DIVERGENCE',
                    'indicator': 'OBV',
                    'description': f'Giá tăng {price_change_5d:.1f}% nhưng OBV giảm {obv_change_5d:.1f}%',
                    'interpretation': '🔴 Phân phối ẩn - Smart money đang bán',
                    'strength': 'MẠNH',
                    'action': 'CẨN THẬN / CHỐT LỜI',
                    'score': -3
                })
            
            # Giá giảm nhưng OBV tăng → Bullish
            if price_change_5d < -2 and obv_change_5d > 5:
                advanced_signals['divergences'].append({
                    'type': 'BULLISH_DIVERGENCE',
                    'indicator': 'OBV',
                    'description': f'Giá giảm {price_change_5d:.1f}% nhưng OBV tăng {obv_change_5d:.1f}%',
                    'interpretation': '🟢 Tích lũy ẩn - Smart money đang mua',
                    'strength': 'MẠNH',
                    'action': 'CÂN NHẮC MUA',
                    'score': 3
                })
    
    # MFI Divergence
    if available.get('MFI') and 'MFI' in df.columns:
        if n >= 10:
            price_recent_low = df['close'].tail(5).min()
            price_prev_low = df['close'].tail(10).head(5).min()
            mfi_recent_low = df['MFI'].tail(5).min()
            mfi_prev_low = df['MFI'].tail(10).head(5).min()
            
            if price_recent_low < price_prev_low and mfi_recent_low > mfi_prev_low:
                advanced_signals['divergences'].append({
                    'type': 'BULLISH_DIVERGENCE',
                    'indicator': 'MFI',
                    'description': 'Giá tạo đáy thấp hơn nhưng MFI tạo đáy cao hơn',
                    'interpretation': '🟢 Dòng tiền đang vào dù giá giảm',
                    'strength': 'TRUNG BÌNH',
                    'action': 'THEO DÕI MUA',
                    'score': 2
                })
    
    # ============================================================
    # 2. XU HƯỚNG CHỈ BÁO (INDICATOR TREND)
    # ============================================================
    
    # RSI Trend
    if available.get('RSI') and 'RSI' in df.columns and n >= 5:
        rsi_now = ind['RSI']
        rsi_3d = df['RSI'].iloc[-4] if pd.notna(df['RSI'].iloc[-4]) else rsi_now
        rsi_5d = df['RSI'].iloc[-6] if pd.notna(df['RSI'].iloc[-6]) else rsi_now
        
        rsi_trend_3d = rsi_now - rsi_3d
        rsi_trend_5d = rsi_now - rsi_5d
        
        if rsi_trend_3d > 5 and rsi_trend_5d > 8:
            advanced_signals['indicator_trends'].append({
                'indicator': 'RSI',
                'trend': 'TĂNG MẠNH',
                'value': f'{rsi_now:.0f} (↑{rsi_trend_5d:.0f} trong 5 ngày)',
                'interpretation': '📈 Động lượng đang tăng mạnh',
                'score': 2
            })
        elif rsi_trend_3d > 2:
            advanced_signals['indicator_trends'].append({
                'indicator': 'RSI',
                'trend': 'TĂNG',
                'value': f'{rsi_now:.0f} (↑{rsi_trend_3d:.0f} trong 3 ngày)',
                'interpretation': '📈 Động lượng đang cải thiện',
                'score': 1
            })
        elif rsi_trend_3d < -5 and rsi_trend_5d < -8:
            advanced_signals['indicator_trends'].append({
                'indicator': 'RSI',
                'trend': 'GIẢM MẠNH',
                'value': f'{rsi_now:.0f} (↓{abs(rsi_trend_5d):.0f} trong 5 ngày)',
                'interpretation': '📉 Động lượng đang suy yếu mạnh',
                'score': -2
            })
        elif rsi_trend_3d < -2:
            advanced_signals['indicator_trends'].append({
                'indicator': 'RSI',
                'trend': 'GIẢM',
                'value': f'{rsi_now:.0f} (↓{abs(rsi_trend_3d):.0f} trong 3 ngày)',
                'interpretation': '📉 Động lượng đang suy yếu',
                'score': -1
            })
    
    # MACD Histogram Trend (Momentum của Momentum)
    if available.get('MACD') and 'MACD_Hist' in df.columns and n >= 5:
        hist_now = ind['MACD_Hist']
        hist_3d = df['MACD_Hist'].iloc[-4] if pd.notna(df['MACD_Hist'].iloc[-4]) else hist_now
        hist_5d = df['MACD_Hist'].iloc[-6] if pd.notna(df['MACD_Hist'].iloc[-6]) else hist_now
        
        # Histogram đang tăng dần (dù âm hay dương)
        if hist_now > hist_3d > hist_5d:
            if hist_now > 0:
                advanced_signals['indicator_trends'].append({
                    'indicator': 'MACD Histogram',
                    'trend': 'TĂNG DẦN (DƯƠNG)',
                    'value': f'{hist_now:.3f}',
                    'interpretation': '📈 Động lượng tăng đang mạnh lên',
                    'score': 2
                })
            else:
                advanced_signals['indicator_trends'].append({
                    'indicator': 'MACD Histogram',
                    'trend': 'BỚT ÂM DẦN',
                    'value': f'{hist_now:.3f}',
                    'interpretation': '📈 Động lượng giảm đang yếu đi → Sắp đảo chiều?',
                    'score': 1
                })
        elif hist_now < hist_3d < hist_5d:
            if hist_now < 0:
                advanced_signals['indicator_trends'].append({
                    'indicator': 'MACD Histogram',
                    'trend': 'GIẢM DẦN (ÂM)',
                    'value': f'{hist_now:.3f}',
                    'interpretation': '📉 Động lượng giảm đang mạnh lên',
                    'score': -2
                })
            else:
                advanced_signals['indicator_trends'].append({
                    'indicator': 'MACD Histogram',
                    'trend': 'BỚT DƯƠNG DẦN',
                    'value': f'{hist_now:.3f}',
                    'interpretation': '📉 Động lượng tăng đang yếu đi → Sắp đảo chiều?',
                    'score': -1
                })
    
    # ADX Trend
    if available.get('ADX') and 'ADX' in df.columns and n >= 5:
        adx_now = ind['ADX']
        adx_5d = df['ADX'].iloc[-6] if pd.notna(df['ADX'].iloc[-6]) else adx_now
        adx_change = adx_now - adx_5d
        
        if adx_change > 5:
            advanced_signals['indicator_trends'].append({
                'indicator': 'ADX',
                'trend': 'TĂNG',
                'value': f'{adx_now:.0f} (↑{adx_change:.0f})',
                'interpretation': '📈 Xu hướng đang MẠNH LÊN',
                'score': 1 if ind.get('Plus_DI', 0) > ind.get('Minus_DI', 0) else -1
            })
        elif adx_change < -5:
            advanced_signals['indicator_trends'].append({
                'indicator': 'ADX',
                'trend': 'GIẢM',
                'value': f'{adx_now:.0f} (↓{abs(adx_change):.0f})',
                'interpretation': '📉 Xu hướng đang YẾU ĐI → Có thể sideway',
                'score': 0
            })
    
    # OBV Trend
    if available.get('OBV') and 'OBV' in df.columns and n >= 10:
        obv_now = ind['OBV']
        obv_5d = df['OBV'].iloc[-6] if pd.notna(df['OBV'].iloc[-6]) else obv_now
        obv_10d = df['OBV'].iloc[-11] if pd.notna(df['OBV'].iloc[-11]) else obv_now
        
        obv_trend_5d = (obv_now - obv_5d) / abs(obv_5d) * 100 if obv_5d != 0 else 0
        obv_trend_10d = (obv_now - obv_10d) / abs(obv_10d) * 100 if obv_10d != 0 else 0
        
        if obv_trend_5d > 10 and obv_trend_10d > 15:
            advanced_signals['indicator_trends'].append({
                'indicator': 'OBV',
                'trend': 'TÍCH LŨY MẠNH',
                'value': f'+{obv_trend_10d:.0f}% (10 ngày)',
                'interpretation': '📈 Dòng tiền đang VÀO MẠNH',
                'score': 2
            })
        elif obv_trend_5d < -10 and obv_trend_10d < -15:
            advanced_signals['indicator_trends'].append({
                'indicator': 'OBV',
                'trend': 'PHÂN PHỐI MẠNH',
                'value': f'{obv_trend_10d:.0f}% (10 ngày)',
                'interpretation': '📉 Dòng tiền đang RA MẠNH',
                'score': -2
            })
    
    # ============================================================
    # 3. CROSSOVER TIMING
    # ============================================================
    
    # MACD Cross Timing
    if available.get('MACD') and 'MACD_Hist' in df.columns and n >= 5:
        # Tìm ngày cross gần nhất
        cross_day = None
        for i in range(1, min(10, n)):
            prev_hist = df['MACD_Hist'].iloc[-(i+1)]
            curr_hist = df['MACD_Hist'].iloc[-i]
            if pd.notna(prev_hist) and pd.notna(curr_hist):
                if prev_hist < 0 and curr_hist > 0:
                    cross_day = i
                    cross_type = 'BULLISH'
                    break
                elif prev_hist > 0 and curr_hist < 0:
                    cross_day = i
                    cross_type = 'BEARISH'
                    break
        
        if cross_day:
            if cross_day <= 2:
                freshness = 'RẤT MỚI'
                score_mult = 1.5
            elif cross_day <= 5:
                freshness = 'MỚI'
                score_mult = 1.0
            else:
                freshness = 'CŨ'
                score_mult = 0.5
            
            advanced_signals['crossover_timing'].append({
                'indicator': 'MACD',
                'type': cross_type,
                'days_ago': cross_day,
                'freshness': freshness,
                'interpretation': f"{'🔼' if cross_type == 'BULLISH' else '🔽'} MACD Cross {cross_type} cách đây {cross_day} ngày ({freshness})",
                'score': int((2 if cross_type == 'BULLISH' else -2) * score_mult)
            })
    
    # Stochastic Cross Timing
    if available.get('STOCH') and 'Stoch_K' in df.columns and 'Stoch_D' in df.columns and n >= 5:
        cross_day = None
        for i in range(1, min(10, n)):
            prev_k = df['Stoch_K'].iloc[-(i+1)]
            prev_d = df['Stoch_D'].iloc[-(i+1)]
            curr_k = df['Stoch_K'].iloc[-i]
            curr_d = df['Stoch_D'].iloc[-i]
            if all(pd.notna([prev_k, prev_d, curr_k, curr_d])):
                if prev_k < prev_d and curr_k > curr_d:
                    cross_day = i
                    cross_type = 'BULLISH'
                    cross_zone = 'OVERSOLD' if curr_k < 30 else ('OVERBOUGHT' if curr_k > 70 else 'NEUTRAL')
                    break
                elif prev_k > prev_d and curr_k < curr_d:
                    cross_day = i
                    cross_type = 'BEARISH'
                    cross_zone = 'OVERBOUGHT' if curr_k > 70 else ('OVERSOLD' if curr_k < 30 else 'NEUTRAL')
                    break
        
        if cross_day:
            if cross_day <= 2:
                freshness = 'RẤT MỚI'
            elif cross_day <= 5:
                freshness = 'MỚI'
            else:
                freshness = 'CŨ'
            
            # Cross ở vùng quá bán/mua có ý nghĩa hơn
            if cross_type == 'BULLISH' and cross_zone == 'OVERSOLD':
                score = 3
                note = '(Vùng quá bán - Tín hiệu MẠNH)'
            elif cross_type == 'BEARISH' and cross_zone == 'OVERBOUGHT':
                score = -3
                note = '(Vùng quá mua - Tín hiệu MẠNH)'
            else:
                score = 1 if cross_type == 'BULLISH' else -1
                note = ''
            
            advanced_signals['crossover_timing'].append({
                'indicator': 'Stochastic',
                'type': cross_type,
                'days_ago': cross_day,
                'freshness': freshness,
                'zone': cross_zone,
                'interpretation': f"{'🔼' if cross_type == 'BULLISH' else '🔽'} Stoch K/D Cross {cross_type} cách đây {cross_day} ngày {note}",
                'score': score
            })
    
    # DI+/DI- Cross
    if available.get('ADX') and 'Plus_DI' in df.columns and 'Minus_DI' in df.columns and n >= 5:
        cross_day = None
        for i in range(1, min(10, n)):
            prev_plus = df['Plus_DI'].iloc[-(i+1)]
            prev_minus = df['Minus_DI'].iloc[-(i+1)]
            curr_plus = df['Plus_DI'].iloc[-i]
            curr_minus = df['Minus_DI'].iloc[-i]
            if all(pd.notna([prev_plus, prev_minus, curr_plus, curr_minus])):
                if prev_plus < prev_minus and curr_plus > curr_minus:
                    cross_day = i
                    cross_type = 'BULLISH'
                    break
                elif prev_plus > prev_minus and curr_plus < curr_minus:
                    cross_day = i
                    cross_type = 'BEARISH'
                    break
        
        if cross_day:
            adx_value = ind['ADX']
            if adx_value > 25:
                strength = 'MẠNH' if adx_value > 40 else 'TRUNG BÌNH'
                score = (3 if cross_type == 'BULLISH' else -3) if adx_value > 40 else (2 if cross_type == 'BULLISH' else -2)
            else:
                strength = 'YẾU (ADX thấp)'
                score = 1 if cross_type == 'BULLISH' else -1
            
            advanced_signals['crossover_timing'].append({
                'indicator': 'DI+/DI-',
                'type': cross_type,
                'days_ago': cross_day,
                'adx': adx_value,
                'strength': strength,
                'interpretation': f"{'🔼' if cross_type == 'BULLISH' else '🔽'} DI Cross {cross_type} (ADX={adx_value:.0f}) - {strength}",
                'score': score
            })
    
    # ============================================================
    # 4. BOLLINGER BANDS ANALYSIS
    # ============================================================
    
    if available.get('BB') and 'BB_Upper' in df.columns and 'BB_Lower' in df.columns and n >= 20:
        bb_upper = ind['BB_Upper']
        bb_lower = ind['BB_Lower']
        bb_middle = ind.get('BB_Middle', (bb_upper + bb_lower) / 2)
        
        # Bandwidth (độ rộng band)
        bandwidth = (bb_upper - bb_lower) / bb_middle * 100 if bb_middle > 0 else 0
        
        # So sánh với bandwidth trung bình 20 ngày
        if 'BB_Upper' in df.columns and 'BB_Lower' in df.columns:
            bb_widths = []
            for i in range(min(20, n)):
                u = df['BB_Upper'].iloc[-(i+1)]
                l = df['BB_Lower'].iloc[-(i+1)]
                m = (u + l) / 2
                if pd.notna(u) and pd.notna(l) and m > 0:
                    bb_widths.append((u - l) / m * 100)
            
            if bb_widths:
                avg_bandwidth = np.mean(bb_widths)
                min_bandwidth = min(bb_widths)
                
                advanced_signals['bb_analysis'] = {
                    'bandwidth': bandwidth,
                    'avg_bandwidth': avg_bandwidth,
                    'bandwidth_percentile': sum(1 for w in bb_widths if w < bandwidth) / len(bb_widths) * 100
                }
                
                # BB Squeeze Detection
                if bandwidth < avg_bandwidth * 0.7:
                    advanced_signals['bb_analysis']['squeeze'] = True
                    advanced_signals['bb_analysis']['squeeze_interpretation'] = '🔥 BB SQUEEZE - Chuẩn bị BREAKOUT mạnh!'
                    # Hướng breakout dự đoán dựa trên các chỉ báo khác
                    if ind.get('RSI', 50) > 50 and ind.get('MACD_Hist', 0) > 0:
                        advanced_signals['bb_analysis']['expected_direction'] = 'UP'
                        advanced_signals['bb_analysis']['score'] = 2
                    elif ind.get('RSI', 50) < 50 and ind.get('MACD_Hist', 0) < 0:
                        advanced_signals['bb_analysis']['expected_direction'] = 'DOWN'
                        advanced_signals['bb_analysis']['score'] = -2
                    else:
                        advanced_signals['bb_analysis']['expected_direction'] = 'UNCERTAIN'
                        advanced_signals['bb_analysis']['score'] = 0
                else:
                    advanced_signals['bb_analysis']['squeeze'] = False
                
                # %B (vị trí trong band)
                percent_b = (close - bb_lower) / (bb_upper - bb_lower) * 100 if (bb_upper - bb_lower) > 0 else 50
                advanced_signals['bb_analysis']['percent_b'] = percent_b
                
                # Band Walking (giá đi dọc theo band)
                if percent_b > 95:
                    advanced_signals['bb_analysis']['band_walk'] = 'UPPER'
                    advanced_signals['bb_analysis']['band_walk_interpretation'] = '📈 Đang "đi bộ" trên BB Upper - Uptrend mạnh'
                elif percent_b < 5:
                    advanced_signals['bb_analysis']['band_walk'] = 'LOWER'
                    advanced_signals['bb_analysis']['band_walk_interpretation'] = '📉 Đang "đi bộ" dưới BB Lower - Downtrend mạnh'
    
    # ============================================================
    # 5. PATTERN RECOGNITION (Nhận diện mô hình)
    # ============================================================
    
    if n >= 10:
        closes = df['close'].tail(10).values
        highs = df['high'].tail(10).values
        lows = df['low'].tail(10).values
        
        # Higher Highs, Higher Lows (Uptrend)
        recent_highs = [highs[-1], highs[-3], highs[-5]]
        recent_lows = [lows[-1], lows[-3], lows[-5]]
        
        if recent_highs[0] > recent_highs[1] > recent_highs[2] and recent_lows[0] > recent_lows[1] > recent_lows[2]:
            advanced_signals['pattern_signals'].append({
                'pattern': 'HIGHER_HIGHS_HIGHER_LOWS',
                'interpretation': '📈 Đỉnh cao hơn + Đáy cao hơn → UPTREND rõ ràng',
                'action': 'THEO TREND TĂNG',
                'score': 2
            })
        
        # Lower Highs, Lower Lows (Downtrend)
        elif recent_highs[0] < recent_highs[1] < recent_highs[2] and recent_lows[0] < recent_lows[1] < recent_lows[2]:
            advanced_signals['pattern_signals'].append({
                'pattern': 'LOWER_HIGHS_LOWER_LOWS',
                'interpretation': '📉 Đỉnh thấp hơn + Đáy thấp hơn → DOWNTREND rõ ràng',
                'action': 'TRÁNH MUA / CHỜ ĐÁY',
                'score': -2
            })
        
        # Double Bottom Detection (đơn giản)
        if n >= 15:
            lows_15d = df['low'].tail(15).values
            min1_idx = np.argmin(lows_15d[:7])  # Đáy 1 trong 7 ngày đầu
            min2_idx = np.argmin(lows_15d[8:]) + 8  # Đáy 2 trong 7 ngày sau
            
            min1 = lows_15d[min1_idx]
            min2 = lows_15d[min2_idx]
            
            # Hai đáy gần bằng nhau (±3%)
            if abs(min1 - min2) / min1 < 0.03 and min2_idx - min1_idx >= 5:
                # Kiểm tra giá hiện tại đã vượt neckline chưa
                neckline = max(df['high'].tail(15).values[min1_idx:min2_idx])
                if close > neckline:
                    advanced_signals['pattern_signals'].append({
                        'pattern': 'DOUBLE_BOTTOM_BREAKOUT',
                        'interpretation': '📈 DOUBLE BOTTOM đã breakout → Tín hiệu MUA MẠNH',
                        'action': 'MUA',
                        'score': 3
                    })
                elif close > min2 * 1.02:
                    advanced_signals['pattern_signals'].append({
                        'pattern': 'DOUBLE_BOTTOM_FORMING',
                        'interpretation': '📈 DOUBLE BOTTOM đang hình thành → Theo dõi breakout',
                        'action': 'CHUẨN BỊ MUA',
                        'score': 1
                    })
    
    # ============================================================
    # 6. VÙNG HỘI TỤ (CONFLUENCE ZONES)
    # ============================================================
    
    support_levels = []
    resistance_levels = []
    
    # Thu thập các mức hỗ trợ/kháng cự từ các chỉ báo
    if available.get('BB'):
        support_levels.append(('BB Lower', ind.get('BB_Lower', close * 0.95)))
        resistance_levels.append(('BB Upper', ind.get('BB_Upper', close * 1.05)))
        support_levels.append(('BB Middle', ind.get('BB_Middle', close)))
        resistance_levels.append(('BB Middle', ind.get('BB_Middle', close)))
    
    if available.get('SMA'):
        for sma_key in ['SMA_20', 'SMA_50', 'SMA_100', 'SMA_200']:
            if available.get(sma_key) and ind.get(sma_key):
                if ind[sma_key] < close:
                    support_levels.append((sma_key, ind[sma_key]))
                else:
                    resistance_levels.append((sma_key, ind[sma_key]))
    
    if available.get('EMA'):
        for ema_key in ['EMA_12', 'EMA_26', 'EMA_50']:
            if available.get(ema_key) and ind.get(ema_key):
                if ind[ema_key] < close:
                    support_levels.append((ema_key, ind[ema_key]))
                else:
                    resistance_levels.append((ema_key, ind[ema_key]))
    
    if available.get('VWAP') and ind.get('VWAP'):
        if ind['VWAP'] < close:
            support_levels.append(('VWAP', ind['VWAP']))
        else:
            resistance_levels.append(('VWAP', ind['VWAP']))
    
    if available.get('SAR') and ind.get('SAR'):
        if ind['SAR'] < close:
            support_levels.append(('SAR', ind['SAR']))
        else:
            resistance_levels.append(('SAR', ind['SAR']))
    
    # Tìm vùng hội tụ (nhiều mức gần nhau)
    def find_confluence(levels, tolerance_pct=1.5):
        if not levels:
            return []
        
        sorted_levels = sorted(levels, key=lambda x: x[1])
        confluences = []
        
        i = 0
        while i < len(sorted_levels):
            cluster = [sorted_levels[i]]
            base_price = sorted_levels[i][1]
            
            j = i + 1
            while j < len(sorted_levels):
                if abs(sorted_levels[j][1] - base_price) / base_price * 100 < tolerance_pct:
                    cluster.append(sorted_levels[j])
                    j += 1
                else:
                    break
            
            if len(cluster) >= 2:
                avg_price = np.mean([c[1] for c in cluster])
                confluences.append({
                    'price': avg_price,
                    'indicators': [c[0] for c in cluster],
                    'count': len(cluster),
                    'strength': 'MẠNH' if len(cluster) >= 3 else 'TRUNG BÌNH'
                })
            
            i = j if j > i + 1 else i + 1
        
        return confluences
    
    support_confluences = find_confluence(support_levels)
    resistance_confluences = find_confluence(resistance_levels)
    
    for conf in support_confluences:
        distance_pct = (close - conf['price']) / close * 100
        advanced_signals['confluence_zones'].append({
            'type': 'SUPPORT',
            'price': conf['price'],
            'indicators': conf['indicators'],
            'strength': conf['strength'],
            'distance_pct': distance_pct,
            'interpretation': f"📍 Vùng HỖ TRỢ {conf['strength']} tại {conf['price']:,.0f} ({distance_pct:.1f}% dưới giá) - {', '.join(conf['indicators'])}"
        })
    
    for conf in resistance_confluences:
        distance_pct = (conf['price'] - close) / close * 100
        advanced_signals['confluence_zones'].append({
            'type': 'RESISTANCE',
            'price': conf['price'],
            'indicators': conf['indicators'],
            'strength': conf['strength'],
            'distance_pct': distance_pct,
            'interpretation': f"📍 Vùng KHÁNG CỰ {conf['strength']} tại {conf['price']:,.0f} ({distance_pct:.1f}% trên giá) - {', '.join(conf['indicators'])}"
        })
    
    # ============================================================
    # 7. TÍNH ĐIỂM TỔNG HỢP TỪ PHÂN TÍCH NÂNG CAO
    # ============================================================
    
    total_score = 0
    
    # Điểm từ phân kỳ
    for div in advanced_signals['divergences']:
        total_score += div['score']
    
    # Điểm từ xu hướng chỉ báo
    for trend in advanced_signals['indicator_trends']:
        total_score += trend['score']
    
    # Điểm từ crossover
    for cross in advanced_signals['crossover_timing']:
        total_score += cross['score']
    
    # Điểm từ BB analysis
    if advanced_signals['bb_analysis'].get('score'):
        total_score += advanced_signals['bb_analysis']['score']
    
    # Điểm từ pattern
    for pattern in advanced_signals['pattern_signals']:
        total_score += pattern['score']
    
    advanced_signals['strength_score'] = total_score
    
    return advanced_signals


# ============================================================
# HÀM DỰ BÁO ĐA KHUNG THỜI GIAN: T0-T5, W1-W4, M1-M3
# VỚI HỆ THỐNG ĐIỂM CÓ TRỌNG SỐ - ĐẦY ĐỦ 26 CHỈ BÁO
# ============================================================

def forecast_multi_timeframe(df, symbol):
    """
    Dự báo đa khung thời gian với:
    - T0 (hiện tại) đến T5 (5 ngày)
    - W1-W4 (1-4 tuần)
    - M1-M3 (1-3 tháng)
    - Hệ thống điểm có trọng số - ĐẦY ĐỦ 26 CHỈ BÁO
    - Giải thích chi tiết từng chỉ báo
    - Hành động dựa trên NHIỀU chỉ báo
    """
    
    if df is None or len(df) < 30:
        return {'symbol': symbol, 'error': 'Không đủ dữ liệu (cần ≥30 ngày)'}
    
    df = df.sort_values('time').reset_index(drop=True)
    current = df.iloc[-1]
    close = float(current['close'])
    
    # ============================================================
    # HÀM LẤY GIÁ TRỊ AN TOÀN
    # ============================================================
    
    def safe_get(col, default=None):
        if col in df.columns and pd.notna(current.get(col)):
            return float(current[col])
        return default
    
    # ============================================================
    # LẤY GIÁ TRỊ 26 CHỈ BÁO
    # ============================================================
    
    ind = {
        # Xu hướng
        'SMA_5': safe_get('SMA_5', close),
        'SMA_10': safe_get('SMA_10', close),
        'SMA_20': safe_get('SMA_20', close),
        'SMA_50': safe_get('SMA_50', close),
        'SMA_100': safe_get('SMA_100', close),
        'SMA_200': safe_get('SMA_200', close),
        'EMA_12': safe_get('EMA_12', close),
        'EMA_26': safe_get('EMA_26', close),
        'EMA_50': safe_get('EMA_50', close),
        'WMA_10': safe_get('WMA_10', close),
        'WMA_20': safe_get('WMA_20', close),
        'TEMA_20': safe_get('TEMA_20', close),
        'DEMA_20': safe_get('DEMA_20', close),
        'MACD': safe_get('MACD', 0),
        'MACD_Signal': safe_get('MACD_Signal', 0),
        'MACD_Hist': safe_get('MACD_Hist', 0),
        'SAR': safe_get('SAR', close),
        # Động lượng
        'RSI': safe_get('RSI', 50),
        'Stoch_K': safe_get('Stoch_K', 50),
        'Stoch_D': safe_get('Stoch_D', 50),
        'StochRSI': safe_get('StochRSI', 50),
        'ROC': safe_get('ROC', 0),
        'MOM': safe_get('Momentum', 0),
        # Dao động
        'CCI': safe_get('CCI', 0),
        'Williams_R': safe_get('Williams_R', -50),
        'ADX': safe_get('ADX', 25),
        'Plus_DI': safe_get('Plus_DI', 25),
        'Minus_DI': safe_get('Minus_DI', 25),
        'ATR': safe_get('ATR', close * 0.02),
        'BB_Upper': safe_get('BB_Upper', close * 1.05),
        'BB_Middle': safe_get('BB_Middle', close),
        'BB_Lower': safe_get('BB_Lower', close * 0.95),
        # Khối lượng
        'OBV': safe_get('OBV', 0),
        'MFI': safe_get('MFI', 50),
        'CMF': safe_get('CMF', 0),
        'AD': safe_get('AD', 0),
        'VWAP': safe_get('VWAP', close),
        'FI': safe_get('FI', 0)
    }
    # ============================================================
    # KIỂM TRA CHỈ BÁO CÓ SẴN TRONG DATA (ĐẦY ĐỦ)
    # ============================================================
    
    available = {
        # Xu hướng (7)
        'SMA': any(col in df.columns for col in ['SMA_5', 'SMA_10', 'SMA_20', 'SMA_50', 'SMA_100', 'SMA_200']),
        'SMA_5': 'SMA_5' in df.columns,
        'SMA_10': 'SMA_10' in df.columns,
        'SMA_20': 'SMA_20' in df.columns,
        'SMA_50': 'SMA_50' in df.columns,
        'SMA_100': 'SMA_100' in df.columns,
        'SMA_200': 'SMA_200' in df.columns,
        
        'EMA': any(col in df.columns for col in ['EMA_12', 'EMA_26', 'EMA_50']),
        'EMA_12': 'EMA_12' in df.columns,
        'EMA_26': 'EMA_26' in df.columns,
        'EMA_50': 'EMA_50' in df.columns,
        
        'WMA': any(col in df.columns for col in ['WMA_10', 'WMA_20']),
        'WMA_10': 'WMA_10' in df.columns,
        'WMA_20': 'WMA_20' in df.columns,
        
        'TEMA': 'TEMA_20' in df.columns,
        'DEMA': 'DEMA_20' in df.columns,
        
        'MACD': 'MACD' in df.columns,
        'MACD_Signal': 'MACD_Signal' in df.columns,
        'MACD_Hist': 'MACD_Hist' in df.columns,
        'MACD_Cross': 'MACD_Cross' in df.columns,
        
        'SAR': 'SAR' in df.columns,
        
        # Động lượng (5)
        'RSI': 'RSI' in df.columns,
        
        'STOCH': 'Stoch_K' in df.columns,
        'Stoch_K': 'Stoch_K' in df.columns,
        'Stoch_D': 'Stoch_D' in df.columns,
        
        'STOCHRSI': 'StochRSI' in df.columns,
        'ROC': 'ROC' in df.columns,
        'MOM': 'Momentum' in df.columns,
        
        # Dao động (5)
        'CCI': 'CCI' in df.columns,
        'WILLR': 'Williams_R' in df.columns,
        
        'ADX': 'ADX' in df.columns,
        'Plus_DI': 'Plus_DI' in df.columns,
        'Minus_DI': 'Minus_DI' in df.columns,
        
        'ATR': 'ATR' in df.columns,
        
        'BB': 'BB_Upper' in df.columns,
        'BB_Upper': 'BB_Upper' in df.columns,
        'BB_Middle': 'BB_Middle' in df.columns,
        'BB_Lower': 'BB_Lower' in df.columns,
        
        # Khối lượng (6)
        'OBV': 'OBV' in df.columns,
        'MFI': 'MFI' in df.columns,
        'CMF': 'CMF' in df.columns,
        'AD': 'AD' in df.columns,
        'VWAP': 'VWAP' in df.columns,
        'FI': 'FI' in df.columns,
    }
    
    # Đếm số chỉ báo chính có sẵn (23 nhóm)
    main_indicators = ['SMA', 'EMA', 'WMA', 'TEMA', 'DEMA', 'MACD', 'SAR',
                       'RSI', 'STOCH', 'STOCHRSI', 'ROC', 'MOM',
                       'CCI', 'WILLR', 'ADX', 'ATR', 'BB',
                       'OBV', 'MFI', 'CMF', 'AD', 'VWAP', 'FI']
    available_count = sum(1 for ind in main_indicators if available.get(ind, False))
    total_main = len(main_indicators)  # 23
    
    # ============================================================
    # KHỞI TẠO BIẾN CHỈ BÁO ĐỂ DÙNG TOÀN HÀM
    # ============================================================
    
    # Khởi tạo biến từ dictionary ind
    rsi = ind['RSI']
    macd = ind['MACD']
    macd_signal = ind['MACD_Signal']
    macd_hist = ind['MACD_Hist']
    stoch_k = ind['Stoch_K']
    stoch_d = ind['Stoch_D']
    stoch_rsi = ind['StochRSI']
    mfi = ind['MFI']
    obv = ind['OBV']
    adx = ind['ADX']
    plus_di = ind['Plus_DI']
    minus_di = ind['Minus_DI']
    cci = ind['CCI']
    willr = ind['Williams_R']
    roc = ind['ROC']
    mom = ind['MOM']
    
    sma_5 = ind['SMA_5']
    sma_10 = ind['SMA_10']
    sma_20 = ind['SMA_20']
    sma_50 = ind['SMA_50']
    sma_100 = ind['SMA_100']
    sma_200 = ind['SMA_200']
    ema_12 = ind['EMA_12']
    ema_26 = ind['EMA_26']
    ema_50 = ind['EMA_50']
    wma_10 = ind['WMA_10']
    wma_20 = ind['WMA_20']
    tema_20 = ind['TEMA_20']
    dema_20 = ind['DEMA_20']
    
    bb_upper = ind['BB_Upper']
    bb_middle = ind['BB_Middle']
    bb_lower = ind['BB_Lower']
    
    atr = ind['ATR']
    sar = ind['SAR']
    
    vwap = ind['VWAP']
    cmf = ind['CMF']
    ad = ind['AD']
    fi = ind['FI']


    
    # ============================================================
    # TÍNH ĐIỂM CÓ TRỌNG SỐ CHO TẤT CẢ 26 CHỈ BÁO
    # ============================================================
    
    weighted_scores = {}
    
    # ----- 1. SMA (trọng số 5) -----
    if available['SMA']:
        sma_score = 50
        sma_reasons = []
        
        if available['SMA_20'] and close > ind['SMA_20']:
            sma_score += 10
            sma_reasons.append("Giá > SMA20")
        elif available['SMA_20']:
            sma_score -= 10
            sma_reasons.append("Giá < SMA20")
        
        if available['SMA_50'] and close > ind['SMA_50']:
            sma_score += 15
            sma_reasons.append("Giá > SMA50")
        elif available['SMA_50']:
            sma_score -= 15
            sma_reasons.append("Giá < SMA50")
        
        if available['SMA_200'] and close > ind['SMA_200']:
            sma_score += 20
            sma_reasons.append("Giá > SMA200")
        elif available['SMA_200']:
            sma_score -= 20
            sma_reasons.append("Giá < SMA200")
        
        # Golden/Death Cross
        if available['SMA_50'] and available['SMA_200'] and len(df) > 1:
            prev_sma50 = df['SMA_50'].iloc[-2] if pd.notna(df['SMA_50'].iloc[-2]) else ind['SMA_50']
            prev_sma200 = df['SMA_200'].iloc[-2] if pd.notna(df['SMA_200'].iloc[-2]) else ind['SMA_200']
            
            if prev_sma50 < prev_sma200 and ind['SMA_50'] > ind['SMA_200']:
                sma_score += 25
                sma_reasons.append("🌟 GOLDEN CROSS")
            elif prev_sma50 > prev_sma200 and ind['SMA_50'] < ind['SMA_200']:
                sma_score -= 25
                sma_reasons.append("💀 DEATH CROSS")
        
        sma_score = max(0, min(100, sma_score))
        weighted_scores['SMA'] = {
            'score': sma_score,
            'weight': INDICATOR_WEIGHTS['SMA'],
            'signal': 'BULLISH' if sma_score > 50 else 'BEARISH',
            'reason': "; ".join(sma_reasons) if sma_reasons else "Không đủ dữ liệu SMA"
        }
    
    # ----- 2. EMA (trọng số 4) -----
    if available['EMA']:
        ema_score = 50
        ema_reasons = []
        
        if available['EMA_12'] and available['EMA_26']:
            if ind['EMA_12'] > ind['EMA_26']:
                ema_score += 25
                ema_reasons.append("EMA12 > EMA26")
            else:
                ema_score -= 25
                ema_reasons.append("EMA12 < EMA26")
        
        if available['EMA_50']:
            if close > ind['EMA_50']:
                ema_score += 15
                ema_reasons.append("Giá > EMA50")
            else:
                ema_score -= 15
                ema_reasons.append("Giá < EMA50")
        
        ema_score = max(0, min(100, ema_score))
        weighted_scores['EMA'] = {
            'score': ema_score,
            'weight': INDICATOR_WEIGHTS['EMA'],
            'signal': 'BULLISH' if ema_score > 50 else 'BEARISH',
            'reason': "; ".join(ema_reasons) if ema_reasons else "Không đủ dữ liệu EMA"
        }
    
    # ----- 3. WMA (trọng số 2) -----
    if available['WMA']:
        wma_score = 50
        wma_reasons = []
        
        if available['WMA_10'] and available['WMA_20']:
            if ind['WMA_10'] > ind['WMA_20']:
                wma_score += 25
                wma_reasons.append("WMA10 > WMA20")
            else:
                wma_score -= 25
                wma_reasons.append("WMA10 < WMA20")
        
        if available['WMA_10']:
            if close > ind['WMA_10']:
                wma_score += 15
                wma_reasons.append("Giá > WMA10")
            else:
                wma_score -= 15
                wma_reasons.append("Giá < WMA10")
        
        wma_score = max(0, min(100, wma_score))
        weighted_scores['WMA'] = {
            'score': wma_score,
            'weight': INDICATOR_WEIGHTS['WMA'],
            'signal': 'BULLISH' if wma_score > 50 else 'BEARISH',
            'reason': "; ".join(wma_reasons) if wma_reasons else "Không đủ dữ liệu WMA"
        }
    
    
    # ----- 4. TEMA (trọng số 2) -----
    if 'TEMA_20' in df.columns:
        tema_score = 50
        tema_reasons = []
        
        if close > ind['TEMA_20']:
            tema_score += 30
            tema_reasons.append(f"Giá > TEMA20")
        else:
            tema_score -= 30
            tema_reasons.append(f"Giá < TEMA20")
        
        tema_score = max(0, min(100, tema_score))
        weighted_scores['TEMA'] = {
            'score': tema_score,
            'weight': INDICATOR_WEIGHTS['TEMA'],
            'signal': 'BULLISH' if tema_score > 50 else 'BEARISH',
            'reason': "; ".join(tema_reasons)
        }
    
    # ----- 5. DEMA (trọng số 2) -----
    if 'DEMA_20' in df.columns:
        dema_score = 50
        dema_reasons = []
        
        if close > ind['DEMA_20']:
            dema_score += 30
            dema_reasons.append(f"Giá > DEMA20")
        else:
            dema_score -= 30
            dema_reasons.append(f"Giá < DEMA20")
        
        dema_score = max(0, min(100, dema_score))
        weighted_scores['DEMA'] = {
            'score': dema_score,
            'weight': INDICATOR_WEIGHTS['DEMA'],
            'signal': 'BULLISH' if dema_score > 50 else 'BEARISH',
            'reason': "; ".join(dema_reasons)
        }
    
    # ----- 6. MACD (trọng số 8) -----
    if available['MACD']:
        macd_score = 50
        macd_reasons = []
        
        if available['MACD_Signal']:
            if ind['MACD'] > ind['MACD_Signal']:
                macd_score += 20
                macd_reasons.append("MACD > Signal")
            else:
                macd_score -= 20
                macd_reasons.append("MACD < Signal")
        
        if available['MACD_Hist']:
            if ind['MACD_Hist'] > 0:
                macd_score += 15
                macd_reasons.append("Histogram > 0")
            else:
                macd_score -= 15
                macd_reasons.append("Histogram < 0")
            
            # MACD Crossover
            if len(df) > 1:
                prev_hist = df['MACD_Hist'].iloc[-2] if pd.notna(df['MACD_Hist'].iloc[-2]) else 0
                if prev_hist < 0 and ind['MACD_Hist'] > 0:
                    macd_score += 20
                    macd_reasons.append("🔼 MACD Cross Up")
                elif prev_hist > 0 and ind['MACD_Hist'] < 0:
                    macd_score -= 20
                    macd_reasons.append("🔽 MACD Cross Down")
        
        macd_score = max(0, min(100, macd_score))
        weighted_scores['MACD'] = {
            'score': macd_score,
            'weight': INDICATOR_WEIGHTS['MACD'],
            'signal': 'BULLISH' if macd_score > 50 else 'BEARISH',
            'reason': "; ".join(macd_reasons) if macd_reasons else "Không đủ dữ liệu MACD"
        }
    
    # ----- 7. SAR (trọng số 3) -----
    if 'SAR' in df.columns:
        sar_score = 50
        sar_reasons = []
        
        if close > ind['SAR']:
            sar_score += 30
            sar_reasons.append("Giá > SAR (Uptrend)")
        else:
            sar_score -= 30
            sar_reasons.append("Giá < SAR (Downtrend)")
        
        # Đảo chiều SAR
        if len(df) > 1:
            prev_sar = df['SAR'].iloc[-2] if pd.notna(df['SAR'].iloc[-2]) else ind['SAR']
            prev_close = df['close'].iloc[-2]
            if prev_close < prev_sar and close > ind['SAR']:
                sar_score += 20
                sar_reasons.append("🔄 SAR đảo chiều LÊN")
            elif prev_close > prev_sar and close < ind['SAR']:
                sar_score -= 20
                sar_reasons.append("🔄 SAR đảo chiều XUỐNG")
        
        sar_score = max(0, min(100, sar_score))
        weighted_scores['SAR'] = {
            'score': sar_score,
            'weight': INDICATOR_WEIGHTS['SAR'],
            'signal': 'BULLISH' if sar_score > 50 else 'BEARISH',
            'reason': "; ".join(sar_reasons)
        }
    
    # ----- 8. RSI (trọng số 8) -----
    rsi = ind['RSI']
    rsi_reasons = []
    
    if rsi < 30:
        rsi_score = 85
        rsi_reasons.append(f"RSI={rsi:.1f} < 30: QUÁ BÁN")
    elif rsi < 40:
        rsi_score = 70
        rsi_reasons.append(f"RSI={rsi:.1f}: Vùng thấp")
    elif rsi < 50:
        rsi_score = 55
        rsi_reasons.append(f"RSI={rsi:.1f}: Hơi yếu")
    elif rsi < 60:
        rsi_score = 50
        rsi_reasons.append(f"RSI={rsi:.1f}: Trung lập")
    elif rsi < 70:
        rsi_score = 40
        rsi_reasons.append(f"RSI={rsi:.1f}: Hơi mạnh")
    elif rsi < 80:
        rsi_score = 25
        rsi_reasons.append(f"RSI={rsi:.1f} > 70: QUÁ MUA")
    else:
        rsi_score = 10
        rsi_reasons.append(f"RSI={rsi:.1f} > 80: RẤT QUÁ MUA")
    
    if 'RSI' in df.columns:
        weighted_scores['RSI'] = {
            'score': rsi_score,
            'weight': INDICATOR_WEIGHTS['RSI'],
            'signal': 'OVERSOLD' if rsi < 30 else ('OVERBOUGHT' if rsi > 70 else 'NEUTRAL'),
            'reason': "; ".join(rsi_reasons)
        }
    
    # ----- 9. STOCH (trọng số 6) -----
    if available['STOCH']:
        stoch_k = ind['Stoch_K']
        stoch_d = ind['Stoch_D'] if available['Stoch_D'] else stoch_k
        stoch_reasons = []
        
        if stoch_k < 20:
            stoch_score = 85
            stoch_reasons.append(f"Stoch_K={stoch_k:.1f} < 20: QUÁ BÁN")
        elif stoch_k < 50:
            stoch_score = 60
            stoch_reasons.append(f"Stoch_K={stoch_k:.1f}: Vùng thấp")
        elif stoch_k < 80:
            stoch_score = 40
            stoch_reasons.append(f"Stoch_K={stoch_k:.1f}: Vùng cao")
        else:
            stoch_score = 15
            stoch_reasons.append(f"Stoch_K={stoch_k:.1f} > 80: QUÁ MUA")
        
        if available['Stoch_D']:
            if stoch_k > stoch_d:
                stoch_score += 10
                stoch_reasons.append("K > D")
            else:
                stoch_score -= 10
                stoch_reasons.append("K < D")
        
        stoch_score = max(0, min(100, stoch_score))
        weighted_scores['STOCH'] = {
            'score': stoch_score,
            'weight': INDICATOR_WEIGHTS['STOCH'],
            'signal': 'OVERSOLD' if stoch_k < 20 else ('OVERBOUGHT' if stoch_k > 80 else 'NEUTRAL'),
            'reason': "; ".join(stoch_reasons)
        }
    
    # ----- 10. STOCHRSI (trọng số 4) -----
    if 'StochRSI' in df.columns:
        stoch_rsi = ind['StochRSI']
        stoch_rsi_reasons = []
        
        if stoch_rsi < 20:
            stoch_rsi_score = 85
            stoch_rsi_reasons.append(f"StochRSI={stoch_rsi:.1f}: QUÁ BÁN")
        elif stoch_rsi < 50:
            stoch_rsi_score = 60
            stoch_rsi_reasons.append(f"StochRSI={stoch_rsi:.1f}: Vùng thấp")
        elif stoch_rsi < 80:
            stoch_rsi_score = 40
            stoch_rsi_reasons.append(f"StochRSI={stoch_rsi:.1f}: Vùng cao")
        else:
            stoch_rsi_score = 15
            stoch_rsi_reasons.append(f"StochRSI={stoch_rsi:.1f}: QUÁ MUA")
        
        weighted_scores['STOCHRSI'] = {
            'score': stoch_rsi_score,
            'weight': INDICATOR_WEIGHTS['STOCHRSI'],
            'signal': 'OVERSOLD' if stoch_rsi < 20 else ('OVERBOUGHT' if stoch_rsi > 80 else 'NEUTRAL'),
            'reason': "; ".join(stoch_rsi_reasons)
        }
    
    # ----- 11. ROC (trọng số 3) -----
    if 'ROC' in df.columns:
        roc = ind['ROC']
        roc_reasons = []
        
        if roc > 10:
            roc_score = 80
            roc_reasons.append(f"ROC={roc:.2f}%: Momentum MẠNH")
        elif roc > 3:
            roc_score = 65
            roc_reasons.append(f"ROC={roc:.2f}%: Tăng")
        elif roc > 0:
            roc_score = 55
            roc_reasons.append(f"ROC={roc:.2f}%: Tăng nhẹ")
        elif roc > -3:
            roc_score = 45
            roc_reasons.append(f"ROC={roc:.2f}%: Giảm nhẹ")
        elif roc > -10:
            roc_score = 35
            roc_reasons.append(f"ROC={roc:.2f}%: Giảm")
        else:
            roc_score = 20
            roc_reasons.append(f"ROC={roc:.2f}%: Momentum YẾU")
        
        weighted_scores['ROC'] = {
            'score': roc_score,
            'weight': INDICATOR_WEIGHTS['ROC'],
            'signal': 'POSITIVE' if roc > 0 else 'NEGATIVE',
            'reason': "; ".join(roc_reasons)
        }
    
    # ----- 12. MOM (trọng số 4) -----
    if 'Momentum' in df.columns:
        mom = ind['MOM']
        mom_reasons = []
        
        mom_pct = (mom / close) * 100 if close > 0 else 0
        if mom > 0:
            mom_score = 50 + min(mom_pct * 10, 40)
            mom_reasons.append(f"Momentum={mom:.2f} (+{mom_pct:.2f}%)")
        else:
            mom_score = 50 - min(abs(mom_pct) * 10, 40)
            mom_reasons.append(f"Momentum={mom:.2f} ({mom_pct:.2f}%)")
        
        mom_score = max(0, min(100, mom_score))
        weighted_scores['MOM'] = {
            'score': mom_score,
            'weight': INDICATOR_WEIGHTS['MOM'],
            'signal': 'POSITIVE' if mom > 0 else 'NEGATIVE',
            'reason': "; ".join(mom_reasons)
        }
    
    # ----- 13. CCI (trọng số 4) -----
    if 'CCI' in df.columns:
        cci = ind['CCI']
        cci_reasons = []
        
        if cci < -200:
            cci_score = 90
            cci_reasons.append(f"CCI={cci:.1f}: CỰC KỲ QUÁ BÁN")
        elif cci < -100:
            cci_score = 75
            cci_reasons.append(f"CCI={cci:.1f}: Quá bán")
        elif cci < 0:
            cci_score = 55
            cci_reasons.append(f"CCI={cci:.1f}: Vùng âm")
        elif cci < 100:
            cci_score = 45
            cci_reasons.append(f"CCI={cci:.1f}: Vùng dương")
        elif cci < 200:
            cci_score = 25
            cci_reasons.append(f"CCI={cci:.1f}: Quá mua")
        else:
            cci_score = 10
            cci_reasons.append(f"CCI={cci:.1f}: CỰC KỲ QUÁ MUA")
        
        weighted_scores['CCI'] = {
            'score': cci_score,
            'weight': INDICATOR_WEIGHTS['CCI'],
            'signal': 'OVERSOLD' if cci < -100 else ('OVERBOUGHT' if cci > 100 else 'NEUTRAL'),
            'reason': "; ".join(cci_reasons)
        }
    
    # ----- 14. WILLR (trọng số 4) -----
    if 'Williams_R' in df.columns:
        willr = ind['Williams_R']
        willr_reasons = []
        
        if willr < -80:
            willr_score = 80
            willr_reasons.append(f"Williams %R={willr:.1f}: Quá bán")
        elif willr < -50:
            willr_score = 60
            willr_reasons.append(f"Williams %R={willr:.1f}: Vùng thấp")
        elif willr < -20:
            willr_score = 40
            willr_reasons.append(f"Williams %R={willr:.1f}: Vùng cao")
        else:
            willr_score = 20
            willr_reasons.append(f"Williams %R={willr:.1f}: Quá mua")
        
        weighted_scores['WILLR'] = {
            'score': willr_score,
            'weight': INDICATOR_WEIGHTS['WILLR'],
            'signal': 'OVERSOLD' if willr < -80 else ('OVERBOUGHT' if willr > -20 else 'NEUTRAL'),
            'reason': "; ".join(willr_reasons)
        }
    
    # ----- 15. ADX (trọng số 6) -----
    if available['ADX']:
        adx = ind['ADX']
        plus_di = ind['Plus_DI'] if available['Plus_DI'] else 25
        minus_di = ind['Minus_DI'] if available['Minus_DI'] else 25
        adx_reasons = []
        
        if adx < 20:
            adx_score = 50
            adx_reasons.append(f"ADX={adx:.1f}: Không có trend")
        elif adx < 25:
            adx_score = 55 if plus_di > minus_di else 45
            adx_reasons.append(f"ADX={adx:.1f}: Trend yếu")
        elif adx < 50:
            if available['Plus_DI'] and available['Minus_DI']:
                if plus_di > minus_di:
                    adx_score = 70
                    adx_reasons.append(f"ADX={adx:.1f}: Uptrend mạnh (+DI > -DI)")
                else:
                    adx_score = 30
                    adx_reasons.append(f"ADX={adx:.1f}: Downtrend mạnh (-DI > +DI)")
            else:
                adx_score = 50
                adx_reasons.append(f"ADX={adx:.1f}: Trend mạnh (thiếu DI)")
        else:
            if available['Plus_DI'] and available['Minus_DI']:
                if plus_di > minus_di:
                    adx_score = 80
                    adx_reasons.append(f"ADX={adx:.1f}: Uptrend RẤT MẠNH")
                else:
                    adx_score = 20
                    adx_reasons.append(f"ADX={adx:.1f}: Downtrend RẤT MẠNH")
            else:
                adx_score = 50
                adx_reasons.append(f"ADX={adx:.1f}: Trend rất mạnh (thiếu DI)")
        
        weighted_scores['ADX'] = {
            'score': adx_score,
            'weight': INDICATOR_WEIGHTS['ADX'],
            'signal': 'STRONG_UP' if adx_score > 60 else ('STRONG_DOWN' if adx_score < 40 else 'WEAK'),
            'reason': "; ".join(adx_reasons)
        }
    
    # ----- 16. ATR (trọng số 5) -----
    if 'ATR' in df.columns:
        atr = ind['ATR']
        atr_pct = (atr / close) * 100 if close > 0 else 2
        atr_reasons = []
        
        if atr_pct < 1.5:
            atr_score = 60
            atr_reasons.append(f"ATR={atr_pct:.2f}%: Biến động thấp")
        elif atr_pct < 3:
            atr_score = 50
            atr_reasons.append(f"ATR={atr_pct:.2f}%: Biến động TB")
        elif atr_pct < 5:
            atr_score = 40
            atr_reasons.append(f"ATR={atr_pct:.2f}%: Biến động cao")
        else:
            atr_score = 30
            atr_reasons.append(f"ATR={atr_pct:.2f}%: Biến động RẤT CAO")
        
        weighted_scores['ATR'] = {
            'score': atr_score,
            'weight': INDICATOR_WEIGHTS['ATR'],
            'signal': 'LOW_VOL' if atr_pct < 2 else ('HIGH_VOL' if atr_pct > 4 else 'NORMAL'),
            'reason': "; ".join(atr_reasons)
        }
    
    # ----- 17. BB (trọng số 6) -----
    if available['BB']:
        bb_upper = ind['BB_Upper'] if available['BB_Upper'] else close * 1.05
        bb_lower = ind['BB_Lower'] if available['BB_Lower'] else close * 0.95
        bb_middle = ind['BB_Middle'] if available['BB_Middle'] else close
        bb_reasons = []
        
        bb_range = bb_upper - bb_lower
        bb_position = (close - bb_lower) / bb_range * 100 if bb_range > 0 else 50
        
        if close < bb_lower:
            bb_score = 90
            bb_reasons.append("Giá DƯỚI BB Lower")
        elif bb_position < 20:
            bb_score = 75
            bb_reasons.append(f"Gần BB Lower ({bb_position:.0f}%)")
        elif bb_position < 40:
            bb_score = 60
            bb_reasons.append(f"Dưới middle ({bb_position:.0f}%)")
        elif bb_position < 60:
            bb_score = 50
            bb_reasons.append(f"Quanh middle ({bb_position:.0f}%)")
        elif bb_position < 80:
            bb_score = 40
            bb_reasons.append(f"Trên middle ({bb_position:.0f}%)")
        elif close > bb_upper:
            bb_score = 10
            bb_reasons.append("Giá TRÊN BB Upper")
        else:
            bb_score = 25
            bb_reasons.append(f"Gần BB Upper ({bb_position:.0f}%)")
        
        weighted_scores['BB'] = {
            'score': bb_score,
            'weight': INDICATOR_WEIGHTS['BB'],
            'signal': 'OVERSOLD' if bb_position < 20 else ('OVERBOUGHT' if bb_position > 80 else 'NEUTRAL'),
            'reason': "; ".join(bb_reasons)
        }
    
    # ----- 18. OBV (trọng số 5) -----
    if 'OBV' in df.columns and len(df) > 5:
        obv_now = ind['OBV']
        obv_5d = df['OBV'].iloc[-6] if pd.notna(df['OBV'].iloc[-6]) else obv_now
        obv_reasons = []
        
        obv_change = (obv_now - obv_5d) / abs(obv_5d) * 100 if obv_5d != 0 else 0
        
        if obv_change > 10:
            obv_score = 75
            obv_reasons.append(f"OBV tăng {obv_change:.1f}%: Tích lũy mạnh")
        elif obv_change > 0:
            obv_score = 60
            obv_reasons.append(f"OBV tăng {obv_change:.1f}%")
        elif obv_change > -10:
            obv_score = 40
            obv_reasons.append(f"OBV giảm {obv_change:.1f}%")
        else:
            obv_score = 25
            obv_reasons.append(f"OBV giảm {obv_change:.1f}%: Phân phối mạnh")
        
        weighted_scores['OBV'] = {
            'score': obv_score,
            'weight': INDICATOR_WEIGHTS['OBV'],
            'signal': 'ACCUMULATION' if obv_score > 60 else ('DISTRIBUTION' if obv_score < 40 else 'NEUTRAL'),
            'reason': "; ".join(obv_reasons)
        }
    
    # ----- 19. MFI (trọng số 5) -----
    if 'MFI' in df.columns:
        mfi = ind['MFI']
        mfi_reasons = []
        
        if mfi < 20:
            mfi_score = 85
            mfi_reasons.append(f"MFI={mfi:.1f}: Dòng tiền QUÁ BÁN")
        elif mfi < 40:
            mfi_score = 65
            mfi_reasons.append(f"MFI={mfi:.1f}: Dòng tiền yếu")
        elif mfi < 60:
            mfi_score = 50
            mfi_reasons.append(f"MFI={mfi:.1f}: Trung lập")
        elif mfi < 80:
            mfi_score = 35
            mfi_reasons.append(f"MFI={mfi:.1f}: Dòng tiền mạnh")
        else:
            mfi_score = 15
            mfi_reasons.append(f"MFI={mfi:.1f}: Dòng tiền QUÁ MUA")
        
        weighted_scores['MFI'] = {
            'score': mfi_score,
            'weight': INDICATOR_WEIGHTS['MFI'],
            'signal': 'OVERSOLD' if mfi < 20 else ('OVERBOUGHT' if mfi > 80 else 'NEUTRAL'),
            'reason': "; ".join(mfi_reasons)
        }
    
    # ----- 20. CMF (trọng số 3) -----
    if 'CMF' in df.columns:
        cmf = ind['CMF']
        cmf_reasons = []
        
        if cmf > 0.2:
            cmf_score = 80
            cmf_reasons.append(f"CMF={cmf:.3f}: Áp lực mua MẠNH")
        elif cmf > 0.05:
            cmf_score = 65
            cmf_reasons.append(f"CMF={cmf:.3f}: Áp lực mua")
        elif cmf > -0.05:
            cmf_score = 50
            cmf_reasons.append(f"CMF={cmf:.3f}: Cân bằng")
        elif cmf > -0.2:
            cmf_score = 35
            cmf_reasons.append(f"CMF={cmf:.3f}: Áp lực bán")
        else:
            cmf_score = 20
            cmf_reasons.append(f"CMF={cmf:.3f}: Áp lực bán MẠNH")
        
        weighted_scores['CMF'] = {
            'score': cmf_score,
            'weight': INDICATOR_WEIGHTS['CMF'],
            'signal': 'BUYING' if cmf > 0.05 else ('SELLING' if cmf < -0.05 else 'NEUTRAL'),
            'reason': "; ".join(cmf_reasons)
        }
    
    # ----- 21. AD (trọng số 3) -----
    if 'AD' in df.columns and len(df) > 5:
        ad_now = ind['AD']
        ad_5d = df['AD'].iloc[-6] if pd.notna(df['AD'].iloc[-6]) else ad_now
        ad_reasons = []
        
        if ad_now > ad_5d:
            ad_change = (ad_now - ad_5d) / abs(ad_5d) * 100 if ad_5d != 0 else 0
            ad_score = 60 + min(ad_change / 5, 25)
            ad_reasons.append(f"A/D tăng: Tích lũy")
        else:
            ad_change = (ad_5d - ad_now) / abs(ad_5d) * 100 if ad_5d != 0 else 0
            ad_score = 40 - min(ad_change / 5, 25)
            ad_reasons.append(f"A/D giảm: Phân phối")
        
        ad_score = max(0, min(100, ad_score))
        weighted_scores['AD'] = {
            'score': ad_score,
            'weight': INDICATOR_WEIGHTS['AD'],
            'signal': 'ACCUMULATION' if ad_score > 50 else 'DISTRIBUTION',
            'reason': "; ".join(ad_reasons)
        }
    
    # ----- 22. VWAP (trọng số 2) -----
    if 'VWAP' in df.columns:
        vwap = ind['VWAP']
        vwap_reasons = []
        
        vwap_diff = (close - vwap) / vwap * 100 if vwap > 0 else 0
        
        if close > vwap:
            vwap_score = 55 + min(vwap_diff * 5, 30)
            vwap_reasons.append(f"Giá > VWAP (+{vwap_diff:.2f}%)")
        else:
            vwap_score = 45 - min(abs(vwap_diff) * 5, 30)
            vwap_reasons.append(f"Giá < VWAP ({vwap_diff:.2f}%)")
        
        vwap_score = max(0, min(100, vwap_score))
        weighted_scores['VWAP'] = {
            'score': vwap_score,
            'weight': INDICATOR_WEIGHTS['VWAP'],
            'signal': 'BULLISH' if vwap_score > 50 else 'BEARISH',
            'reason': "; ".join(vwap_reasons)
        }
    
    # ----- 23. FI (trọng số 2) -----
    if 'FI' in df.columns:
        fi = ind['FI']
        fi_reasons = []
        
        if fi > 0:
            fi_score = 60
            fi_reasons.append(f"Force Index={fi:,.0f}: Lực mua")
        else:
            fi_score = 40
            fi_reasons.append(f"Force Index={fi:,.0f}: Lực bán")
        
        # Trend FI
        if len(df) >= 3:
            fi_3d = df['FI'].iloc[-4] if pd.notna(df['FI'].iloc[-4]) else fi
            if fi > fi_3d and fi > 0:
                fi_score += 15
                fi_reasons.append("FI tăng")
            elif fi < fi_3d and fi < 0:
                fi_score -= 15
                fi_reasons.append("FI giảm")
        
        fi_score = max(0, min(100, fi_score))
        weighted_scores['FI'] = {
            'score': fi_score,
            'weight': INDICATOR_WEIGHTS['FI'],
            'signal': 'BUYING' if fi > 0 else 'SELLING',
            'reason': "; ".join(fi_reasons)
        }
    
    # ============================================================
    # TÍNH TỔNG ĐIỂM CÓ TRỌNG SỐ
    # ============================================================
    
    total_weighted_score = 0
    total_weight_used = 0
    
    for indicator, data in weighted_scores.items():
        total_weighted_score += data['score'] * data['weight']
        total_weight_used += data['weight']
    
    final_score = total_weighted_score / total_weight_used if total_weight_used > 0 else 50
    
     # ============================================================
    # ĐẾM TÍN HIỆU MUA/BÁN TỪ TẤT CẢ 26 CHỈ BÁO
    # ============================================================
    
    buy_signals = 0
    sell_signals = 0
    hold_signals = 0
    signal_details = []
    
    # ===== NHÓM 1: XU HƯỚNG (7 chỉ báo) =====
    
    # 1. SMA - CHỈ TÍNH NẾU CÓ
    if available['SMA']:
        sma_buy = 0
        sma_sell = 0
        
        # SMA20
        if available['SMA_20']:
            if close > ind['SMA_20']:
                sma_buy += 1
            else:
                sma_sell += 1
        
        # SMA50
        if available['SMA_50']:
            if close > ind['SMA_50']:
                sma_buy += 1
            else:
                sma_sell += 1
        
        # SMA200
        if available['SMA_200']:
            if close > ind['SMA_200']:
                sma_buy += 2
            else:
                sma_sell += 2
        
        # Golden/Death Cross
        if available['SMA_50'] and available['SMA_200'] and len(df) > 1:
            prev_sma50 = df['SMA_50'].iloc[-2] if pd.notna(df['SMA_50'].iloc[-2]) else ind['SMA_50']
            prev_sma200 = df['SMA_200'].iloc[-2] if pd.notna(df['SMA_200'].iloc[-2]) else ind['SMA_200']
            
            if prev_sma50 < prev_sma200 and ind['SMA_50'] > ind['SMA_200']:
                sma_buy += 3
                signal_details.append("🌟 Golden Cross (+3)")
            elif prev_sma50 > prev_sma200 and ind['SMA_50'] < ind['SMA_200']:
                sma_sell += 3
                signal_details.append("💀 Death Cross (-3)")
        
        if sma_buy > sma_sell:
            buy_signals += 1
            signal_details.append(f"SMA bullish (+1)")
        elif sma_sell > sma_buy:
            sell_signals += 1
            signal_details.append(f"SMA bearish (-1)")
        else:
            hold_signals += 1
    
    # 2. EMA - CHỈ TÍNH NẾU CÓ
    if available['EMA']:
        ema_buy = 0
        ema_sell = 0
        
        if available['EMA_12'] and available['EMA_26']:
            if ind['EMA_12'] > ind['EMA_26']:
                ema_buy += 1
            else:
                ema_sell += 1
        
        if available['EMA_50']:
            if close > ind['EMA_50']:
                ema_buy += 1
            else:
                ema_sell += 1
        
        if ema_buy > ema_sell:
            buy_signals += 1
            signal_details.append("EMA bullish (+1)")
        elif ema_sell > ema_buy:
            sell_signals += 1
        else:
            hold_signals += 1
    
    # 3. WMA - CHỈ TÍNH NẾU CÓ
    if available['WMA']:
        wma_buy = 0
        wma_sell = 0
        
        if available['WMA_10'] and available['WMA_20']:
            if ind['WMA_10'] > ind['WMA_20']:
                wma_buy += 1
            else:
                wma_sell += 1
        
        if available['WMA_10']:
            if close > ind['WMA_10']:
                wma_buy += 1
            else:
                wma_sell += 1
        
        if wma_buy > wma_sell:
            buy_signals += 1
            signal_details.append("WMA bullish (+1)")
        elif wma_sell > wma_buy:
            sell_signals += 1
        else:
            hold_signals += 1
    
    # 4. TEMA - CHỈ TÍNH NẾU CÓ
    if available['TEMA']:
        if close > ind['TEMA_20']:
            buy_signals += 1
            signal_details.append("TEMA bullish (+1)")
        else:
            sell_signals += 1
    
    # 5. DEMA - CHỈ TÍNH NẾU CÓ
    if available['DEMA']:
        if close > ind['DEMA_20']:
            buy_signals += 1
            signal_details.append("DEMA bullish (+1)")
        else:
            sell_signals += 1
    
    # 6. MACD - CHỈ TÍNH NẾU CÓ
    if available['MACD']:
        macd_buy = 0
        macd_sell = 0
        
        if available['MACD_Signal']:
            if ind['MACD'] > ind['MACD_Signal']:
                macd_buy += 1
            else:
                macd_sell += 1
        
        if available['MACD_Hist']:
            if ind['MACD_Hist'] > 0:
                macd_buy += 1
            else:
                macd_sell += 1
            
            # MACD Crossover
            if len(df) > 1:
                prev_hist = df['MACD_Hist'].iloc[-2] if pd.notna(df['MACD_Hist'].iloc[-2]) else 0
                if prev_hist < 0 and ind['MACD_Hist'] > 0:
                    macd_buy += 2
                    signal_details.append("🔼 MACD Cross Up (+2)")
                elif prev_hist > 0 and ind['MACD_Hist'] < 0:
                    macd_sell += 2
                    signal_details.append("🔽 MACD Cross Down (-2)")
        
        if macd_buy > macd_sell:
            buy_signals += 1
            signal_details.append("MACD bullish (+1)")
        elif macd_sell > macd_buy:
            sell_signals += 1
        else:
            hold_signals += 1
    
    # 7. SAR - CHỈ TÍNH NẾU CÓ
    if available['SAR']:
        if close > ind['SAR']:
            buy_signals += 1
            signal_details.append("SAR uptrend (+1)")
        else:
            sell_signals += 1
            signal_details.append("SAR downtrend (-1)")
        
        # Đảo chiều SAR
        if len(df) > 1:
            prev_sar = df['SAR'].iloc[-2] if pd.notna(df['SAR'].iloc[-2]) else ind['SAR']
            prev_close = df['close'].iloc[-2]
            if prev_close < prev_sar and close > ind['SAR']:
                buy_signals += 1
                signal_details.append("🔄 SAR đảo chiều LÊN (+1)")
            elif prev_close > prev_sar and close < ind['SAR']:
                sell_signals += 1
                signal_details.append("🔄 SAR đảo chiều XUỐNG (-1)")
    
    # ===== NHÓM 2: ĐỘNG LƯỢNG (5 chỉ báo) =====
    
    # 8. RSI - CHỈ TÍNH NẾU CÓ
    if available['RSI']:
        if rsi < 30:
            buy_signals += 2
            signal_details.append(f"RSI={rsi:.0f} quá bán (+2)")
        elif rsi < 40:
            buy_signals += 1
            signal_details.append(f"RSI={rsi:.0f} vùng thấp (+1)")
        elif rsi > 70:
            sell_signals += 2
            signal_details.append(f"RSI={rsi:.0f} quá mua (-2)")
        elif rsi > 60:
            sell_signals += 1
            signal_details.append(f"RSI={rsi:.0f} vùng cao (-1)")
        else:
            hold_signals += 1
    
    # 9. Stochastic - CHỈ TÍNH NẾU CÓ
    if available['STOCH']:
        stoch_k = ind['Stoch_K']
        stoch_d = ind['Stoch_D'] if available['Stoch_D'] else stoch_k
        
        if stoch_k < 20:
            buy_signals += 2
            signal_details.append(f"Stoch={stoch_k:.0f} quá bán (+2)")
        elif stoch_k < 40:
            buy_signals += 1
        elif stoch_k > 80:
            sell_signals += 2
            signal_details.append(f"Stoch={stoch_k:.0f} quá mua (-2)")
        elif stoch_k > 60:
            sell_signals += 1
        else:
            hold_signals += 1
        
        # K/D crossover
        if available['Stoch_D']:
            if stoch_k > stoch_d and stoch_k < 30:
                buy_signals += 1
                signal_details.append("Stoch K>D vùng thấp (+1)")
            elif stoch_k < stoch_d and stoch_k > 70:
                sell_signals += 1
                signal_details.append("Stoch K<D vùng cao (-1)")
    
    # 10. StochRSI - CHỈ TÍNH NẾU CÓ
    if available['STOCHRSI']:
        stoch_rsi = ind['StochRSI']
        if stoch_rsi < 20:
            buy_signals += 2
            signal_details.append(f"StochRSI={stoch_rsi:.0f} quá bán (+2)")
        elif stoch_rsi < 40:
            buy_signals += 1
        elif stoch_rsi > 80:
            sell_signals += 2
            signal_details.append(f"StochRSI={stoch_rsi:.0f} quá mua (-2)")
        elif stoch_rsi > 60:
            sell_signals += 1
        else:
            hold_signals += 1
    
    # 11. ROC - CHỈ TÍNH NẾU CÓ
    if available['ROC']:
        roc = ind['ROC']
        if roc > 5:
            buy_signals += 1
            signal_details.append(f"ROC={roc:.1f}% momentum mạnh (+1)")
        elif roc > 0:
            buy_signals += 1
        elif roc < -5:
            sell_signals += 1
            signal_details.append(f"ROC={roc:.1f}% momentum yếu (-1)")
        elif roc < 0:
            sell_signals += 1
        else:
            hold_signals += 1
    
    # 12. Momentum - CHỈ TÍNH NẾU CÓ
    if available['MOM']:
        mom = ind['MOM']
        if mom > 0:
            buy_signals += 1
            signal_details.append(f"Momentum={mom:.2f} dương (+1)")
        else:
            sell_signals += 1
            signal_details.append(f"Momentum={mom:.2f} âm (-1)")
    
    # ===== NHÓM 3: DAO ĐỘNG (5 chỉ báo) =====
    
    # 13. CCI - CHỈ TÍNH NẾU CÓ
    if available['CCI']:
        cci = ind['CCI']
        if cci < -200:
            buy_signals += 2
            signal_details.append(f"CCI={cci:.0f} cực kỳ quá bán (+2)")
        elif cci < -100:
            buy_signals += 1
            signal_details.append(f"CCI={cci:.0f} quá bán (+1)")
        elif cci > 200:
            sell_signals += 2
            signal_details.append(f"CCI={cci:.0f} cực kỳ quá mua (-2)")
        elif cci > 100:
            sell_signals += 1
            signal_details.append(f"CCI={cci:.0f} quá mua (-1)")
        else:
            hold_signals += 1
    
    # 14. Williams %R - CHỈ TÍNH NẾU CÓ
    if available['WILLR']:
        willr = ind['Williams_R']
        if willr < -80:
            buy_signals += 1
            signal_details.append(f"Williams %R={willr:.0f} quá bán (+1)")
        elif willr > -20:
            sell_signals += 1
            signal_details.append(f"Williams %R={willr:.0f} quá mua (-1)")
        else:
            hold_signals += 1
    
    # 15. ADX + DI - CHỈ TÍNH NẾU CÓ
    if available['ADX']:
        adx = ind['ADX']
        plus_di = ind['Plus_DI'] if available['Plus_DI'] else 25
        minus_di = ind['Minus_DI'] if available['Minus_DI'] else 25
        
        if adx > 25:
            if plus_di > minus_di:
                buy_signals += 1
                signal_details.append(f"ADX={adx:.0f} uptrend (+DI>{minus_di:.0f}) (+1)")
            else:
                sell_signals += 1
                signal_details.append(f"ADX={adx:.0f} downtrend (-DI>{plus_di:.0f}) (-1)")
        else:
            hold_signals += 1
            signal_details.append(f"ADX={adx:.0f} không trend")
    
    # 16. ATR - KHÔNG TẠO TÍN HIỆU MUA/BÁN, DÙNG CHO SIZING
    # ATR chỉ đánh giá độ biến động, không tạo tín hiệu trực tiếp
    if available['ATR']:
        atr = ind['ATR']
        atr_pct = (atr / close) * 100 if close > 0 else 2
        if atr_pct > 5:
            signal_details.append(f"⚠️ ATR={atr_pct:.1f}% biến động CAO")
        elif atr_pct < 1.5:
            signal_details.append(f"ATR={atr_pct:.1f}% biến động thấp")
    
    # 17. Bollinger Bands - CHỈ TÍNH NẾU CÓ
    if available['BB']:
        bb_upper = ind['BB_Upper'] if available['BB_Upper'] else close * 1.05
        bb_lower = ind['BB_Lower'] if available['BB_Lower'] else close * 0.95
        bb_range = bb_upper - bb_lower
        bb_pos = (close - bb_lower) / bb_range * 100 if bb_range > 0 else 50
        
        if close < bb_lower:
            buy_signals += 2
            signal_details.append(f"Dưới BB Lower (+2)")
        elif bb_pos < 20:
            buy_signals += 1
            signal_details.append(f"BB vị trí {bb_pos:.0f}% (+1)")
        elif close > bb_upper:
            sell_signals += 2
            signal_details.append(f"Trên BB Upper (-2)")
        elif bb_pos > 80:
            sell_signals += 1
            signal_details.append(f"BB vị trí {bb_pos:.0f}% (-1)")
        else:
            hold_signals += 1
    
    # ===== NHÓM 4: KHỐI LƯỢNG (6 chỉ báo) =====
    
    # 18. OBV - CHỈ TÍNH NẾU CÓ
    if available['OBV'] and len(df) > 5:
        obv_now = ind['OBV']
        obv_5d = df['OBV'].iloc[-6] if pd.notna(df['OBV'].iloc[-6]) else obv_now
        obv_change = (obv_now - obv_5d) / abs(obv_5d) * 100 if obv_5d != 0 else 0
        
        if obv_change > 10:
            buy_signals += 1
            signal_details.append(f"OBV +{obv_change:.0f}% tích lũy mạnh (+1)")
        elif obv_now > obv_5d:
            buy_signals += 1
        elif obv_change < -10:
            sell_signals += 1
            signal_details.append(f"OBV {obv_change:.0f}% phân phối mạnh (-1)")
        else:
            sell_signals += 1
    
    # 19. MFI - CHỈ TÍNH NẾU CÓ
    if available['MFI']:
        mfi = ind['MFI']
        if mfi < 20:
            buy_signals += 2
            signal_details.append(f"MFI={mfi:.0f} dòng tiền quá bán (+2)")
        elif mfi < 40:
            buy_signals += 1
        elif mfi > 80:
            sell_signals += 2
            signal_details.append(f"MFI={mfi:.0f} dòng tiền quá mua (-2)")
        elif mfi > 60:
            sell_signals += 1
        else:
            hold_signals += 1
    
    # 20. CMF - CHỈ TÍNH NẾU CÓ
    if available['CMF']:
        cmf = ind['CMF']
        if cmf > 0.1:
            buy_signals += 1
            signal_details.append(f"CMF={cmf:.2f} áp lực mua mạnh (+1)")
        elif cmf > 0.05:
            buy_signals += 1
        elif cmf < -0.1:
            sell_signals += 1
            signal_details.append(f"CMF={cmf:.2f} áp lực bán mạnh (-1)")
        elif cmf < -0.05:
            sell_signals += 1
        else:
            hold_signals += 1
    
    # 21. A/D Line - CHỈ TÍNH NẾU CÓ
    if available['AD'] and len(df) > 5:
        ad_now = ind['AD']
        ad_5d = df['AD'].iloc[-6] if pd.notna(df['AD'].iloc[-6]) else ad_now
        
        if ad_now > ad_5d:
            buy_signals += 1
            signal_details.append("A/D tích lũy (+1)")
        else:
            sell_signals += 1
            signal_details.append("A/D phân phối (-1)")
    
    # 22. VWAP - CHỈ TÍNH NẾU CÓ
    if available['VWAP']:
        vwap = ind['VWAP']
        vwap_diff = (close - vwap) / vwap * 100 if vwap > 0 else 0
        
        if close > vwap:
            buy_signals += 1
            signal_details.append(f"Giá > VWAP +{vwap_diff:.1f}% (+1)")
        else:
            sell_signals += 1
            signal_details.append(f"Giá < VWAP {vwap_diff:.1f}% (-1)")
    
    # 23. Force Index - CHỈ TÍNH NẾU CÓ
    if available['FI']:
        fi = ind['FI']
        if fi > 0:
            buy_signals += 1
            signal_details.append(f"Force Index={fi:,.0f} lực mua (+1)")
        else:
            sell_signals += 1
            signal_details.append(f"Force Index={fi:,.0f} lực bán (-1)")
        
        # Trend FI
        if len(df) >= 3:
            fi_3d = df['FI'].iloc[-4] if pd.notna(df['FI'].iloc[-4]) else fi
            if fi > fi_3d and fi > 0:
                buy_signals += 1
                signal_details.append("FI tăng (+1)")
            elif fi < fi_3d and fi < 0:
                sell_signals += 1
                signal_details.append("FI giảm (-1)")

    
    # ============================================================
    # XÁC ĐỊNH PHA THỊ TRƯỜNG
    # ============================================================
    
    phase, phase_confidence = detect_market_phase(
        df, rsi, stoch_k, ind['MACD_Hist'], close, ind['SMA_20'], ind['SMA_50']
    )
    
    # ============================================================
    # HÀNH ĐỘNG ĐANG GIỮ vs CHƯA CÓ
    # ============================================================
    
    net_signal = buy_signals - sell_signals
    
    # ĐANG GIỮ
    if sell_signals >= buy_signals + 4:
        hanh_dong_dang_giu = "BAN_MANH"
        hanh_dong_dang_giu_display = "🔴 BÁN MẠNH"
    elif sell_signals >= buy_signals + 2:
        hanh_dong_dang_giu = "BAN"
        hanh_dong_dang_giu_display = "🔴 BÁN / CHỐT LỜI"
    elif sell_signals > buy_signals:
        hanh_dong_dang_giu = "CAN_NHAC_BAN"
        hanh_dong_dang_giu_display = "🟡 CÂN NHẮC BÁN"
    elif buy_signals >= sell_signals + 4:
        hanh_dong_dang_giu = "MUA_THEM"
        hanh_dong_dang_giu_display = "🟢 MUA THÊM"
    elif buy_signals >= sell_signals + 2:
        hanh_dong_dang_giu = "GIU_MUA_THEM"
        hanh_dong_dang_giu_display = "🟢 GIỮ + MUA THÊM"
    else:
        hanh_dong_dang_giu = "GIU"
        hanh_dong_dang_giu_display = "⚪ GIỮ / THEO DÕI"
    
    # CHƯA CÓ
    if buy_signals >= sell_signals + 5:
        hanh_dong_chua_co = "MUA_MANH"
        hanh_dong_chua_co_display = "🟢 MUA MẠNH"
    elif buy_signals >= sell_signals + 3:
        hanh_dong_chua_co = "MUA"
        hanh_dong_chua_co_display = "🟢 MUA"
    elif buy_signals > sell_signals:
        hanh_dong_chua_co = "CAN_NHAC_MUA"
        hanh_dong_chua_co_display = "🟢 CÂN NHẮC MUA"
    elif sell_signals > buy_signals + 3:
        hanh_dong_chua_co = "TRANH"
        hanh_dong_chua_co_display = "🔴 TRÁNH / CHỜ"
    elif sell_signals > buy_signals:
        hanh_dong_chua_co = "CHO"
        hanh_dong_chua_co_display = "🟡 CHỜ TÍN HIỆU"
    else:
        hanh_dong_chua_co = "THEO_DOI"
        hanh_dong_chua_co_display = "⚪ THEO DÕI"
    
    # ============================================================
    # TÍNH GIÁ MUA/BÁN/CẮT LỖ
    # ============================================================
    
    atr_val = ind['ATR']
    high_20d = df['high'].tail(20).max()
    low_20d = df['low'].tail(20).min()
    
    support_levels = [low_20d, ind['BB_Lower'], ind['SMA_50'] * 0.98]
    support_levels = [s for s in support_levels if s > 0]
    
    resistance_levels = [high_20d, ind['BB_Upper'], ind['SMA_50'] * 1.02]
    resistance_levels = [r for r in resistance_levels if r > 0]
    
    gia_mua_tot = round(np.median(support_levels), 2) if support_levels else round(close * 0.95, 2)
    gia_ban_1 = round(np.median(resistance_levels), 2) if resistance_levels else round(close * 1.05, 2)
    gia_ban_2 = round(max(resistance_levels), 2) if resistance_levels else round(close * 1.10, 2)
    cat_lo = round(gia_mua_tot - (atr_val * 2), 2)
    
    # ============================================================
    # DỰ BÁO T0-T5, W1-W4, M1-M3
    # ============================================================
    
    bias = (final_score - 50) / 50
    
    prices = {'T0': close}
    changes = {'T0': 0}
    
    daily_change_base = (atr_val / close) * bias if close > 0 else 0.01 * bias
    daily_change_base = max(min(daily_change_base, 0.03), -0.03)
    
    # T1-T5
    cumulative = 0
    for i in range(1, 6):
        decay = 1 - (i - 1) * 0.1
        daily = daily_change_base * decay
        cumulative += daily
        prices[f'T{i}'] = round(close * (1 + cumulative), 2)
        changes[f'T{i}'] = round(cumulative * 100, 2)
    
    # W1-W4
    for i in range(1, 5):
        weeks_change = cumulative + daily_change_base * i * 3 * 0.5
        prices[f'W{i}'] = round(close * (1 + weeks_change), 2)
        changes[f'W{i}'] = round(weeks_change * 100, 2)
    
    # M1-M3
    w4_change = changes['W4'] / 100
    for i in range(1, 4):
        months_change = w4_change + daily_change_base * i * 10 * 0.3
        prices[f'M{i}'] = round(close * (1 + months_change), 2)
        changes[f'M{i}'] = round(months_change * 100, 2)
    
    # ============================================================
    # ĐỘ TIN CẬY
    # ============================================================
    
    available_indicators = len(weighted_scores)
    total_indicators = 23  # Số chỉ báo tối đa có thể tính điểm
    data_confidence_pct = (available_indicators / total_indicators) * 100
    
    if data_confidence_pct >= 80:
        data_confidence = "CAO"
    elif data_confidence_pct >= 50:
        data_confidence = "TRUNG BÌNH"
    else:
        data_confidence = "THẤP"
    
    # Rủi ro
    if phase in ['ĐÁY', 'GIẢM'] and rsi < 40:
        rui_ro_mua_duoi = 'THẤP'
    elif phase == 'TÍCH_LŨY':
        rui_ro_mua_duoi = 'TRUNG BÌNH'
    else:
        rui_ro_mua_duoi = 'CAO'
    
    change_t25 = (changes.get('T2', 0) + changes.get('T3', 0)) / 2
    if change_t25 > 1:
        rui_ro_T25 = f'THẤP (+{change_t25:.1f}%)'
        rui_ro_T25_level = 'THẤP'
    elif change_t25 > -0.5:
        rui_ro_T25 = f'TB ({change_t25:+.1f}%)'
        rui_ro_T25_level = 'TRUNG BÌNH'
    else:
        rui_ro_T25 = f'CAO ({change_t25:+.1f}%)'
        rui_ro_T25_level = 'CAO'
    
    # ============================================================
    # CHI TIẾT CHỈ BÁO
    # ============================================================
    
    indicator_details = []
    for ind_name, data in weighted_scores.items():
        indicator_details.append({
            'indicator': ind_name,
            'score': data['score'],
            'weight': data['weight'],
            'weighted_score': round(data['score'] * data['weight'], 1),
            'signal': data['signal'],
            'reason': data['reason']
        })
    
    indicator_details.sort(key=lambda x: x['weight'], reverse=True)
    
    # ============================================================
    # PHÂN TÍCH CHI TIẾT TÍN HIỆU (MỚI)
    # ============================================================
    
    signal_analysis = analyze_indicator_signals(ind, available, close, df)
    action_analysis = generate_action_analysis(
        phase, 
        signal_analysis['buy_points'], 
        signal_analysis['sell_points'], 
        signal_analysis['hold_points'],
        signal_analysis['conflicts'],
        ind,
        close
    )

    # ============================================================
    # PHÂN TÍCH NÂNG CAO (MỚI)
    # ============================================================
    
    advanced_analysis = advanced_indicator_analysis(df, ind, available, close)
    
    # Cập nhật điểm số với phân tích nâng cao
    advanced_score_adjustment = advanced_analysis['strength_score'] * 2  # Mỗi điểm = 2% final score
    final_score = max(0, min(100, final_score + advanced_score_adjustment))

    # Cập nhật hành động từ phân tích mới
    hanh_dong_chua_co = action_analysis['action_chua_co']
    hanh_dong_chua_co_display = action_analysis['action_chua_co_display']
    hanh_dong_dang_giu = action_analysis['action_dang_giu']
    hanh_dong_dang_giu_display = action_analysis['action_dang_giu_display']

    # ============================================================
    # DỰ BÁO 12 PHƯƠNG PHÁP CHO 26 CHỈ BÁO (MỚI)
    # ============================================================
    
    try:
        # Chạy dự báo 12 phương pháp
        indicator_forecast_results = forecast_all_26_indicators(df, list(weighted_scores.keys()))
        
        # Tính điểm tổng hợp theo ngày
        daily_composite = calculate_daily_composite_score(indicator_forecast_results, final_score)
        
        # Cập nhật prices và changes từ kết quả mới
        new_forecasts = indicator_forecast_results.get('combined_forecast', {}).get('forecasts', {})
        
        # Nếu có dự báo mới, kết hợp với dự báo cũ (trọng số 50-50)
        for key in ['T1', 'T2', 'T3', 'T4', 'T5']:
            if key in new_forecasts and new_forecasts[key]:
                old_price = prices.get(key, close)
                new_indicator_value = new_forecasts[key]
                
                # Chuyển đổi giá trị chỉ báo tổng hợp thành giá
                # Sử dụng bias từ chỉ báo để điều chỉnh giá
                indicator_bias = (new_indicator_value - 50) / 50  # -1 to 1
                price_adjustment = atr * indicator_bias * int(key[1])  # T1=1 ATR, T5=5 ATR
                
                # Kết hợp giá cũ và mới
                combined_price = old_price * 0.6 + (close + price_adjustment) * 0.4
                prices[key] = round(combined_price, 2)
                changes[key] = round((combined_price - close) / close * 100, 2)
        
    except Exception as e:
        print(f"Lỗi dự báo 12 phương pháp cho {symbol}: {e}")
        indicator_forecast_results = None
        daily_composite = None
    
    # ============================================================
    # TRẢ VỀ KẾT QUẢ
    # ============================================================

    return {
        'symbol': symbol,
        'gia_hien_tai': close,
        'prices': prices,
        'changes': changes,
        'pha_hien_tai': phase,
        'do_tin_cay_pha': phase_confidence,
        
        # Hành động
        'hanh_dong_chua_co': hanh_dong_chua_co,
        'hanh_dong_chua_co_display': hanh_dong_chua_co_display,
        'hanh_dong_dang_giu': hanh_dong_dang_giu,
        'hanh_dong_dang_giu_display': hanh_dong_dang_giu_display,
        
        # Lý do và chiến lược
        'action_reason_chua_co': action_analysis['action_reason_chua_co'],
        'action_reason_dang_giu': action_analysis['action_reason_dang_giu'],
        'strategy_chua_co': action_analysis['strategy_chua_co'],
        'strategy_dang_giu': action_analysis['strategy_dang_giu'],
        
        # Mức giá
        'gia_mua_tot': gia_mua_tot,
        'gia_ban_muc_tieu_1': gia_ban_1,
        'gia_ban_muc_tieu_2': gia_ban_2,
        'cat_lo': cat_lo,
        'support': support_levels,
        'resistance': resistance_levels,
        
        # Điểm số
        'final_score': round(final_score, 1),
        'buy_signals': signal_analysis['buy_points'],
        'sell_signals': signal_analysis['sell_points'],
        'hold_signals': signal_analysis['hold_points'],
        
        # Chi tiết tín hiệu (MỚI)
        'signals_table': signal_analysis['signals_table'],
        'buy_group': signal_analysis['buy_group'],
        'sell_group': signal_analysis['sell_group'],
        'hold_group': signal_analysis['hold_group'],
        'conflicts': signal_analysis['conflicts'],
        
        # Phân tích hành động (MỚI)
        'action_factors': action_analysis['factors'],
        'action_analysis': action_analysis['action_analysis'],
        'action_warnings': action_analysis['warnings'],
        'action_summary': action_analysis['summary'],
        
        # Phân tích nâng cao
        'advanced_analysis': advanced_analysis,
        'divergences': advanced_analysis['divergences'],
        'indicator_trends': advanced_analysis['indicator_trends'],
        'crossover_timing': advanced_analysis['crossover_timing'],
        'bb_analysis': advanced_analysis['bb_analysis'],
        'pattern_signals': advanced_analysis['pattern_signals'],
        'confluence_zones': advanced_analysis['confluence_zones'],

        # Giữ nguyên các trường cũ
        'signal_details': signal_details,
        'indicator_details': indicator_details,
        'weighted_scores': weighted_scores,
        'rui_ro_mua_duoi': rui_ro_mua_duoi,
        'rui_ro_T25': rui_ro_T25,
        'rui_ro_T25_level': rui_ro_T25_level,
        'indicators': {
            'rsi': round(rsi, 1),
            'stoch_k': round(stoch_k, 1),
            'macd_hist': round(ind['MACD_Hist'], 3),
            'mfi': round(mfi, 1),
            'adx': round(adx, 1),
            'cci': round(cci, 1),
        },
        'data_confidence': data_confidence,
        'available_indicators': available_indicators,
        'total_indicators': total_indicators

    }

# ============================================================
# HÀM XUẤT FILE (GIỮ NGUYÊN NHƯ CODE GỐC)
# ============================================================

def create_output_dir():
    """Tạo thư mục output theo ngày"""
    date_str = datetime.now().strftime('%Y%m%d')
    output_dir = os.path.join(BASE_OUTPUT_DIR, date_str)
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(os.path.join(output_dir, "Vietstock"), exist_ok=True)
    os.makedirs(os.path.join(output_dir, "MetaStock"), exist_ok=True)
    return output_dir


def export_excel(state):
    """Xuất file Excel đầy đủ"""
    if not state or 'forecasts' not in state or not state['forecasts']:
        return None, "❌ Chưa có dữ liệu. Chạy phân tích trước!"
    
    try:
        output_dir = create_output_dir()
        ts = datetime.now().strftime('%Y%m%d_%H%M')
        excel_path = os.path.join(output_dir, f"BaoCao_{ts}.xlsx")
        
        forecasts = state['forecasts']
        all_data = state.get('all_data', {})
        df_market = pd.DataFrame(state.get('df_market', {}))
        
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            # Sheet Tổng hợp
            summary_data = []
            for f in forecasts:
                summary_data.append({
                    'Mã': f['symbol'],
                    'Giá': f['gia_hien_tai'],
                    'Pha': f['pha_hien_tai'],
                    'Điểm': f['final_score'],
                    'Nếu CHƯA CÓ': f['hanh_dong_chua_co_display'],
                    'Nếu ĐANG GIỮ': f['hanh_dong_dang_giu_display'],
                    'Tín hiệu Mua': f['buy_signals'],
                    'Tín hiệu Bán': f['sell_signals'],
                    'RSI': f['indicators']['rsi'],
                    'Stoch_K': f['indicators']['stoch_k'],
                    'RR Mua đuổi': f['rui_ro_mua_duoi'],
                    'Độ tin cậy': f['data_confidence'],
                })
            pd.DataFrame(summary_data).to_excel(writer, sheet_name='Tổng hợp', index=False)
            
            # Sheet Dự báo T0-T5
            short_data = []
            for f in forecasts:
                short_data.append({
                    'Mã': f['symbol'], 
                    'T0': f['prices']['T0'],
                    'T1': f['prices']['T1'], '%T1': f['changes']['T1'],
                    'T2': f['prices']['T2'], '%T2': f['changes']['T2'],
                    'T3': f['prices']['T3'], '%T3': f['changes']['T3'],
                    'T4': f['prices']['T4'], '%T4': f['changes']['T4'],
                    'T5': f['prices']['T5'], '%T5': f['changes']['T5'],
                })
            pd.DataFrame(short_data).to_excel(writer, sheet_name='Dự báo T0-T5', index=False)
            
            # Sheet W1-W4
            mid_data = []
            for f in forecasts:
                mid_data.append({
                    'Mã': f['symbol'], 'Giá': f['gia_hien_tai'],
                    'W1': f['prices']['W1'], '%W1': f['changes']['W1'],
                    'W2': f['prices']['W2'], '%W2': f['changes']['W2'],
                    'W3': f['prices']['W3'], '%W3': f['changes']['W3'],
                    'W4': f['prices']['W4'], '%W4': f['changes']['W4'],
                })
            pd.DataFrame(mid_data).to_excel(writer, sheet_name='Dự báo W1-W4', index=False)
            
            # Sheet M1-M3
            long_data = []
            for f in forecasts:
                long_data.append({
                    'Mã': f['symbol'], 'Giá': f['gia_hien_tai'],
                    'M1': f['prices']['M1'], '%M1': f['changes']['M1'],
                    'M2': f['prices']['M2'], '%M2': f['changes']['M2'],
                    'M3': f['prices']['M3'], '%M3': f['changes']['M3'],
                })
            pd.DataFrame(long_data).to_excel(writer, sheet_name='Dự báo M1-M3', index=False)
            
            # Sheet Khuyến nghị
            rec_data = []
            for f in forecasts:
                rec_data.append({
                    'Mã': f['symbol'], 'Giá': f['gia_hien_tai'], 'Pha': f['pha_hien_tai'],
                    'Điểm': f['final_score'],
                    'Nếu CHƯA CÓ': f['hanh_dong_chua_co_display'],
                    'Nếu ĐANG GIỮ': f['hanh_dong_dang_giu_display'],
                    'Giá mua tốt': f['gia_mua_tot'],
                    'Mục tiêu 1': f['gia_ban_muc_tieu_1'],
                    'Mục tiêu 2': f['gia_ban_muc_tieu_2'],
                    'Cắt lỗ': f['cat_lo'],
                    'RR Mua đuổi': f['rui_ro_mua_duoi'],
                    'RR T+2.5': f['rui_ro_T25'],
                })
            pd.DataFrame(rec_data).to_excel(writer, sheet_name='Khuyến nghị', index=False)
            
            # Sheet Chi tiết chỉ báo
            indicator_data = []
            for f in forecasts:
                for detail in f.get('indicator_details', []):
                    indicator_data.append({
                        'Mã': f['symbol'],
                        'Chỉ báo': detail['indicator'],
                        'Điểm': detail['score'],
                        'Trọng số': detail['weight'],
                        'Điểm×Trọng số': detail['weighted_score'],
                        'Tín hiệu': detail['signal'],
                        'Giải thích': detail['reason'],
                    })
            if indicator_data:
                pd.DataFrame(indicator_data).to_excel(writer, sheet_name='Chi tiết chỉ báo', index=False)
            
            # Sheet Thị trường
            if len(df_market) > 0:
                df_market.to_excel(writer, sheet_name='Thị trường', index=False)
            
            # Sheet từng mã
            for sym, df_sym in all_data.items():
                if df_sym is not None and len(df_sym) > 0:
                    df_exp = df_sym.sort_values('time', ascending=False).copy()
                    df_exp['Ngày'] = df_exp['time'].dt.strftime('%d/%m/%Y')
                    cols = ['Ngày', 'open', 'high', 'low', 'close', 'volume']
                    for c in df_exp.columns:
                        if c not in cols + ['time']:
                            cols.append(c)
                    cols_exist = [c for c in cols if c in df_exp.columns]
                    df_exp[cols_exist].to_excel(writer, sheet_name=sym[:31], index=False)
        
        return excel_path, f"✅ Đã xuất: {excel_path}"
    except Exception as e:
        return None, f"❌ Lỗi: {str(e)}"


def export_word(state):
    """Xuất file Word đầy đủ"""
    if not state or 'forecasts' not in state or not state['forecasts']:
        return None, "❌ Chưa có dữ liệu. Chạy phân tích trước!"
    
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        output_dir = create_output_dir()
        ts = datetime.now().strftime('%Y%m%d_%H%M')
        word_path = os.path.join(output_dir, f"BaoCao_{ts}.docx")
        
        forecasts = state['forecasts']
        df_market = pd.DataFrame(state.get('df_market', {}))
        start_date = state.get('start_date', '')
        end_date = state.get('end_date', '')
        
        doc = Document()
        
        title = doc.add_heading('BÁO CÁO PHÂN TÍCH KỸ THUẬT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Ngày: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        doc.add_paragraph(f'Dữ liệu: {start_date} → {end_date}')
        doc.add_paragraph(f'Số mã: {len(forecasts)}')
        
        # I. Thị trường
        doc.add_heading('I. CHỈ SỐ THỊ TRƯỜNG', level=1)
        if len(df_market) > 0:
            table = doc.add_table(rows=1, cols=len(df_market.columns))
            table.style = 'Table Grid'
            for i, col in enumerate(df_market.columns):
                table.rows[0].cells[i].text = str(col)
            for _, row in df_market.iterrows():
                row_cells = table.add_row().cells
                for i, col in enumerate(df_market.columns):
                    val = row[col]
                    row_cells[i].text = f'{val:.2f}' if isinstance(val, float) and pd.notna(val) else str(val) if pd.notna(val) else 'N/A'
        
        # II. Tóm tắt
        doc.add_heading('II. TÓM TẮT HÀNH ĐỘNG', level=1)
        
        mua = [f for f in forecasts if 'MUA' in f['hanh_dong_chua_co']]
        cho = [f for f in forecasts if 'CHO' in f['hanh_dong_chua_co'] or 'THEO' in f['hanh_dong_chua_co']]
        tranh = [f for f in forecasts if 'TRANH' in f['hanh_dong_chua_co']]
        
        doc.add_heading('A. Dành cho người CHƯA CÓ:', level=2)
        if mua:
            doc.add_paragraph(f"🟢 MUA: {', '.join([f['symbol'] for f in mua])}")
        if cho:
            doc.add_paragraph(f"🟡 CHỜ: {', '.join([f['symbol'] for f in cho])}")
        if tranh:
            doc.add_paragraph(f"🔴 TRÁNH: {', '.join([f['symbol'] for f in tranh])}")
        
        ban = [f for f in forecasts if 'BAN' in f['hanh_dong_dang_giu']]
        giu = [f for f in forecasts if 'GIU' in f['hanh_dong_dang_giu'] and 'BAN' not in f['hanh_dong_dang_giu']]
        
        doc.add_heading('B. Dành cho người ĐANG GIỮ:', level=2)
        if ban:
            doc.add_paragraph(f"🔴 BÁN: {', '.join([f['symbol'] for f in ban])}")
        if giu:
            doc.add_paragraph(f"🟢 GIỮ: {', '.join([f['symbol'] for f in giu])}")
        
        # III-V. Dự báo
        doc.add_heading('III. DỰ BÁO NGẮN HẠN (T0-T5)', level=1)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        for i, h in enumerate(['Mã', 'T0', 'T1', 'T3', 'T5', '%T5', 'CHƯA CÓ', 'ĐANG GIỮ']):
            table.rows[0].cells[i].text = h
        for f in forecasts:
            row = table.add_row().cells
            row[0].text = f['symbol']
            row[1].text = f"{f['prices']['T0']:,.0f}"
            row[2].text = f"{f['prices']['T1']:,.0f}"
            row[3].text = f"{f['prices']['T3']:,.0f}"
            row[4].text = f"{f['prices']['T5']:,.0f}"
            row[5].text = f"{f['changes']['T5']:+.1f}%"
            row[6].text = f['hanh_dong_chua_co_display']
            row[7].text = f['hanh_dong_dang_giu_display']
        
        # VI. Khuyến nghị chi tiết
        doc.add_heading('VI. KHUYẾN NGHỊ CHI TIẾT', level=1)
        for f in forecasts:
            doc.add_heading(f"{f['symbol']} - {f['gia_hien_tai']:,.0f} ({f['pha_hien_tai']})", level=2)
            
            # Điểm và độ tin cậy
            doc.add_paragraph(f"📍 PHA {f['pha_hien_tai']}: Cổ phiếu đang ở vùng {f['pha_hien_tai'].lower()}")
            doc.add_paragraph(f"💯 ĐIỂM SỐ: {f['final_score']}/100 | Độ tin cậy: {f['data_confidence']} ({f['available_indicators']}/{f['total_indicators']} chỉ báo)")
            
            # Yếu tố chính
            doc.add_paragraph("🔑 YẾU TỐ CHÍNH:")
            for factor in f.get('action_factors', [])[:4]:
                doc.add_paragraph(f"   • {factor}")
            
            # Phân tích hành động
            doc.add_paragraph("📋 PHÂN TÍCH HÀNH ĐỘNG:")
            for analysis in f.get('action_analysis', []):
                doc.add_paragraph(f"   {analysis}")
            
            # Cảnh báo
            if f.get('action_warnings'):
                doc.add_paragraph("⚠️ CẢNH BÁO:")
                for warning in f['action_warnings']:
                    doc.add_paragraph(f"   {warning}")
            
            # Hành động đề xuất
            doc.add_paragraph(f"👤 CHƯA CÓ: {f['hanh_dong_chua_co_display']}")
            doc.add_paragraph(f"   └── Lý do: {f.get('action_reason_chua_co', '')}")
            doc.add_paragraph(f"   └── Chiến lược: {f.get('strategy_chua_co', '')}")
            
            doc.add_paragraph(f"👤 ĐANG GIỮ: {f['hanh_dong_dang_giu_display']}")
            doc.add_paragraph(f"   └── Lý do: {f.get('action_reason_dang_giu', '')}")
            doc.add_paragraph(f"   └── Chiến lược: {f.get('strategy_dang_giu', '')}")
            
            # Mức giá
            doc.add_paragraph(f"💰 MỨC GIÁ:")
            doc.add_paragraph(f"   Mua tốt: {f['gia_mua_tot']:,.0f} | Bán T1: {f['gia_ban_muc_tieu_1']:,.0f} | Bán T2: {f['gia_ban_muc_tieu_2']:,.0f} | Cắt lỗ: {f['cat_lo']:,.0f}")
            
            # Rủi ro
            doc.add_paragraph(f"⚠️ RỦI RO: Mua đuổi={f['rui_ro_mua_duoi']} | T+2.5={f['rui_ro_T25']}")
            
            # Phân tích tín hiệu chi tiết
            doc.add_paragraph("📊 PHÂN TÍCH TÍN HIỆU:")
            
            # Nhóm MUA
            if f.get('buy_group'):
                doc.add_paragraph(f"Nhóm MUA ({f['buy_signals']} điểm):")
                for signal in f['buy_group'][:5]:
                    doc.add_paragraph(f"   • {signal}")
            
            # Nhóm BÁN
            if f.get('sell_group'):
                doc.add_paragraph(f"Nhóm BÁN ({f['sell_signals']} điểm):")
                for signal in f['sell_group'][:5]:
                    doc.add_paragraph(f"   • {signal}")
            
            # Nhóm GIỮ
            if f.get('hold_group'):
                doc.add_paragraph(f"Nhóm GIỮ ({f['hold_signals']} điểm):")
                for signal in f['hold_group'][:3]:
                    doc.add_paragraph(f"   • {signal}")
            
            # Kết luận
            doc.add_paragraph(f"📈 KẾT LUẬN: {f.get('action_summary', '')}")
            
            # Xung đột
            if f.get('conflicts'):
                doc.add_paragraph("⚠️ XUNG ĐỘT PHÁT HIỆN:")
                for conflict in f['conflicts']:
                    doc.add_paragraph(f"   • {conflict['type']}: {conflict['description']}")
                    doc.add_paragraph(f"     → {conflict['interpretation']}")
                    doc.add_paragraph(f"     → Đề xuất: {conflict['suggestion']}")
            
            doc.add_paragraph("---")
        
        # VII. Dự báo T0-T5 chi tiết
        doc.add_heading('VII. DỰ BÁO NGẮN HẠN (T0-T5)', level=1)
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        for i, h in enumerate(['Mã', 'T0', 'T1', 'T2', 'T3', 'T4', 'T5', 'CHƯA CÓ', 'ĐANG GIỮ']):
            table.rows[0].cells[i].text = h
        for f in forecasts:
            row = table.add_row().cells
            row[0].text = f['symbol']
            row[1].text = f"{f['prices']['T0']:,.0f}"
            row[2].text = f"{f['prices']['T1']:,.0f} ({f['changes']['T1']:+.1f}%)"
            row[3].text = f"{f['prices']['T2']:,.0f} ({f['changes']['T2']:+.1f}%)"
            row[4].text = f"{f['prices']['T3']:,.0f} ({f['changes']['T3']:+.1f}%)"
            row[5].text = f"{f['prices']['T4']:,.0f} ({f['changes']['T4']:+.1f}%)"
            row[6].text = f"{f['prices']['T5']:,.0f} ({f['changes']['T5']:+.1f}%)"
            row[7].text = f['hanh_dong_chua_co_display']
            row[8].text = f['hanh_dong_dang_giu_display']
        
        # VIII. Dự báo W1-W4
        doc.add_heading('VIII. DỰ BÁO TRUNG HẠN (W1-W4)', level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        for i, h in enumerate(['Mã', 'Giá', 'W1', 'W2', 'W3', 'W4']):
            table.rows[0].cells[i].text = h
        for f in forecasts:
            row = table.add_row().cells
            row[0].text = f['symbol']
            row[1].text = f"{f['gia_hien_tai']:,.0f}"
            row[2].text = f"{f['prices']['W1']:,.0f} ({f['changes']['W1']:+.1f}%)"
            row[3].text = f"{f['prices']['W2']:,.0f} ({f['changes']['W2']:+.1f}%)"
            row[4].text = f"{f['prices']['W3']:,.0f} ({f['changes']['W3']:+.1f}%)"
            row[5].text = f"{f['prices']['W4']:,.0f} ({f['changes']['W4']:+.1f}%)"
        
        # IX. Dự báo M1-M3
        doc.add_heading('IX. DỰ BÁO DÀI HẠN (M1-M3)', level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        for i, h in enumerate(['Mã', 'Giá', 'M1', 'M2', 'M3']):
            table.rows[0].cells[i].text = h
        for f in forecasts:
            row = table.add_row().cells
            row[0].text = f['symbol']
            row[1].text = f"{f['gia_hien_tai']:,.0f}"
            row[2].text = f"{f['prices']['M1']:,.0f} ({f['changes']['M1']:+.1f}%)"
            row[3].text = f"{f['prices']['M2']:,.0f} ({f['changes']['M2']:+.1f}%)"
            row[4].text = f"{f['prices']['M3']:,.0f} ({f['changes']['M3']:+.1f}%)"
        
        # X. Hướng dẫn đọc báo cáo
        doc.add_heading('X. HƯỚNG DẪN ĐỌC BÁO CÁO', level=1)
        
        doc.add_paragraph("📌 PHÂN BIỆT TÍN HIỆU vs HÀNH ĐỘNG:")
        doc.add_paragraph("• TÍN HIỆU: Kết quả phân tích từ từng chỉ báo riêng lẻ")
        doc.add_paragraph("• HÀNH ĐỘNG: Khuyến nghị cuối cùng dựa trên tổng hợp nhiều tín hiệu")
        
        doc.add_paragraph("")
        doc.add_paragraph("📌 Ý NGHĨA TÍN HIỆU:")
        doc.add_paragraph("• 📉 QUÁ BÁN = GIÁ RẺ / CƠ HỘI MUA")
        doc.add_paragraph("• 📈 QUÁ MUA = GIÁ CAO / CẨN THẬN")
        doc.add_paragraph("• ⚪ TRUNG LẬP = Chờ tín hiệu rõ hơn")
        
        doc.add_paragraph("")
        doc.add_paragraph("📌 Ý NGHĨA HÀNH ĐỘNG:")
        doc.add_paragraph("• 🟢 MUA MẠNH: Đa số chỉ báo đồng thuận tích cực")
        doc.add_paragraph("• 🟢 MUA: Tín hiệu tích cực, có thể mua")
        doc.add_paragraph("• 🟢 CÂN NHẮC MUA: Có tín hiệu tích cực nhưng chưa đủ mạnh")
        doc.add_paragraph("• ⚪ THEO DÕI: Xung đột tín hiệu, chờ rõ hơn")
        doc.add_paragraph("• 🟡 CHỜ: Tín hiệu tiêu cực nhẹ")
        doc.add_paragraph("• 🔴 TRÁNH: Tín hiệu tiêu cực, không nên vào")
        doc.add_paragraph("• 🔴 BÁN: Nên thoát vị thế")
        doc.save(word_path)
        return word_path, f"✅ Đã xuất: {word_path}"
    
    except Exception as e:
        return None, f"❌ Lỗi: {str(e)}"

def export_vietstock_metastock(state):
    """Xuất Vietstock và MetaStock"""
    if not state or 'all_data' not in state or not state['all_data']:
        return None, "❌ Chưa có dữ liệu."
    
    try:
        output_dir = create_output_dir()
        ts = datetime.now().strftime('%Y%m%d_%H%M')
        vs_dir = os.path.join(output_dir, "Vietstock")
        ms_dir = os.path.join(output_dir, "MetaStock")
        
        count = 0
        for sym, df_sym in state['all_data'].items():
            if df_sym is None or len(df_sym) == 0:
                continue
            
            df_sorted = df_sym.sort_values('time', ascending=False).copy()
            
            # Vietstock
            vs_df = pd.DataFrame()
            vs_df['Ngày'] = df_sorted['time'].dt.strftime('%d/%m/%Y')
            vs_df['Mã'] = sym
            for col in ['open', 'high', 'low', 'close', 'volume']:
                if col in df_sorted.columns:
                    vs_df[col.capitalize()] = df_sorted[col]
            vs_df.to_excel(os.path.join(vs_dir, f"{sym}_{ts}.xlsx"), index=False)
            
            # MetaStock
            ms_df = pd.DataFrame({
                'DATE': df_sorted['time'].dt.strftime('%Y%m%d'),
                'TICKER': sym,
                'OPEN': df_sorted['open'],
                'HIGH': df_sorted['high'],
                'LOW': df_sorted['low'],
                'CLOSE': df_sorted['close'],
                'VOLUME': df_sorted['volume'].astype(int),
            })
            ms_df.to_csv(os.path.join(ms_dir, f"{sym}_{ts}.csv"), index=False)
            count += 1
        
        return output_dir, f"✅ Đã xuất {count} mã"
    except Exception as e:
        return None, f"❌ Lỗi: {str(e)}"


def export_zip(state):
    """Xuất ZIP"""
    if not state or 'forecasts' not in state or not state['forecasts']:
        return None, "❌ Chưa có dữ liệu."
    
    try:
        excel_path, _ = export_excel(state)
        word_path, _ = export_word(state)
        export_vietstock_metastock(state)
        
        output_dir = create_output_dir()
        ts = datetime.now().strftime('%Y%m%d_%H%M')
        zip_path = os.path.join(output_dir, f"BaoCao_Full_{ts}.zip")
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            if excel_path and os.path.exists(excel_path):
                zf.write(excel_path, os.path.basename(excel_path))
            if word_path and os.path.exists(word_path):
                zf.write(word_path, os.path.basename(word_path))
            
            for folder in ["Vietstock", "MetaStock"]:
                folder_path = os.path.join(output_dir, folder)
                if os.path.exists(folder_path):
                    for file in os.listdir(folder_path):
                        zf.write(os.path.join(folder_path, file), f"{folder}/{file}")
        
        return zip_path, f"✅ Đã tạo: {zip_path}"
    except Exception as e:
        return None, f"❌ Lỗi: {str(e)}"


# ============================================================
# GIAO DIỆN GRADIO
# ============================================================
# ============================================================
# PHẦN 5: HIỂN THỊ KẾT QUẢ DỰ BÁO 12 PHƯƠNG PHÁP TRÊN GRADIO
# ============================================================

def create_interface():
    """Tạo giao diện Gradio - CẬP NHẬT VỚI TAB DỰ BÁO 12 PHƯƠNG PHÁP"""
    
    with gr.Blocks(title="Phân tích CK Việt Nam", theme=gr.themes.Soft()) as demo:
        
        gr.Markdown("""
        # 🚀 HỆ THỐNG PHÂN TÍCH KỸ THUẬT CHỨNG KHOÁN VIỆT NAM
        **Phiên bản 3.3** - 26 chỉ báo | 12 phương pháp dự báo | Dự báo T0-T5, W1-W4, M1-M3
        
        ⚡ **MỚI:** Dự báo xu hướng bằng 12 phương pháp toán học cho từng chỉ báo
        """)
        
        analysis_state = gr.State({})
        
        with gr.Tabs():
            # TAB CẤU HÌNH (giữ nguyên như cũ)
            with gr.TabItem("⚙️ Cấu hình"):
                with gr.Row():
                    with gr.Column():
                        gr.Markdown("### 📋 MÃ CỔ PHIẾU")
                        mode_selector = gr.Radio(
                            choices=[("Quét tất cả (max 50)", "all"), ("Nhập mã tùy chọn", "custom")],
                            value="custom", label="Chế độ"
                        )
                        custom_symbols = gr.Textbox(
                            value="VNM, FPT, VIC, HPG, MWG, TCB, VCB, ACB, VPB, DGW, VCK",
                            label="Danh sách mã", lines=2
                        )
                        get_all_btn = gr.Button("📥 Lấy tất cả mã")
                    
                    with gr.Column():
                        gr.Markdown("### ⏰ THỜI GIAN")
                        time_mode = gr.Radio(
                            choices=[("Số ngày", "days"), ("Từ - Đến", "range")],
                            value="days", label="Chế độ"
                        )
                        days_back = gr.Slider(30, 720, value=90, step=5, label="Số ngày")
                        with gr.Row():
                            start_date = gr.Textbox(
                                value=(datetime.now() - timedelta(days=90)).strftime('%Y-%m-%d'),
                                label="Từ"
                            )
                            end_date = gr.Textbox(
                                value=datetime.now().strftime('%Y-%m-%d'),
                                label="Đến"
                            )
                        min_data = gr.Slider(1, 50, value=5, step=1, label="Số ngày tối thiểu")
                
                gr.Markdown("### 📈 CHỈ BÁO KỸ THUẬT (26 chỉ báo)")
                gr.Markdown("*8 chỉ báo bắt buộc tự động được bật: RSI, MACD, STOCH, BB, SMA, ATR, MFI, OBV*")
                
                with gr.Row():
                    select_all_btn = gr.Button("✅ Chọn tất cả", size="sm")
                    deselect_all_btn = gr.Button("❌ Bỏ chọn", size="sm")
                
                indicator_cbs = {}
                with gr.Row():
                    with gr.Column():
                        gr.Markdown("**📈 XU HƯỚNG**")
                        for k, v in INDICATOR_GROUPS['📈 XU HƯỚNG'].items():
                            is_required = k in REQUIRED_INDICATORS
                            default_value = True if is_required else v['default']
                            indicator_cbs[k] = gr.Checkbox(value=default_value, label=v['name'])
                    with gr.Column():
                        gr.Markdown("**⚡ ĐỘNG LƯỢNG**")
                        for k, v in INDICATOR_GROUPS['⚡ ĐỘNG LƯỢNG'].items():
                            is_required = k in REQUIRED_INDICATORS
                            default_value = True if is_required else v['default']
                            indicator_cbs[k] = gr.Checkbox(value=default_value, label=v['name'])
                    with gr.Column():
                        gr.Markdown("**🔄 DAO ĐỘNG**")
                        for k, v in INDICATOR_GROUPS['🔄 DAO ĐỘNG'].items():
                            is_required = k in REQUIRED_INDICATORS
                            default_value = True if is_required else v['default']
                            indicator_cbs[k] = gr.Checkbox(value=default_value, label=v['name'])
                    with gr.Column():
                        gr.Markdown("**📊 KHỐI LƯỢNG**")
                        for k, v in INDICATOR_GROUPS['📊 KHỐI LƯỢNG'].items():
                            is_required = k in REQUIRED_INDICATORS
                            default_value = True if is_required else v['default']
                            indicator_cbs[k] = gr.Checkbox(value=default_value, label=v['name'])

                # THÊM: Tùy chọn dự báo 12 phương pháp
                gr.Markdown("### 🔮 DỰ BÁO 12 PHƯƠNG PHÁP")
                with gr.Row():
                    enable_12_methods = gr.Checkbox(value=True, label="✅ Bật dự báo 12 phương pháp cho 26 chỉ báo")
                    show_detailed_report = gr.Checkbox(value=True, label="📊 Hiển thị báo cáo chi tiết")

                include_market = gr.Checkbox(value=True, label="📊 Bao gồm chỉ số thị trường")
                run_btn = gr.Button("🚀 BẮT ĐẦU PHÂN TÍCH", variant="primary", size="lg")
            
            # TAB KẾT QUẢ (giữ nguyên)
            with gr.TabItem("📊 Kết quả"):
                output_summary = gr.Markdown("*Nhấn 'Bắt đầu phân tích'*")
                output_market = gr.Dataframe(label="Chỉ số thị trường")
                
                gr.Markdown("### TÓM TẮT HÀNH ĐỘNG")
                output_actions = gr.Markdown()
                
                with gr.Tabs():
                    with gr.TabItem("T0-T5"):
                        output_short = gr.Dataframe(label="Dự báo ngắn hạn")
                    with gr.TabItem("W1-W4"):
                        output_mid = gr.Dataframe(label="Dự báo trung hạn")
                    with gr.TabItem("M1-M3"):
                        output_long = gr.Dataframe(label="Dự báo dài hạn")
                
                gr.Markdown("### KHUYẾN NGHỊ CHI TIẾT")
                output_rec = gr.Markdown()
            
            # ===== TAB MỚI: DỰ BÁO 12 PHƯƠNG PHÁP =====
            with gr.TabItem("🔮 Dự báo 12 PP"):
                gr.Markdown("""
                ### 📊 DỰ BÁO 12 PHƯƠNG PHÁP CHO 26 CHỈ BÁO
                
                Hệ thống sử dụng 12 phương pháp toán học để dự báo xu hướng của từng chỉ báo,
                sau đó tổng hợp thành dự báo chung có trọng số.
                """)
                
                # Chọn mã để xem chi tiết
                with gr.Row():
                    symbol_selector = gr.Dropdown(
                        choices=[],
                        label="Chọn mã cổ phiếu để xem chi tiết",
                        interactive=True
                    )
                    refresh_btn = gr.Button("🔄 Làm mới", size="sm")
                
                # Tóm tắt dự báo
                with gr.Row():
                    with gr.Column(scale=1):
                        forecast_summary = gr.Markdown("*Chọn mã để xem dự báo*")
                    with gr.Column(scale=1):
                        forecast_chart_data = gr.Dataframe(
                            label="Dự báo theo thời gian",
                            headers=["Thời gian", "Điểm", "Thay đổi", "Hướng", "Đảo chiều"]
                        )
                
                # Bảng chi tiết 12 phương pháp
                gr.Markdown("### 📋 Chi tiết 12 phương pháp")
                methods_detail = gr.Dataframe(
                    label="Kết quả từ 12 phương pháp",
                    headers=["Phương pháp", "Xu hướng", "Độ tin cậy", "Dự báo T5", "Giải thích"]
                )
                
                # Báo cáo chi tiết (Markdown)
                gr.Markdown("### 📝 Báo cáo chi tiết")
                detailed_report = gr.Markdown("*Chọn mã để xem báo cáo chi tiết*")
                
                # Bảng so sánh các chỉ báo
                gr.Markdown("### 📊 So sánh các chỉ báo")
                indicators_comparison = gr.Dataframe(
                    label="Dự báo từng chỉ báo",
                    headers=["Chỉ báo", "Giá trị", "Xu hướng", "Tin cậy", "T1", "T3", "T5"]
                )
            
            # TAB XUẤT FILE (cập nhật)
            with gr.TabItem("💾 Xuất file"):
                gr.Markdown("### Xuất báo cáo")
                with gr.Row():
                    export_excel_btn = gr.Button("📊 Excel", variant="primary")
                    export_word_btn = gr.Button("📝 Word", variant="primary")
                with gr.Row():
                    export_vs_btn = gr.Button("📈 Vietstock & MetaStock")
                    export_zip_btn = gr.Button("📦 ZIP (Tất cả)")
                
                # THÊM: Xuất báo cáo 12 phương pháp
                gr.Markdown("### 🔮 Xuất báo cáo dự báo 12 phương pháp")
                with gr.Row():
                    export_forecast_md_btn = gr.Button("📝 Markdown", variant="secondary")
                    export_forecast_excel_btn = gr.Button("📊 Excel chi tiết", variant="secondary")
                
                output_export = gr.Markdown()
                output_file = gr.File(label="📥 Tải file")
        
        # ============================================================
        # XỬ LÝ SỰ KIỆN CHO TAB DỰ BÁO 12 PHƯƠNG PHÁP
        # ============================================================
        
        def update_symbol_dropdown(state):
            """Cập nhật dropdown khi có dữ liệu mới"""
            if state and 'forecasts' in state:
                symbols = [f['symbol'] for f in state['forecasts']]
                return gr.Dropdown(choices=symbols, value=symbols[0] if symbols else None)
            return gr.Dropdown(choices=[], value=None)
        
        def show_forecast_details(symbol, state):
            """Hiển thị chi tiết dự báo cho mã được chọn"""
            if not symbol or not state or 'forecasts' not in state:
                return (
                    "*Không có dữ liệu*",
                    pd.DataFrame(),
                    pd.DataFrame(),
                    "*Không có dữ liệu*",
                    pd.DataFrame()
                )
            
            # Tìm forecast cho symbol
            forecast = None
            for f in state['forecasts']:
                if f['symbol'] == symbol:
                    forecast = f
                    break
            
            if not forecast:
                return (
                    f"*Không tìm thấy dữ liệu cho {symbol}*",
                    pd.DataFrame(),
                    pd.DataFrame(),
                    f"*Không tìm thấy dữ liệu cho {symbol}*",
                    pd.DataFrame()
                )
            
            # Lấy kết quả 12 phương pháp
            forecast_12 = forecast.get('forecast_12_methods', {})
            
            if not forecast_12 or not forecast_12.get('success'):
                return (
                    f"## {symbol}\n\n❌ Chưa chạy dự báo 12 phương pháp hoặc có lỗi",
                    pd.DataFrame(),
                    pd.DataFrame(),
                    f"## {symbol}\n\n❌ Chưa chạy dự báo 12 phương pháp",
                    pd.DataFrame()
                )
            
            indicator_forecasts = forecast_12.get('indicator_forecasts', {})
            daily_composite = forecast_12.get('daily_composite', {})
            markdown_report = forecast_12.get('markdown_report', '')
            
            # 1. Tóm tắt
            summary = indicator_forecasts.get('summary', {})
            overall_trend = summary.get('overall_trend', 'N/A')
            confidence = summary.get('overall_confidence', 0)
            
            if 'TĂNG' in overall_trend:
                trend_emoji = '🟢'
            elif 'GIẢM' in overall_trend:
                trend_emoji = '🔴'
            else:
                trend_emoji = '🟡'
            
            summary_md = f"""
## {symbol} - {forecast.get('gia_hien_tai', 0):,.0f}

### {trend_emoji} Xu hướng: **{overall_trend}**
- Độ tin cậy: **{confidence:.1f}%**
- Số chỉ báo: **{summary.get('indicators_analyzed', 0)}**
- Tỷ lệ tăng: {summary.get('trend_breakdown', {}).get('up_pct', 0):.1f}%
- Tỷ lệ giảm: {summary.get('trend_breakdown', {}).get('down_pct', 0):.1f}%
            """
            
            # 2. Bảng dự báo theo ngày (SỬA: dùng daily_results thay vì daily_scores)
            daily_results = daily_composite.get('daily_results', {})
            chart_data = []
            for key in ['T0', 'T1', 'T2', 'T3', 'T4', 'T5', 'W1', 'W2', 'W3', 'W4', 'M1', 'M2', 'M3']:
                data = daily_results.get(key, {})
                chart_data.append({
                    'Thời gian': key,
                    'Điểm': f"{data.get('score', 0):.1f}" if isinstance(data.get('score'), (int, float)) else '-',
                    'Giá': f"{data.get('price', 0):,.0f}" if isinstance(data.get('price'), (int, float)) else '-',
                    'Thay đổi': f"{data.get('change_price_pct', 0):+.1f}%" if isinstance(data.get('change_price_pct'), (int, float)) else '-',
                    'Hướng': data.get('direction', '-'),
                    'Đảo chiều': data.get('reversal_type', '-') or '-'
                })
            df_chart = pd.DataFrame(chart_data)
            
            # 3. Bảng chi tiết 12 phương pháp (lấy từ 1 chỉ báo đại diện - RSI)
            methods_data = []
            individual = indicator_forecasts.get('individual_forecasts', {})
            
            # Lấy RSI hoặc chỉ báo đầu tiên có dữ liệu
            sample_indicator = individual.get('RSI') or (list(individual.values())[0] if individual else None)
            
            if sample_indicator:
                detailed = sample_indicator.get('detailed', {})
                methods = detailed.get('methods', {})
                
                method_names = {
                    'linear': 'Linear Regression',
                    'polynomial': 'Polynomial Regression',
                    'velocity': 'First Derivative',
                    'acceleration': 'Second Derivative',
                    'peak_trough': 'Peak/Trough Detection',
                    'multi_timeframe': 'Multi-Timeframe',
                    'pattern': 'Pattern Matching',
                    'fourier': 'Fourier Transform',
                    'statistics': 'Probability & Statistics',
                    'fibonacci': 'Fibonacci Levels',
                    'logical': 'Logical Rules',
                    'ensemble': 'ML Ensemble'
                }
                
                for method_key, method_name in method_names.items():
                    method_data = methods.get(method_key, {})
                    if method_data:
                        forecasts = method_data.get('forecasts', {})
                        methods_data.append({
                            'Phương pháp': method_name,
                            'Xu hướng': '-',  # Có thể thêm logic
                            'Độ tin cậy': f"{method_data.get('confidence', 0):.0f}%",
                            'Dự báo T5': f"{forecasts.get('T5', '-'):.1f}" if isinstance(forecasts.get('T5'), (int, float)) else '-',
                            'Giải thích': (method_data.get('explanation', '')[:80] + '...') if len(method_data.get('explanation', '')) > 80 else method_data.get('explanation', '')
                        })
            
            df_methods = pd.DataFrame(methods_data) if methods_data else pd.DataFrame()
            
            # 4. Bảng so sánh các chỉ báo
            indicators_data = []
            for ind_name, ind_data in individual.items():
                final = ind_data.get('final', {})
                forecasts = final.get('forecasts', {})
                indicators_data.append({
                    'Chỉ báo': ind_name,
                    'Giá trị': f"{final.get('current_value', '-'):.2f}" if isinstance(final.get('current_value'), (int, float)) else '-',
                    'Xu hướng': final.get('trend', '-'),
                    'Tin cậy': f"{final.get('confidence', 0):.0f}%",
                    'T1': f"{forecasts.get('T1', '-'):.1f}" if isinstance(forecasts.get('T1'), (int, float)) else '-',
                    'T3': f"{forecasts.get('T3', '-'):.1f}" if isinstance(forecasts.get('T3'), (int, float)) else '-',
                    'T5': f"{forecasts.get('T5', '-'):.1f}" if isinstance(forecasts.get('T5'), (int, float)) else '-',
                })
            df_indicators = pd.DataFrame(indicators_data) if indicators_data else pd.DataFrame()
            
            return (
                summary_md,
                df_chart,
                df_methods,
                markdown_report,
                df_indicators
            )
        
        def export_forecast_markdown(state):
            """Xuất báo cáo dự báo ra file Markdown"""
            if not state or 'forecasts' not in state:
                return None, "❌ Chưa có dữ liệu"
            
            try:
                output_dir = create_output_dir()
                ts = datetime.now().strftime('%Y%m%d_%H%M')
                md_path = os.path.join(output_dir, f"DuBao_12PP_{ts}.md")
                
                content = []
                content.append("# BÁO CÁO DỰ BÁO 12 PHƯƠNG PHÁP")
                content.append(f"\nNgày: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
                content.append(f"\nSố mã: {len(state['forecasts'])}")
                content.append("\n---\n")
                
                for f in state['forecasts']:
                    forecast_12 = f.get('forecast_12_methods', {})
                    if forecast_12 and forecast_12.get('success'):
                        content.append(forecast_12.get('markdown_report', ''))
                        content.append("\n---\n")
                
                with open(md_path, 'w', encoding='utf-8') as file:
                    file.write("\n".join(content))
                
                return md_path, f"✅ Đã xuất: {md_path}"
            except Exception as e:
                return None, f"❌ Lỗi: {str(e)}"
        
        def export_forecast_excel_detail(state):
            """Xuất báo cáo dự báo ra file Excel chi tiết"""
            if not state or 'forecasts' not in state:
                return None, "❌ Chưa có dữ liệu"
            
            try:
                output_dir = create_output_dir()
                ts = datetime.now().strftime('%Y%m%d_%H%M')
                excel_path = os.path.join(output_dir, f"DuBao_12PP_{ts}.xlsx")
                
                with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                    # Sheet tổng hợp
                    summary_data = []
                    for f in state['forecasts']:
                        forecast_12 = f.get('forecast_12_methods', {})
                        if forecast_12 and forecast_12.get('success'):
                            indicator_forecasts = forecast_12.get('indicator_forecasts', {})
                            summary = indicator_forecasts.get('summary', {})
                            summary_data.append({
                                'Mã': f['symbol'],
                                'Giá': f.get('gia_hien_tai', 0),
                                'Xu hướng': summary.get('overall_trend', 'N/A'),
                                'Độ tin cậy': summary.get('overall_confidence', 0),
                                'Số chỉ báo': summary.get('indicators_analyzed', 0),
                                '% Tăng': summary.get('trend_breakdown', {}).get('up_pct', 0),
                                '% Giảm': summary.get('trend_breakdown', {}).get('down_pct', 0),
                            })
                    
                    if summary_data:
                        pd.DataFrame(summary_data).to_excel(writer, sheet_name='Tổng hợp', index=False)
                    
                    # Sheet chi tiết từng mã
                    for f in state['forecasts']:
                        forecast_12 = f.get('forecast_12_methods', {})
                        if forecast_12 and forecast_12.get('success'):
                            export_forecast_to_excel_sheet(
                                writer,
                                f['symbol'],
                                forecast_12.get('indicator_forecasts'),
                                forecast_12.get('daily_composite')
                            )
                
                return excel_path, f"✅ Đã xuất: {excel_path}"
            except Exception as e:
                return None, f"❌ Lỗi: {str(e)}"
        
        # ============================================================
        # CẬP NHẬT HÀM RUN_ANALYSIS
        # ============================================================
        
        def run_analysis(mode, symbols_text, time_mode_val, days, start_dt, end_dt, min_pts, 
                         enable_12_methods_val, show_detailed_val, include_mkt, *ind_values):
            """Hàm phân tích chính - CẬP NHẬT với 12 phương pháp"""
            
            if mode == "all":
                symbols = get_all_symbols()[:50]
            else:
                symbols = [s.strip().upper() for s in symbols_text.split(',') if s.strip()]
            
            if not symbols:
                return "❌ Nhập ít nhất 1 mã!", None, "", None, None, None, "", {}
            
            if time_mode_val == "days":
                end_parsed = datetime.now()
                start_parsed = end_parsed - timedelta(days=days)
            else:
                try:
                    start_parsed = datetime.strptime(start_dt, '%Y-%m-%d')
                    end_parsed = datetime.strptime(end_dt, '%Y-%m-%d')
                except:
                    return "❌ Định dạng ngày sai!", None, "", None, None, None, "", {}
            
            START = start_parsed.strftime('%Y-%m-%d')
            END = end_parsed.strftime('%Y-%m-%d')
            
            ind_keys = list(indicator_cbs.keys())
            selected_inds = [k for k, v in zip(ind_keys, ind_values) if v]
            if not selected_inds:
                return "❌ Chọn ít nhất 1 chỉ báo!", None, "", None, None, None, "", {}
            if 'ATR' not in selected_inds:
                selected_inds.append('ATR')
            
            df_market = pd.DataFrame()
            market_ctx = ""
            if include_mkt:
                df_market = get_market_indices()
                if len(df_market) > 0:
                    market_ctx, _ = analyze_market(df_market)
            
            all_data = {}
            forecasts = []
            
            for sym in symbols:
                try:
                    df = get_stock_data(sym, START, END)
                    if df is None or len(df) < min_pts:
                        continue
                    
                    for col in ['close', 'high', 'low', 'open', 'volume']:
                        df[col] = df[col].astype(float)
                    df['time'] = pd.to_datetime(df['time'])
                    
                    inds = calculate_indicators(df, selected_inds)
                    for col, vals in inds.items():
                        df[col] = vals
                    
                    df = df.sort_values('time', ascending=False).reset_index(drop=True)
                    all_data[sym] = df
                    
                    fc = forecast_multi_timeframe(df.copy(), sym)
                    if 'error' not in fc:
                        # THÊM: Chạy 12 phương pháp nếu được bật
                        if enable_12_methods_val:
                            # Lấy giá và ATR
                            current_price = float(df['close'].iloc[-1])
                            atr_value = float(df['ATR'].iloc[-1]) if 'ATR' in df.columns else current_price * 0.02
                            
                            forecast_12 = run_12_methods_forecast_v2(
                                df.copy(),
                                sym,
                                fc.get('weighted_scores', {}),
                                fc.get('final_score', 50),
                                current_price,
                                atr_value
                            )
                            fc['forecast_12_methods'] = forecast_12
                        
                        forecasts.append(fc)
                    
                    time.sleep(0.3)
                except Exception as e:
                    print(f"Lỗi {sym}: {e}")
            
            if not forecasts:
                return "❌ Không có dữ liệu!", None, "", None, None, None, "", {}
            
            # Output (giữ nguyên như cũ)
            summary = f"## 📊 KẾT QUẢ\n\n**Thời gian:** {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n"
            summary += f"**Dữ liệu:** {START} → {END}\n\n**Số mã:** {len(forecasts)}\n\n"
            if market_ctx:
                summary += f"**Thị trường:** {market_ctx}\n"
            
            # THÊM: Thông tin 12 phương pháp
            if enable_12_methods_val:
                success_count = sum(1 for f in forecasts if f.get('forecast_12_methods', {}).get('success'))
                summary += f"\n**Dự báo 12 phương pháp:** {success_count}/{len(forecasts)} mã thành công"
            
            mua = [f for f in forecasts if 'MUA' in f['hanh_dong_chua_co']]
            cho = [f for f in forecasts if 'CHO' in f['hanh_dong_chua_co'] or 'THEO' in f['hanh_dong_chua_co']]
            tranh = [f for f in forecasts if 'TRANH' in f['hanh_dong_chua_co']]
            ban = [f for f in forecasts if 'BAN' in f['hanh_dong_dang_giu']]
            giu = [f for f in forecasts if 'GIU' in f['hanh_dong_dang_giu'] and 'BAN' not in f['hanh_dong_dang_giu']]
            
            actions = "### 👤 CHƯA CÓ:\n"
            if mua: actions += f"🟢 **MUA:** {', '.join([f['symbol'] for f in mua])}\n\n"
            if cho: actions += f"🟡 **CHỜ:** {', '.join([f['symbol'] for f in cho])}\n\n"
            if tranh: actions += f"🔴 **TRÁNH:** {', '.join([f['symbol'] for f in tranh])}\n\n"
            actions += "### 👤 ĐANG GIỮ:\n"
            if ban: actions += f"🔴 **BÁN:** {', '.join([f['symbol'] for f in ban])}\n\n"
            if giu: actions += f"🟢 **GIỮ:** {', '.join([f['symbol'] for f in giu])}\n\n"
            
            # Bảng T0-T5
            df_short = pd.DataFrame([{
                'Mã': f['symbol'], 
                'T0': f['prices']['T0'],
                'Pha': f['pha_hien_tai'],
                'Điểm': f['final_score'],
                'CHƯA CÓ': f['hanh_dong_chua_co_display'], 
                'ĐANG GIỮ': f['hanh_dong_dang_giu_display'],
                'T1': f"{f['prices']['T1']:,.0f} ({f['changes']['T1']:+.1f}%)",
                'T2': f"{f['prices']['T2']:,.0f} ({f['changes']['T2']:+.1f}%)",
                'T3': f"{f['prices']['T3']:,.0f} ({f['changes']['T3']:+.1f}%)",
                'T4': f"{f['prices']['T4']:,.0f} ({f['changes']['T4']:+.1f}%)",
                'T5': f"{f['prices']['T5']:,.0f} ({f['changes']['T5']:+.1f}%)",
                'Tín hiệu': f"M:{f['buy_signals']} B:{f['sell_signals']}",
            } for f in forecasts])
            
            # Bảng W1-W4
            df_mid = pd.DataFrame([{
                'Mã': f['symbol'], 
                'Giá': f['gia_hien_tai'],
                'W1': f"{f['prices']['W1']:,.0f} ({f['changes']['W1']:+.1f}%)", 
                'W2': f"{f['prices']['W2']:,.0f} ({f['changes']['W2']:+.1f}%)",
                'W3': f"{f['prices']['W3']:,.0f} ({f['changes']['W3']:+.1f}%)", 
                'W4': f"{f['prices']['W4']:,.0f} ({f['changes']['W4']:+.1f}%)",
            } for f in forecasts])
            
            # Bảng M1-M3
            df_long = pd.DataFrame([{
                'Mã': f['symbol'], 
                'Giá': f['gia_hien_tai'],
                'M1': f"{f['prices']['M1']:,.0f} ({f['changes']['M1']:+.1f}%)",
                'M2': f"{f['prices']['M2']:,.0f} ({f['changes']['M2']:+.1f}%)",
                'M3': f"{f['prices']['M3']:,.0f} ({f['changes']['M3']:+.1f}%)",
            } for f in forecasts])
            
            # Khuyến nghị chi tiết (giữ nguyên như cũ, có thể thêm thông tin 12 phương pháp)
            rec = ""
            for f in forecasts:
                rec += f"### {f['symbol']} ({f['gia_hien_tai']:,.0f} - {f['pha_hien_tai']})\n"
                rec += f"**Điểm:** {f['final_score']}/100 | **Độ tin cậy:** {f['data_confidence']} ({f['available_indicators']}/{f['total_indicators']} chỉ báo)\n\n"
                
                # THÊM: Thông tin từ 12 phương pháp
                forecast_12 = f.get('forecast_12_methods', {})
                if forecast_12 and forecast_12.get('success'):
                    indicator_forecasts = forecast_12.get('indicator_forecasts', {})
                    summary_12 = indicator_forecasts.get('summary', {})
                    rec += f"**🔮 Dự báo 12 PP:** {summary_12.get('overall_trend', 'N/A')} (Tin cậy: {summary_12.get('overall_confidence', 0):.0f}%)\n"
                    rec += f"   - Tỷ lệ tăng: {summary_12.get('trend_breakdown', {}).get('up_pct', 0):.1f}%\n"
                    rec += f"   - Tỷ lệ giảm: {summary_12.get('trend_breakdown', {}).get('down_pct', 0):.1f}%\n\n"
                
                rec += f"📍 **PHA {f['pha_hien_tai']}:** Cổ phiếu đang ở vùng {f['pha_hien_tai'].lower()}\n"
                
                # ... (giữ nguyên phần còn lại)
                rec += f"**👤 CHƯA CÓ:** {f['hanh_dong_chua_co_display']}\n"
                rec += f"**👤 ĐANG GIỮ:** {f['hanh_dong_dang_giu_display']}\n\n"
                rec += f"💰 **MỨC GIÁ:** Mua tốt: {f['gia_mua_tot']:,.0f} | Bán T1: {f['gia_ban_muc_tieu_1']:,.0f} | Bán T2: {f['gia_ban_muc_tieu_2']:,.0f} | Cắt lỗ: {f['cat_lo']:,.0f}\n\n"
                rec += f"⚠️ **RỦI RO:** Mua đuổi={f['rui_ro_mua_duoi']} | T+2.5={f['rui_ro_T25']}\n\n"
                rec += "---\n\n"
            
            state = {
                'all_data': all_data, 
                'forecasts': forecasts,
                'df_market': df_market.to_dict() if len(df_market) > 0 else {},
                'start_date': START, 
                'end_date': END,
            }
            
            return summary, df_market, actions, df_short, df_mid, df_long, rec, state
        
        # ============================================================
        # KẾT NỐI SỰ KIỆN
        # ============================================================
        
        get_all_btn.click(lambda: ", ".join(get_all_symbols()[:50]), outputs=custom_symbols)
        select_all_btn.click(lambda: [True] * 26, outputs=list(indicator_cbs.values()))
        deselect_all_btn.click(lambda: [False] * 26, outputs=list(indicator_cbs.values()))
        
        # Cập nhật run_analysis với thêm 2 tham số mới
        run_btn.click(
            run_analysis,
            inputs=[mode_selector, custom_symbols, time_mode, days_back, start_date, end_date,
                    min_data, enable_12_methods, show_detailed_report, include_market, 
                    *list(indicator_cbs.values())],
            outputs=[output_summary, output_market, output_actions, output_short, output_mid, output_long, output_rec, analysis_state]
        )
        
        # Cập nhật dropdown sau khi phân tích
        run_btn.click(
            update_symbol_dropdown,
            inputs=[analysis_state],
            outputs=[symbol_selector]
        )
        
        # Hiển thị chi tiết khi chọn mã
        symbol_selector.change(
            show_forecast_details,
            inputs=[symbol_selector, analysis_state],
            outputs=[forecast_summary, forecast_chart_data, methods_detail, detailed_report, indicators_comparison]
        )
        
        refresh_btn.click(
            show_forecast_details,
            inputs=[symbol_selector, analysis_state],
            outputs=[forecast_summary, forecast_chart_data, methods_detail, detailed_report, indicators_comparison]
        )
        
        # Xuất file
        export_excel_btn.click(lambda s: export_excel(s), inputs=analysis_state, outputs=[output_file, output_export])
        export_word_btn.click(lambda s: export_word(s), inputs=analysis_state, outputs=[output_file, output_export])
        export_vs_btn.click(lambda s: (None, export_vietstock_metastock(s)[1]), inputs=analysis_state, outputs=[output_file, output_export])
        export_zip_btn.click(lambda s: export_zip(s), inputs=analysis_state, outputs=[output_file, output_export])
        
        # Xuất báo cáo 12 phương pháp
        export_forecast_md_btn.click(export_forecast_markdown, inputs=analysis_state, outputs=[output_file, output_export])
        export_forecast_excel_btn.click(export_forecast_excel_detail, inputs=analysis_state, outputs=[output_file, output_export])
    
    return demo


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("🚀 HỆ THỐNG PHÂN TÍCH KỸ THUẬT CHỨNG KHOÁN VIỆT NAM")
    print("   Phiên bản 3.2 - 26 chỉ báo có trọng số")
    print("=" * 60)
    
    os.makedirs(BASE_OUTPUT_DIR, exist_ok=True)
    
    demo = create_interface()
    demo.launch(server_name="0.0.0.0", server_port=7860, share=False)
